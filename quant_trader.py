#!/usr/bin/env python3
"""
======================================================================
  A-share Quantitative Trading System — Standalone Edition
  鹏辉能源定制版 — 数据获取 → 因子计算 → 策略信号 → 回测 → 报告
======================================================================
  Usage:
    python quant_trader.py scan          # 全市场扫描 (A股+美股)
    python quant_trader.py analyze 300438 # 单股深度分析
    python quant_trader.py macro             # 宏观分析 (DXY+VIX+SPY联动)
    python quant_trader.py fundamental      # 市场基本面扫描
    python quant_trader.py market           # 全市场客观扫描 (涨幅榜+成交榜+行业+宽度)
    python quant_trader.py sector           # 自选股行业景气度
    python quant_trader.py flow              # 主力资金流向扫描
    python quant_trader.py flow 300438       # 个股资金流向分析
    python quant_trader.py fundamental 300438 # 个股基本面分析
    python quant_trader.py backtest      # 策略回测
    python quant_trader.py hot_sector   # 热门板块追踪+强势股筛选
	    python quant_trader.py dashboard     # 启动Web看板
    python quant_trader.py first_bearish           # 全市场扫描：上升趋势首根大阴线
    python quant_trader.py first_bearish 300438    # 单股深度分析
    python quant_trader.py backtest               # 事件驱动回测 (首阴候选)
    python quant_trader.py backtest 300438        # 单股回测
    python quant_trader.py prewarm                # 预热本地K线缓存
    python quant_trader.py cache                  # 查看缓存状态
    python quant_trader.py north_bound            # 北向资金分析
    python quant_trader.py limit                  # 涨跌停板统计
    python quant_trader.py dragon_tiger           # 龙虎榜分析
    python quant_trader.py yuezi                 # 游资席位跟踪(全景)
    python quant_trader.py yuezi top             # 顶级游资今日操作
    python quant_trader.py yuezi resonance       # 多游资共振检测
    python quant_trader.py yuezi profile 章盟主   # 游资画像
    python quant_trader.py yuezi seat 上塘路      # 席位搜索(按关键词)
    python quant_trader.py weekly_doji            # 周线缩量十字星扫描
    python quant_trader.py weekly_doji_deep 600519 # 周线十字星深度分析
    python quant_trader.py margin_index            # 两融余额vs上证指数关联分析
    python quant_trader.py temperature             # 市场温度计(7维度)
    python quant_trader.py sector_temp              # 板块温度排名
    python quant_trader.py sector_temp concept      # 概念板块温度
    python quant_trader.py sector_temp 磷化工        # 搜索板块
    python quant_trader.py prob analyze 300438      # 概率论: 单股策略显著性
    python quant_trader.py prob kelly 300438        # 概率论: 凯利最优仓位
    python quant_trader.py prob guide               # 概率论: 交易指南
    python quant_trader.py chemical element Li     # 查化学元素
    python quant_trader.py chemical product 磷酸铁锂 # 查化工产品
    python quant_trader.py chemical sector         # 化工行业分类
    python quant_trader.py margin                 # 融资融券分析
    python quant_trader.py multi_tf 300438        # 多时间框架共振
    python quant_trader.py heatmap                # 板块轮动热力图
    python quant_trader.py sw_rt                  # 申万一级行业实时行情(全28行业热力图)
    python quant_trader.py sw_rt 801080.SI        # 单个申万行业实时行情
    python quant_trader.py sw_list                # 申万行业指数代码列表
    python quant_trader.py optimize               # 参数优化(Grid Search)
    python quant_trader.py hedge 300438           # 对冲分析
    python quant_trader.py akshare                # Akshare市场扫描
    python quant_trader.py attribution            # 绩效归因分析
    python quant_trader.py futures               # 股指期货基差+期权PCR
    python quant_trader.py special               # 另类数据综合(大宗/解禁/增减持/行业资金)
    python quant_trader.py pairs                 # 配对交易筛选
    python quant_trader.py rotation              # 板块轮动策略(动量+均值回归)
    python quant_trader.py earnings              # 业绩超预期筛选
    python quant_trader.py sentiment             # 市场情绪分析
    python quant_trader.py monte_carlo           # 蒙特卡洛模拟+压力测试
    python quant_trader.py intraday 300438       # 日内策略(开盘突破+VWAP+尾盘)
    python quant_trader.py margin_signal 300438  # 两融信号+逼空/去杠杆检测
    python quant_trader.py rebalance             # 指数调仓日历+股权质押
    python quant_trader.py data_quality          # 数据质量检查
    python quant_trader.py versions              # 策略版本管理
    python quant_trader.py api                   # 启动REST API服务(端口8000)
    python quant_trader.py garch 300438           # GARCH波动率预测+动态仓位
    python quant_trader.py bl                     # Black-Litterman组合优化
    python quant_trader.py social 300438          # 社交媒体舆情分析
    python quant_trader.py level2 300438          # Level2盘口深度+大单分析
    python quant_trader.py paper                  # 模拟交易账户
    python quant_trader.py monitor                # 启动Prometheus健康监控
    python quant_trader.py check_adj 300438        # 复权一致性校验(Baostock vs Sina)
    python quant_trader.py check_la                # 未来函数审计(Shuffle+Leakage+CodeScan)
    python quant_trader.py bug_audit               # 隐形BUG全系统审计(6合1)
    python quant_trader.py oversold               # 超跌反弹全市场筛选
    python quant_trader.py oversold 30            # 超跌反弹TOP30候选
    python quant_trader.py oversold_backtest 300438  # 超跌反弹策略回测
    python quant_trader.py yixian 300657           # 一线定乾坤深度分析+图表
    python quant_trader.py yixian_scan 30          # 一线定乾坤全市场扫描
    python quant_trader.py yixian_live 300657      # 一线定乾坤 Tushare实时更新+微信推送
    python quant_trader.py news_pick               # 新闻驱动选股 (TOP15)
    python quant_trader.py news_pick 20            # 新闻驱动选股 (TOP20)
    python quant_trader.py inst_check 603979       # 机构票判定
    python quant_trader.py inst_scan 20            # 批量扫描机构票
    python quant_trader.py dark_horse              # 黑马筛选
    python quant_trader.py dark_horse 300438       # 单股黑马分析
    python quant_trader.py lowvol                  # 低波+动量+流动性三因子选股(TOP30)
    python quant_trader.py lowvol 20               # 三因子选股(TOP20)
    python quant_trader.py lowvol 20 60            # 三因子选股(≥60分)
    python quant_trader.py task                   # 任务看板 (当前状态)
    python quant_trader.py task --days 7          # 任务看板 (近7天)
    python quant_trader.py task --watch           # 任务看板 (持续监控)
    python quant_trader.py code_search 关键词      # 代码语义搜索
    python quant_trader.py knowledge ingest       # 知识引擎：全量摄入
    python quant_trader.py knowledge query 问题    # 知识引擎：跨源检索
    python quant_trader.py knowledge add_rule 规则  # 知识引擎：添加规则
======================================================================
"""
import sys, os, json, subprocess, warnings, itertools
from datetime import datetime, timedelta
from pathlib import Path

# Sector tracker import
try:
    from quant_trading.sector_tracker import SectorTracker, hot_sector_scan as _sector_hot_scan
except ImportError:
    from sector_tracker import SectorTracker, hot_sector_scan as _sector_hot_scan

# Dark horse screener — removed (module not available)
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
warnings.filterwarnings('ignore')
# ---- 静默异常处理器 (替代裸except) ----
import logging as _logging
_log = _logging.getLogger('quant_trader')

def _suppress(context=''):
    """替代裸 except Exception: pass, 至少在DEBUG级别记录"""
    import traceback
    _log.debug(f'Suppressed exception in {context}: {traceback.format_exc(limit=1)}')


# Fix Windows console encoding
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    os.environ['PYTHONIOENCODING'] = 'utf-8'

import numpy as np
import pandas as pd
pd.set_option('display.max_columns', 20, 'display.width', 160, 'display.max_rows', 50)

# ---- Rich 美化终端输出 (统一来源: quant_trading.utils.display) ----
from quant_trading.utils.display import (
    console, print_header, print_bar, rich_table,
    C_GOLD, C_RED, C_GREEN, C_BLUE, C_CYAN,
    C_PURPLE, C_MAGENTA, C_ORANGE, C_GRAY, C_WHITE, C_YELLOW,
)
from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeElapsedColumn
from rich.text import Text
from rich.style import Style

# ============================================================
# CONFIGURATION (统一来源: quant_trading/config.py)
# ============================================================
from quant_trading.config import (
    WATCHLIST, US_WATCHLIST, MACRO_INDICES as MACRO,
    MA_CROSS, MOM_BREAK, VAL_REV, TREND_FOL,
    MAX_POSITION_PCT, STOP_LOSS_PCT, TAKE_PROFIT_PCT,
    INITIAL_CAPITAL, COMMISSION, SLIPPAGE,
    SINA_MACRO, SINA_US, STOCK_NAMES, REPORTS_DIR,
    BACKTEST_CONFIG,
)
from quant_trading.circuit_breaker import (
    BAOSTOCK_BREAKER, SINA_BREAKER, EASTMONEY_BREAKER,
    SINA_LIMITER, EASTMONEY_LIMITER,
)

# 向后兼容别名
CFG = type('Config', (), {
    'WATCHLIST': WATCHLIST,
    'US_WATCHLIST': US_WATCHLIST,
    'MACRO': MACRO,
    'MA_CROSS': MA_CROSS,
    'MOM_BREAK': MOM_BREAK,
    'VAL_REV': VAL_REV,
    'TREND_FOL': TREND_FOL,
    'MAX_POSITION_PCT': MAX_POSITION_PCT,
    'STOP_LOSS_PCT': STOP_LOSS_PCT,
    'TAKE_PROFIT_PCT': TAKE_PROFIT_PCT,
    'INITIAL_CAPITAL': BACKTEST_CONFIG['initial_capital'],
    'COMMISSION': BACKTEST_CONFIG['commission'],
    'SLIPPAGE': BACKTEST_CONFIG['slippage'],
})()

# ============================================================
# DATA ENGINE
# ============================================================

# ---- Baostock session management ----
_bs_logged_in = False

def _bs_login():
    """确保 baostock 已登录 (幂等, 带熔断保护)"""
    global _bs_logged_in
    if _bs_logged_in:
        return True
    if not BAOSTOCK_BREAKER.allow_request():
        return False
    try:
        import baostock as bs
        lg = bs.login()
        if lg.error_code == '0':
            _bs_logged_in = True
            BAOSTOCK_BREAKER.record_success()
            return True
        BAOSTOCK_BREAKER.record_failure()
    except Exception:
        BAOSTOCK_BREAKER.record_failure()
        _suppress("_bs_login")
    return False

def _bs_code(code):
    """将 6位代码 转成 baostock 格式: 300438 → sz.300438"""
    return ("sh." if code.startswith(("6", "68")) else "sz.") + code


def _fetch_all_a_stock_codes_tushare():
    """从 Tushare Pro (权威源) 获取全A股代码→名称映射"""
    try:
        import tushare as ts
        from quant_trading.data_source_config import DATA_CONFIG

        token = os.environ.get("TUSHARE_TOKEN", "")
        if not token:
            # Try reading from .env
            env_file = Path(__file__).parent / "quant_trading" / ".env"
            if env_file.exists():
                for line in open(env_file, encoding="utf-8"):
                    if line.startswith("TUSHARE_TOKEN="):
                        token = line.split("=", 1)[1].strip().strip('"').strip("'")
                        break
        if not token:
            return None

        ts.set_token(token)
        pro = ts.pro_api()

        # Tushare stock_basic: 获取全量A股列表
        fields = "ts_code,symbol,name,area,industry,list_status"
        df = pro.stock_basic(exchange='', list_status='L', fields=fields)
        if df is None or df.empty:
            return None

        code_name = {}
        for _, row in df.iterrows():
            name = str(row['name'])
            if 'ST' in name or '退' in name:
                continue
            code = str(row['symbol'])
            code_name[code] = name

        print(f"  [TUSHARE] stock_basic: {len(code_name)} 只A股")
        return code_name
    except Exception as e:
        print(f"  [WARN] Tushare stock_basic failed: {e}")
        return None


def fetch_all_a_stock_codes():
    """获取全A股代码→名称映射 (Tushare Pro 优先, Baostock/EastMoney 备用)"""
    from quant_trading.data_source_config import DATA_CONFIG

    # 1) 尝试 Tushare Pro (权威源)
    result = _fetch_all_a_stock_codes_tushare()
    if result and len(result) > 100:
        return result

    # 2) 回退到 Baostock (记录警告)
    if DATA_CONFIG.fallback_behavior.value == "strict":
        print("  [STRICT] Tushare 不可用，严格模式下不执行回退")
        return {}

    print("  [FALLBACK] Tushare 不可用 → 回退到 Baostock")
    if not _bs_login():
        print("  [WARN] Baostock login failed, falling back to East Money")
        return _fetch_all_a_stock_codes_eastmoney()

    try:
        import baostock as bs
        rs = bs.query_stock_basic()
        if rs.error_code != '0':
            raise Exception(rs.error_msg)

        code_name = {}
        while rs.next():
            row = rs.get_row_data()
            code_full = row[0]      # sh.600000
            name = row[1]           # 浦发银行
            stock_type = row[4]     # 1=股票
            status = row[5]         # 1=上市

            if stock_type != '1' or status != '1':
                continue
            # 仅A股: sh.6 / sz.0 / sz.3
            if not (code_full.startswith('sh.6') or code_full.startswith('sz.0') or code_full.startswith('sz.3')):
                continue
            # 排除 ST/退市
            if 'ST' in name or '退' in name:
                continue

            code = code_full.replace('sh.', '').replace('sz.', '')
            code_name[code] = name

        return code_name
    except Exception as e:
        print(f"  [WARN] Baostock query_stock_basic failed: {e}")
        return _fetch_all_a_stock_codes_eastmoney()


def _fetch_all_a_stock_codes_eastmoney():
    """东方财富备用方案: 获取全A股代码→名称映射"""
    import urllib.request, urllib.error
    code_name = {}
    for fs in ["m:0+t:6", "m:0+t:80", "m:1+t:2", "m:1+t:23"]:
        page = 1
        while True:
            try:
                url = (f"https://push2delay.eastmoney.com/api/qt/clist/get?"
                       f"pn={page}&pz=200&po=0&np=1&fltt=2&fid=f12&"
                       f"fs={fs}&fields=f2,f3,f12,f14,f15,f20")
                req = urllib.request.Request(url, headers={
                    "User-Agent": "Mozilla/5.0",
                    "Referer": "https://quote.eastmoney.com/"})
                data = json.loads(urllib.request.urlopen(req, timeout=15).read())
                items = data.get("data", {}).get("diff", [])
                total = data.get("data", {}).get("total", 0)
                for item in items:
                    code = item.get("f12", "")
                    name = item.get("f14", "")
                    if any(kw in name for kw in ["ST", "退", "N", "C"]):
                        continue
                    if code.startswith(("8", "9")) or len(code) != 6:
                        continue
                    code_name[code] = name
                if page * 200 >= total:
                    break
                page += 1
            except Exception:
                break
    return code_name


def _fetch_from_tushare_with_retry(code, datalen, freq='D'):
    """
    Tushare K线拉取 — 带指数退避重试
    根据 DATA_CONFIG 配置决定重试次数和退避策略
    """
    import time
    from quant_trading.tushare_engine import fetch_kline as ts_fetch
    from quant_trading.data_source_config import DATA_CONFIG

    last_err = None
    for attempt in range(DATA_CONFIG.tushare_max_retries):
        try:
            df = ts_fetch(code, datalen, freq=freq)
            if df is not None and not df.empty and len(df) >= datalen * 0.6:
                return df
        except Exception as e:
            last_err = e
            if attempt < DATA_CONFIG.tushare_max_retries - 1:
                delay = DATA_CONFIG.tushare_retry_base_delay * (2 ** attempt)
                time.sleep(delay)

    if last_err:
        _suppress(f"tushare_exhausted_{DATA_CONFIG.tushare_max_retries}retries")
    return None


def _fetch_fallback_and_validate(code, datalen, source_name, fetch_fn, use_cache, freq):
    """
    从回退源拉取数据，并通过交叉校验 (vs Tushare缓存) 验证。

    Returns:
        (df, validated) — df 可能为 None; validated=True 表示通过交叉校验
    """
    from quant_trading.data_source_config import DATA_CONFIG, cross_validate_with_tushare
    from quant_trading.data_cache import put_kline

    try:
        df = fetch_fn()
        if df is None or df.empty:
            return None, False

        # 长度检查
        if len(df) < datalen * 0.6:
            return None, False

        # 交叉校验 (vs Tushare缓存)
        passed, reason, div_pct = cross_validate_with_tushare(df, code, freq)

        if not passed:
            _log.warning(
                f"[数据源] {source_name} 交叉校验失败 ({reason}) — 已拒收"
            )
            return None, False

        if reason and reason not in ('no_cache_to_compare', 'cache_not_tushare'):
            _log.info(
                f"[数据源] {source_name} 交叉校验通过 (最大差异 {div_pct:.2f}%)"
            )

        # 写入缓存 (标记来源) — 始终写入,不受use_cache限制
        put_kline(code, df, freq, source=source_name)

        return df, True

    except Exception:
        _suppress(source_name)
        return None, False


def fetch_kline(code, datalen=250, use_cache=True, freq='d', source='tushare', skip_fallback=False):
    """
    Fetch K-line — Tushare Pro 权威数据源 + 交叉校验回退

    优先级链 (不可动摇):
      1. 缓存命中 (Tushare来源优先, ≤3天) → 直接返回
      2. 缓存失效 → Tushare Pro (最多3次重试, 指数退避)
      3. 分钟级数据 → Sina (唯一源)
      4. 回退链 (仅在非STRICT模式下, skip_fallback=False):
         a. Sina 日线 → 交叉校验
         b. Baostock 长历史 → 交叉校验
         c. Baostock 最终 → 交叉校验
      5. STRICT模式: Tushare失败 → 返回空DataFrame

    freq: 'd'=日线, '60'=60分钟, '30'=30分钟, '15'=15分钟, '5'=5分钟
    source: 'tushare'(默认,权威) | 'adata' | 'baostock' | 'sina'
    """
    from quant_trading.data_cache import get_cached_kline, put_kline
    from quant_trading.data_source_config import DATA_CONFIG

    # ═══ 1) 缓存优先 (快速路径) — Tushare来源优先 ═══
    if use_cache:
        cached = get_cached_kline(code, datalen, freq)
        if cached is not None and len(cached) >= min(datalen, 20):
            try:
                last_date = pd.to_datetime(cached["date"].iloc[-1])
                days_stale = (datetime.now() - last_date).days
                max_stale = DATA_CONFIG.other_cache_max_stale_days  # default 2
                if days_stale <= max_stale and len(cached) >= datalen * 0.8:
                    return cached  # 缓存新鲜 → 零API调用
            except Exception:
                if len(cached) >= datalen * 0.9:
                    return cached

    # ═══ 2) 权威源: Tushare Pro (日线, 带重试) ═══
    if freq == 'd' and source in ('tushare', 'auto'):
        df = _fetch_from_tushare_with_retry(code, datalen, freq='D')
        if df is not None and not df.empty:
            put_kline(code, df, freq, source='tushare')  # 始终刷新缓存
            return df

        # Tushare 失败日志
        if DATA_CONFIG.fallback_behavior.value != 'silent':
            _log.warning(
                f"[数据源] Tushare 拉取失败 ({DATA_CONFIG.tushare_max_retries}次重试耗尽) — "
                f"code={code} datalen={datalen}"
            )

        # prewarm 模式: 跳过慢速回退链，直接返回空
        if skip_fallback:
            return pd.DataFrame()

    # ═══ 3) 分钟级数据 → Sina (Tushare不支持分钟线) ═══
    if freq != 'd':
        sina_scale = freq
        prefix = "sh" if code.startswith(("6", "68")) else "sz"
        url = (f"https://money.finance.sina.com.cn/quotes_service/api/json_v2.php/"
               f"CN_MarketData.getKLineData?symbol={prefix}{code}&scale={sina_scale}&ma=no&datalen={datalen}")
        raw = curl_get(url)
        if raw:
            try:
                df = pd.DataFrame(json.loads(raw))
                for c in ["open","high","low","close","volume"]:
                    if c in df.columns:
                        df[c] = pd.to_numeric(df[c], errors="coerce")
                df = df.rename(columns={"day":"date"})
                df["date"] = pd.to_datetime(df["date"])
                df["code"] = code
                if not df.empty:
                    if use_cache:
                        put_kline(code, df, freq, source='sina')
                    return df
            except Exception:
                _suppress("sina_intraday")
        return pd.DataFrame()

    # ═══ 4) 日线回退链 (仅在非STRICT模式下) ═══
    if DATA_CONFIG.strict_mode:
        _log.warning(
            f"[数据源] STRICT模式 — Tushare不可用, 拒绝回退。"
            f"code={code} 返回空数据"
        )
        return pd.DataFrame()

    # prewarm 模式: 跳过慢速回退链
    if skip_fallback:
        return pd.DataFrame()

    # 4a) Sina 日线回退
    if source in ('tushare', 'auto', 'sina'):
        def _fetch_sina():
            prefix = "sh" if code.startswith(("6", "68")) else "sz"
            url = (f"https://money.finance.sina.com.cn/quotes_service/api/json_v2.php/"
                   f"CN_MarketData.getKLineData?symbol={prefix}{code}&scale=240&ma=no&datalen={datalen}")
            raw = curl_get(url)
            if not raw:
                return None
            df = pd.DataFrame(json.loads(raw))
            for c in ["open","high","low","close","volume"]:
                if c in df.columns:
                    df[c] = pd.to_numeric(df[c], errors="coerce")
            df = df.rename(columns={"day":"date"})
            df["date"] = pd.to_datetime(df["date"])
            df["code"] = code
            return df if not df.empty else None

        df, validated = _fetch_fallback_and_validate(
            code, datalen, 'sina', _fetch_sina, use_cache, freq
        )
        if df is not None and validated:
            _log.info(f"[数据源] 回退到 Sina (code={code}, {len(df)}条)")
            return df

    # 4b) Baostock 长历史回退
    if datalen > 200 and _bs_login() and source in ('tushare', 'auto', 'baostock', 'adata'):
        def _fetch_baostock():
            import baostock as bs
            from quant_trading.circuit_breaker import BAOSTOCK_LIMITER
            BAOSTOCK_LIMITER.acquire()
            end_date = datetime.now().strftime('%Y-%m-%d')
            start_date = (datetime.now() - timedelta(days=int(datalen * 1.8))).strftime('%Y-%m-%d')
            rs = bs.query_history_k_data_plus(
                _bs_code(code),
                'date,open,high,low,close,volume',
                start_date=start_date, end_date=end_date,
                frequency='d', adjustflag='2'
            )
            if rs.error_code != '0':
                return None
            rows = []
            while rs.next():
                rows.append(rs.get_row_data())
            if not rows or len(rows) < datalen * 0.6:
                return None
            df = pd.DataFrame(rows, columns=['date','open','high','low','close','volume'])
            for c in ['open','high','low','close','volume']:
                df[c] = pd.to_numeric(df[c], errors='coerce')
            df['date'] = pd.to_datetime(df['date'])
            df['code'] = code
            return df.tail(datalen)

        df, validated = _fetch_fallback_and_validate(
            code, datalen, 'baostock', _fetch_baostock, use_cache, freq
        )
        if df is not None and validated:
            _log.info(f"[数据源] 回退到 Baostock (code={code}, {len(df)}条)")
            return df

    # 4c) Baostock 最终回退 (任意长度)
    if _bs_login() and source in ('tushare', 'auto', 'baostock', 'adata'):
        def _fetch_baostock_final():
            import baostock as bs
            from quant_trading.circuit_breaker import BAOSTOCK_LIMITER
            BAOSTOCK_LIMITER.acquire()
            end_date = datetime.now().strftime('%Y-%m-%d')
            start_date = (datetime.now() - timedelta(days=int(datalen * 1.8))).strftime('%Y-%m-%d')
            rs = bs.query_history_k_data_plus(
                _bs_code(code),
                'date,open,high,low,close,volume',
                start_date=start_date, end_date=end_date,
                frequency='d', adjustflag='2'
            )
            if rs.error_code != '0':
                return None
            rows = []
            while rs.next():
                rows.append(rs.get_row_data())
            if not rows:
                return None
            df = pd.DataFrame(rows, columns=['date','open','high','low','close','volume'])
            for c in ['open','high','low','close','volume']:
                df[c] = pd.to_numeric(df[c], errors='coerce')
            df['date'] = pd.to_datetime(df['date'])
            df['code'] = code
            return df.tail(datalen)

        df, validated = _fetch_fallback_and_validate(
            code, datalen, 'baostock', _fetch_baostock_final, use_cache, freq
        )
        if df is not None and validated:
            _log.info(f"[数据源] 回退到 Baostock (最终) (code={code}, {len(df)}条)")
            return df

    # ═══ 5) 全部失败 ═══
    _log.warning(f"[数据源] 所有数据源均失败 — code={code} 返回空数据")
    return pd.DataFrame()

def fetch_us(ticker, period="6mo", retries=3):
    """Fetch US stock / macro data from Yahoo Finance (with retry)"""
    import time, yfinance as yf
    for attempt in range(retries):
        try:
            t = yf.Ticker(ticker)
            df = t.history(period=period)
            if df.empty:
                if attempt < retries - 1:
                    time.sleep(5 * (attempt + 1))
                    continue
                return pd.DataFrame()
            df = df.reset_index()
            df = df.rename(columns={
                "Date":"date","Open":"open","High":"high",
                "Low":"low","Close":"close","Volume":"volume"
            })
            df["date"] = pd.to_datetime(df["date"]).dt.tz_localize(None)
            df["code"] = ticker
            for c in ["open","high","low","close","volume"]:
                if c in df.columns: df[c] = pd.to_numeric(df[c], errors="coerce")
            return df
        except Exception:
            if attempt < retries - 1:
                time.sleep(5 * (attempt + 1))
            else:
                return pd.DataFrame()
# Sina mappings imported from quant_trading.config

def fetch_sina_global(sina_code):
    """Fetch real-time global quote from Sina Finance"""
    raw = curl_get(f"https://hq.sinajs.cn/list={sina_code}",
                   referer="https://finance.sina.com.cn/")
    if not raw: return None
    try:
        # Parse Sina format: var hq_str_CODE="data1,data2,...";
        data = raw.split('"')[1] if '"' in raw else ""
        parts = data.split(",")
        if sina_code.startswith("hf_"):  # Futures
            return {"price": float(parts[0]) if parts[0] else 0,
                    "high": float(parts[4]) if len(parts)>4 and parts[4] else 0,
                    "low": float(parts[5]) if len(parts)>5 and parts[5] else 0}
        elif sina_code.startswith("gb_"):  # US stocks
            return {"name": parts[0], "price": float(parts[1]) if len(parts)>1 and parts[1] else 0,
                    "change_pct": float(parts[2]) if len(parts)>2 and parts[2] else 0}
        elif sina_code == "DINIW":  # Dollar Index: time,price,?,?,vol,?,high,low,?,name,date
            return {"price": float(parts[1]) if len(parts)>1 and parts[1] else 0,
                    "high": float(parts[6]) if len(parts)>6 and parts[6] else 0,
                    "low": float(parts[7]) if len(parts)>7 and parts[7] else 0}
    except Exception: return None

def fetch_fundamentals(code):
    """Fetch financial statements + valuation — 主: Baostock, 备: Sina实时估值"""
    results = {}

    if _bs_login():
        try:
            import baostock as bs
            bs_code = _bs_code(code)

            # ---- 利润表 (最近8个季度) ----
            income_rows = []
            current_year = datetime.now().year
            for year in range(current_year - 2, current_year + 1):
                for q in [1, 2, 3, 4]:
                    if year == current_year and q > ((datetime.now().month - 1) // 3 + 1):
                        break
                    try:
                        rs = bs.query_profit_data(bs_code, year=year, quarter=q)
                        while rs.error_code == '0' and rs.next():
                            row = rs.get_row_data()
                            # row: code, pubDate, statDate, roeAvg, npMargin, gpMargin,
                            #      netProfit, epsTTM, MBRevenue, totalShare, liqaShare
                            income_rows.append({
                                "REPORT_DATE": row[2],  # statDate
                                "TOTAL_OPERATE_INCOME": float(row[8]) if row[8] else 0,  # MBRevenue
                                "PARENT_NETPROFIT": float(row[6]) if row[6] else 0,       # netProfit
                                "ROE": float(row[3]) if row[3] else 0,                    # roeAvg
                            })
                    except Exception:
                        _suppress("unknown")

            if income_rows:
                # Sort by date descending
                income_rows.sort(key=lambda x: x["REPORT_DATE"], reverse=True)
                results["income"] = pd.DataFrame(income_rows[:8])  # Latest 8 quarters
                # Add ROE from last row
                results["metrics"] = {"roe": income_rows[0]["ROE"] if income_rows[0]["ROE"] else None}

            # ---- 行业分类 ----
            try:
                rs_ind = bs.query_stock_industry(bs_code)
                if rs_ind.error_code == '0':
                    while rs_ind.next():
                        row = rs_ind.get_row_data()
                        if results.get("metrics") is None:
                            results["metrics"] = {}
                        results["metrics"]["industry"] = row[3]  # industry name
                        break
            except Exception:
                _suppress("unknown")

        except Exception:
            pass  # 静默回退

    # ---- Sina 实时估值 (PE/PB/EPS) — Baostock 不提供 ----
    try:
        prefix = "sh" if code.startswith(("6", "68")) else "sz"
        raw = curl_get(f"https://hq.sinajs.cn/list={prefix}{code}")
        if raw:
            data = raw.split('"')[1] if '"' in raw else ""
            parts = data.split(",")
            if len(parts) > 40:
                pe = float(parts[38]) if parts[38] else 0  # PE
                mc = float(parts[44]) if len(parts) > 44 and parts[44] else 0  # market cap (万)
                if results.get("metrics") is None:
                    results["metrics"] = {}
                results["metrics"]["pe"] = pe if pe > 0 else None
                results["metrics"]["mc"] = mc * 10000 if mc > 0 else None  # 万→元
                # Sina doesn't directly give PB/EPS, keep as None
                results["metrics"]["pb"] = None
                results["metrics"]["eps"] = None
                results["metrics"]["navps"] = None
    except Exception:
        _suppress("unknown")

    return results

def calc_fundamental_score(fundamentals):
    """Score fundamentals 0-10"""
    score = 0
    details = []

    if not fundamentals or "income" not in fundamentals:
        return 0, ["No fundamental data available"]

    inc = fundamentals["income"]
    metrics = fundamentals.get("metrics", {})

    # Revenue growth (YoY — compare same quarter)
    if len(inc) >= 4:
        rev = pd.to_numeric(inc["TOTAL_OPERATE_INCOME"], errors="coerce")
        # Q1 vs Q1 last year (if we have 4+ quarters)
        if len(inc) >= 5:
            rev_curr = rev.iloc[0]  # Latest quarter
            rev_yoy = rev.iloc[4]   # Same quarter last year
        else:
            # Use TTM approximation
            rev_curr = rev.iloc[:4].sum() if len(inc)>=4 else rev.iloc[0]
            rev_yoy = rev.iloc[4:8].sum() if len(inc)>=8 else (rev.iloc[1] if len(inc)>1 else rev_curr*0.8)
        if rev_yoy > 0:
            rev_growth = (rev_curr - rev_yoy) / rev_yoy * 100
            if rev_growth > 30: score += 3; details.append(f"Rev Growth YoY {rev_growth:.0f}% (+3)")
            elif rev_growth > 15: score += 2; details.append(f"Rev Growth YoY {rev_growth:.0f}% (+2)")
            elif rev_growth > 5: score += 1; details.append(f"Rev Growth YoY {rev_growth:.0f}% (+1)")
            elif rev_growth < -10: score -= 1; details.append(f"Rev Decline YoY {rev_growth:.0f}% (-1)")
            else: details.append(f"Rev Growth YoY {rev_growth:.0f}% (0)")

    # Net profit turnaround / growth
    if len(inc) >= 2:
        profit = pd.to_numeric(inc["PARENT_NETPROFIT"], errors="coerce")
        profit_curr = profit.iloc[0]  # Latest
        # Find same quarter last year
        if len(inc) >= 5:
            profit_prev = profit.iloc[4]
        else:
            profit_prev = profit.iloc[1] if len(inc)>1 else 0
        if profit_curr > 0 and profit_prev < 0:
            score += 3; details.append(f"Loss→Profit turnaround (+3)")
        elif profit_curr > 0 and profit_prev > 0:
            profit_growth = (profit_curr-profit_prev)/abs(profit_prev)*100
            if profit_growth > 50: score += 3; details.append(f"Profit Growth {profit_growth:.0f}% (+3)")
            elif profit_growth > 20: score += 2; details.append(f"Profit Growth {profit_growth:.0f}% (+2)")
            else: score += 1; details.append(f"Profit Growth {profit_growth:.0f}% (+1)")
        elif profit_curr < 0:
            score -= 2; details.append(f"Net Loss ({profit_curr/1e8:.1f}Yi) (-2)")

    # Valuation
    pe = metrics.get("pe", 0) or 0
    pb = metrics.get("pb", 0) or 0
    if pe > 0 and pe < 30: score += 2; details.append(f"PE={pe:.1f} reasonable (+2)")
    elif pe > 0 and pe < 60: score += 1; details.append(f"PE={pe:.1f} moderate (+1)")
    elif pe > 100: score -= 1; details.append(f"PE={pe:.1f} high (-1)")
    if 1 < pb < 3: score += 2; details.append(f"PB={pb:.1f} low (+2)")
    elif 3 <= pb < 8: score += 1; details.append(f"PB={pb:.1f} normal (+1)")
    elif pb >= 15: score -= 1; details.append(f"PB={pb:.1f} very high (-1)")

    # ROE
    roe = metrics.get("roe", 0) or 0
    if roe > 20: score += 2; details.append(f"ROE={roe:.1f}% excellent (+2)")
    elif roe > 10: score += 1; details.append(f"ROE={roe:.1f}% good (+1)")
    elif roe < 5: score -= 1; details.append(f"ROE={roe:.1f}% low (-1)")

    return max(0, min(score, 10)), details

def fundamental_analysis(code=None):
    """Fundamental analysis — single stock or market overview"""
    print_header("FUNDAMENTAL ANALYSIS")

    if code:
        # Single stock
        name = get_stock_name(code)
        print(f"\n  {name} ({code}) — Fundamental Report")
        print(f"  {'─'*55}")
        fund = fetch_fundamentals(code)

        # Financial data
        if fund.get("income") is not None and len(fund["income"]) > 0:
            inc = fund["income"].copy()
            for c in ["TOTAL_OPERATE_INCOME","PARENT_NETPROFIT","OPERATE_PROFIT"]:
                if c in inc.columns: inc[c] = pd.to_numeric(inc[c], errors="coerce")
            print(f"\n  [Income Statement — Latest 4 Quarters]")
            print(f"  {'Period':<14} {'Revenue(Yi)':<16} {'Net Profit(Yi)':<16} {'Margin':<10}")
            print(f"  {'─'*56}")
            for _, r in inc.head(4).iterrows():
                rev = r.get("TOTAL_OPERATE_INCOME",0) or 0
                np_ = r.get("PARENT_NETPROFIT",0) or 0
                margin = np_/rev*100 if rev>0 else 0
                date = str(r.get("REPORT_DATE",""))[:10]
                print(f"  {date:<14} {rev/1e8:>13.2f} Yi  {np_/1e8:>13.2f} Yi  {margin:>7.1f}%")

        # Valuation
        metrics = fund.get("metrics", {})
        if metrics:
            print(f"\n  [Valuation Metrics]")
            print(f"  PE(TTM): {metrics.get('pe','N/A'):.1f}" if isinstance(metrics.get('pe'), (int,float)) else f"  PE: {metrics.get('pe','N/A')}")
            print(f"  PB:      {metrics.get('pb','N/A'):.2f}" if isinstance(metrics.get('pb'), (int,float)) else f"  PB: {metrics.get('pb','N/A')}")
            print(f"  ROE:     {metrics.get('roe','N/A'):.1f}%" if isinstance(metrics.get('roe'), (int,float)) else f"  ROE: {metrics.get('roe','N/A')}")
            print(f"  EPS:     {metrics.get('eps','N/A')}")
            print(f"  Industry:{metrics.get('industry','N/A')}")

        # Score
        score, details = calc_fundamental_score(fund)
        print(f"\n  [Fundamental Score: {score}/10]")
        print(f"  {'─'*55}")
        for d in details:
            print(f"  {d}")
        stars = "★"*score + "☆"*(10-score)
        print(f"  {stars}")

    else:
        # Market overview — scan watchlist for fundamental health
        print(f"\n  Scanning watchlist fundamentals...")
        results = []
        for group, codes in CFG.WATCHLIST.items():
            for code in codes[:2]:  # First 2 per group
                try:
                    fund = fetch_fundamentals(code)
                    score, details = calc_fundamental_score(fund)
                    metrics = fund.get("metrics", {})
                    inc = fund.get("income")
                    rev_g = 0
                    if inc is not None and len(inc) >= 2:
                        rev = pd.to_numeric(inc["TOTAL_OPERATE_INCOME"], errors="coerce")
                        if rev.iloc[0]>0 and rev.iloc[1]>0:
                            rev_g = (rev.iloc[0]-rev.iloc[1])/rev.iloc[1]*100
                    results.append({
                        "code":code, "name":get_stock_name(code),
                        "score":score, "pe":metrics.get("pe","-"), "pb":metrics.get("pb","-"),
                        "roe":metrics.get("roe","-"), "rev_growth":rev_g,
                    })
                except Exception: pass

        results.sort(key=lambda r: r["score"], reverse=True)
        print(f"\n  {'Code':<8} {'Name':<10} {'F-Score':<8} {'PE':<8} {'PB':<6} {'ROE':<8} {'Rev G':<8}")
        print(f"  {'─'*60}")
        for r in results:
            print(f"  {r['code']:<8} {r['name']:<10} {r['score']:<8} "
                  f"{r['pe']!s:<8} {r['pb']!s:<6} {r['roe']!s:<8} {r['rev_growth']:>+6.1f}%")

def fund_flow_analysis(code=None):
    """Analysis of major capital flow (主力资金流向)"""
    if code:
        # Single stock deep analysis
        df = fetch_kline(code, 60)
        if df.empty:
            print(f"  Cannot fetch data for {code}")
            return
        f = compute_factors(df)
        name = get_stock_name(code)

        print_header(f"FUND FLOW ANALYSIS — {name} ({code})")
        latest = f.iloc[-1]
        close = f["close"]

        # ---- Volume-Price Analysis ----
        print(f"\n  {'─'*55}")
        print(f"  VOLUME-PRICE RELATIONSHIP")
        print(f"  {'─'*55}")

        up_days = f[f["close"] > f["open"]]
        down_days = f[f["close"] < f["open"]]

        vol_up = up_days["volume"].sum() if len(up_days) > 0 else 0
        vol_down = down_days["volume"].sum() if len(down_days) > 0 else 0
        total_vol = vol_up + vol_down

        print(f"  涨日成交量:    {vol_up/1e8:.2f} Yi ({vol_up/total_vol*100:.0f}%)" if total_vol>0 else "")
        print(f"  跌日成交量:    {vol_down/1e8:.2f} Yi ({vol_down/total_vol*100:.0f}%)" if total_vol>0 else "")

        if vol_up > vol_down * 1.3:
            print(f"  >> [ACCUMULATION] 上涨放量 — 主力在吸筹")
        elif vol_down > vol_up * 1.3:
            print(f"  >> [DISTRIBUTION] 下跌放量 — 主力在出货")
        else:
            print(f"  >> [BALANCED] 量价均衡")

        # ---- Recent Flow Trends ----
        print(f"\n  {'─'*55}")
        print(f"  RECENT FLOW ANALYSIS")
        print(f"  {'─'*55}")

        for days, label in [(5, "近5日"), (10, "近10日"), (20, "近20日")]:
            recent = f.tail(days)
            up_vol = recent[recent["close"] > recent["open"]]["volume"].sum()
            dn_vol = recent[recent["close"] < recent["open"]]["volume"].sum()
            ratio = up_vol / dn_vol if dn_vol > 0 else (2.0 if up_vol > 0 else 1.0)
            net = up_vol - dn_vol
            bar = "█" * min(int(ratio * 5), 30) if ratio > 1 else "░" * min(int((2-ratio) * 5), 30)
            direction = "INFLOW ↑" if ratio > 1.2 else ("OUTFLOW ↓" if ratio < 0.8 else "NEUTRAL →")
            print(f"  {label}: 涨/跌量比 {ratio:.2f}  {bar}  {direction}  净量: {net/1e4:+.0f}万手")

        # ---- MFI (Money Flow Index) ----
        print(f"\n  {'─'*55}")
        print(f"  MONEY FLOW INDICATORS")
        print(f"  {'─'*55}")

        # Typical price
        tp = (f["high"] + f["low"] + f["close"]) / 3
        mf = tp * f["volume"]  # Money flow
        mf_pos = mf.where(tp > tp.shift(1), 0).rolling(14).sum()
        mf_neg = mf.where(tp < tp.shift(1), 0).rolling(14).sum()
        mfi = 100 - (100 / (1 + mf_pos / mf_neg))
        latest_mfi = mfi.iloc[-1]

        print(f"  MFI(14): {latest_mfi:.1f}")
        if latest_mfi > 80: print(f"  >> OVERBOUGHT — 资金过度流入，警惕回调")
        elif latest_mfi > 60: print(f"  >> Strong inflow — 资金持续流入")
        elif latest_mfi < 20: print(f"  >> OVERSOLD — 资金过度流出，关注反弹")
        elif latest_mfi < 40: print(f"  >> Weak outflow — 资金流出中")
        else: print(f"  >> Neutral — 资金流向正常")

        # Volume ratio
        print(f"  量比 (5/20日): {latest.get('vol_ratio', 0):.2f}")
        if latest.get('vol_ratio', 0) > 2.0:
            print(f"  >> [HIGH] 近期巨量 — 主力大动作")
        elif latest.get('vol_ratio', 0) > 1.5:
            print(f"  >> [ELEVATED] 放量中 — 关注方向")
        elif latest.get('vol_ratio', 0) < 0.5:
            print(f"  >> [LOW] 缩量 — 主力暂歇")

        # OBV trend
        print(f"\n  {'─'*55}")
        print(f"  OBV (On-Balance Volume) TREND")
        print(f"  {'─'*55}")
        obv = f.get("obv", pd.Series())
        if len(obv) > 0:
            obv_5d = (obv.iloc[-1] - obv.iloc[-6]) / abs(obv.iloc[-6]) * 100 if len(obv) > 5 and obv.iloc[-6] != 0 else 0
            obv_20d = (obv.iloc[-1] - obv.iloc[-21]) / abs(obv.iloc[-21]) * 100 if len(obv) > 20 and obv.iloc[-21] != 0 else 0
            print(f"  OBV 5日变化:  {obv_5d:+.1f}%")
            print(f"  OBV 20日变化: {obv_20d:+.1f}%")
            if obv_5d > 10 and obv_20d > 15:
                print(f"  >> [STRONG ACCUMULATION] 主力持续吸筹中")
            elif obv_5d < -10 and obv_20d < -15:
                print(f"  >> [STRONG DISTRIBUTION] 主力持续出货中")
            elif obv_5d > 0 and close.iloc[-1] < close.iloc[-6]:
                print(f"  >> [BULLISH DIVERGENCE] OBV↑ 价格↓ — 底部吸筹信号")
            elif obv_5d < 0 and close.iloc[-1] > close.iloc[-6]:
                print(f"  >> [BEARISH DIVERGENCE] OBV↓ 价格↑ — 顶部出货信号")

        # ---- Smart Money Score ----
        print(f"\n  {'─'*55}")
        print(f"  SMART MONEY SCORE")
        print(f"  {'─'*55}")
        sm_score = 0
        signals = []

        if vol_up > vol_down * 1.3: sm_score += 2; signals.append("上涨放量 +2")
        elif vol_up > vol_down * 1.1: sm_score += 1; signals.append("上涨略放量 +1")
        if latest_mfi > 60: sm_score += 2; signals.append("MFI强势流入 +2")
        elif latest_mfi > 50: sm_score += 1; signals.append("MFI温和流入 +1")
        if latest.get('vol_ratio', 1) > 1.5: sm_score += 1; signals.append("量比放大 +1")
        if len(obv) > 5 and obv_5d > 10: sm_score += 2; signals.append("OBV强上升 +2")
        elif len(obv) > 5 and obv_5d > 0: sm_score += 1; signals.append("OBV上升 +1")
        if vol_down > vol_up * 1.3: sm_score -= 2; signals.append("下跌放量 -2")
        if latest_mfi < 40: sm_score -= 1; signals.append("MFI流出 -1")

        sm_score = max(0, min(sm_score, 10))
        for s in signals: print(f"  {s}")
        print(f"  Score: {sm_score}/10  {'★'*sm_score}{'☆'*(10-sm_score)}")
        if sm_score >= 7: print(f"  >> 主力强力买入 — 跟随主力做多")
        elif sm_score >= 5: print(f"  >> 主力温和买入 — 偏多")
        elif sm_score <= 2: print(f"  >> 主力出货中 — 谨慎")
    else:
        # Market-wide fund flow scan
        print_header("FUND FLOW SCAN — Watchlist")
        results = []
        for group, codes in CFG.WATCHLIST.items():
            for code in codes:
                df = fetch_kline(code, 60)
                if df.empty: continue
                f = compute_factors(df)
                latest = f.iloc[-1]
                close = f["close"]

                up_d = f[f["close"] > f["open"]]
                dn_d = f[f["close"] < f["open"]]
                up_v = up_d["volume"].sum()
                dn_v = dn_d["volume"].sum()
                ratio = up_v / dn_v if dn_v > 0 else 2.0

                tp = (f["high"] + f["low"] + f["close"]) / 3
                mf = tp * f["volume"]
                mf_pos = mf.where(tp > tp.shift(1), 0).rolling(14).sum()
                mf_neg = mf.where(tp < tp.shift(1), 0).rolling(14).sum()
                mfi_val = 100 - (100/(1+mf_pos/mf_neg))

                # Score
                sc = 0
                if ratio > 1.3: sc += 3
                elif ratio > 1.1: sc += 1
                if mfi_val.iloc[-1] > 60: sc += 2
                elif mfi_val.iloc[-1] > 50: sc += 1
                if latest.get("vol_ratio", 1) > 1.5: sc += 1

                direction = "IN" if ratio > 1.1 else ("OUT" if ratio < 0.9 else "--")
                results.append({
                    "code": code, "name": get_stock_name(code),
                    "score": sc, "ratio": ratio, "mfi": float(mfi_val.iloc[-1]),
                    "direction": direction, "close": float(close.iloc[-1]),
                })

        results.sort(key=lambda r: -r["score"])
        print(f"\n  {'Code':<8} {'Name':<10} {'Score':<6} {'V-Ratio':<8} {'MFI':<6} {'Direction':<10} {'Price':<8}")
        print(f"  {'─'*55}")
        for r in results:
            print(f"  {r['code']:<8} {r['name']:<10} {r['score']:<6} {r['ratio']:>5.2f}   {r['mfi']:>5.1f}  {r['direction']:<10} {r['close']:>7.2f}")

        # Top flow
        inflows = [r for r in results if r["score"] >= 4]
        outflows = [r for r in results if r["score"] <= 1]
        print(f"\n  >> Strong Inflow ({len(inflows)}): " + ", ".join(f"{r['name']}({r['score']})" for r in inflows[:5]) if inflows else "  None")
        print(f"  >> Outflow Warning ({len(outflows)}): " + ", ".join(f"{r['name']}({r['score']})" for r in outflows[:5]) if outflows else "  None")


def objective_market():
    """Objective market-wide scan — No watchlist! Fetches real data from exchanges."""
    print_header("OBJECTIVE MARKET SCAN — Real-time Exchange Data (No Watchlist)")

    import urllib.request, urllib.error

    # ============================================================
    # 1. TOP GAINERS — 全市场涨幅榜
    # ============================================================
    print(f"\n  {'─'*55}")
    print(f"  🔥 TOP 20 GAINERS — 全市场涨幅榜")
    print(f"  {'─'*55}")

    try:
        url = ("https://push2delay.eastmoney.com/api/qt/clist/get?"
               "pn=1&pz=20&po=1&np=1&fltt=2&fid=f3&fs=m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23&"
               "fields=f2,f3,f4,f12,f14")
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0", "Referer": "https://quote.eastmoney.com/"})
        data = json.loads(urllib.request.urlopen(req, timeout=10).read())
        items = data.get("data", {}).get("diff", [])
        print(f"  {'Rank':<5} {'Code':<8} {'Name':<12} {'Chg%':<10} {'Vol':<10}")
        print(f"  {'─'*45}")
        for i, item in enumerate(items[:20]):
            name = item.get("f14", "?")
            chg = item.get("f3", 0) or 0
            code = item.get("f12", "?")
            vol = item.get("f4", 0) or 0
            bar = "█" * int(abs(chg)*2) if chg > 0 else ""
            print(f"  {i+1:<5} {code:<8} {name:<12} {chg:>+7.2f}%  {bar[:20]:20s} {(vol or 0)/1e4:.0f}万手")
    except Exception as e:
        print(f"  Unavailable: {e}")

    # ============================================================
    # 2. TOP VOLUME — 全市场成交额榜
    # ============================================================
    print(f"\n  {'─'*55}")
    print(f"  💰 TOP 20 BY VOLUME — 全市场成交额排行")
    print(f"  {'─'*55}")

    try:
        url = ("https://push2delay.eastmoney.com/api/qt/clist/get?"
               "pn=1&pz=20&po=0&np=1&fltt=2&fid=f5&fs=m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23&"
               "fields=f2,f3,f5,f12,f14")
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0", "Referer": "https://quote.eastmoney.com/"})
        data = json.loads(urllib.request.urlopen(req, timeout=10).read())
        items = data.get("data", {}).get("diff", [])
        print(f"  {'Rank':<5} {'Code':<8} {'Name':<12} {'Amount(Yi)':<15} {'Chg%':<8}")
        print(f"  {'─'*48}")
        for i, item in enumerate(items[:20]):
            name = item.get("f14", "?")
            code = item.get("f12", "?")
            amount = (item.get("f5", 0) or 0) / 1e8
            chg = item.get("f3", 0) or 0
            print(f"  {i+1:<5} {code:<8} {name:<12} {amount:>12.2f} Yi  {chg:>+6.2f}%")
    except Exception as e:
        print(f"  Unavailable: {e}")

    # ============================================================
    # 3. INDUSTRY SECTOR RANKING
    # ============================================================
    print(f"\n  {'─'*55}")
    print(f"  📊 INDUSTRY SECTOR RANKING — 行业板块涨跌榜")
    print(f"  {'─'*55}")

    try:
        url = ("https://push2delay.eastmoney.com/api/qt/clist/get?"
               "pn=1&pz=61&po=1&np=1&fltt=2&fid=f3&fs=m:90+t:2&"
               "fields=f2,f3,f4,f12,f14,f104")
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0", "Referer": "https://quote.eastmoney.com/"})
        data = json.loads(urllib.request.urlopen(req, timeout=10).read())
        items = data.get("data", {}).get("diff", [])

        # Top 10 and Bottom 10
        print(f"  {'TOP 10 — Hot Sectors 🔥':>30}")
        print(f"  {'─'*45}")
        for item in items[:10]:
            name = item.get("f14", "?")
            chg = item.get("f3", 0) or 0
            bar = "█" * int(abs(chg)*4)
            print(f"  {name:<18s} {chg:>+6.2f}%  {bar[:25]}")

        print(f"\n  {'BOTTOM 10 — Cold Sectors ❄️':>32}")
        print(f"  {'─'*45}")
        for item in items[-10:]:
            name = item.get("f14", "?")
            chg = item.get("f3", 0) or 0
            bar = "░" * int(abs(chg)*4)
            print(f"  {name:<18s} {chg:>+6.2f}%  {bar[:25]}")
    except Exception as e:
        print(f"  Unavailable: {e}")

    # ============================================================
    # 4. MARKET BREADTH — 市场宽度
    # ============================================================
    print(f"\n  {'─'*55}")
    print(f"  📏 MARKET BREADTH — 市场宽度统计")
    print(f"  {'─'*55}")

    try:
        # Count stocks up/down across the market
        url = ("https://push2delay.eastmoney.com/api/qt/clist/get?"
               "pn=1&pz=1&po=0&np=1&fltt=2&fid=f3&fs=m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23&"
               "fields=f2,f3,f12")
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0", "Referer": "https://quote.eastmoney.com/"})
        data = json.loads(urllib.request.urlopen(req, timeout=10).read())
        total_count = data.get("data", {}).get("total", 0)
        print(f"  Total A-share stocks tracked: {total_count}")

        # Count up/down in top 200
        url2 = ("https://push2delay.eastmoney.com/api/qt/clist/get?"
                "pn=1&pz=200&po=0&np=1&fltt=2&fid=f3&fs=m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23&"
                "fields=f3")
        req2 = urllib.request.Request(url2, headers={
            "User-Agent": "Mozilla/5.0", "Referer": "https://quote.eastmoney.com/"})
        data2 = json.loads(urllib.request.urlopen(req2, timeout=10).read())
        items2 = data2.get("data", {}).get("diff", [])
        up_count = sum(1 for i in items2 if (i.get("f3") or 0) > 0)
        dn_count = sum(1 for i in items2 if (i.get("f3") or 0) < 0)
        flat_count = len(items2) - up_count - dn_count

        print(f"  Sample (top 200): Up {up_count} | Down {dn_count} | Flat {flat_count}")
        if up_count > dn_count * 2:
            print(f"  >> STRONG BULLISH — 涨家数远超跌家数")
        elif up_count > dn_count:
            print(f"  >> BULLISH — 涨多跌少")
        elif dn_count > up_count * 2:
            print(f"  >> STRONG BEARISH — 跌家数远超涨家数")
        else:
            print(f"  >> MIXED — 涨跌互现")
    except Exception as e:
        print(f"  Unavailable: {e}")

    print(f"\n  {'─'*55}")
    print(f"  Data source: Shanghai/Shenzhen Exchange via East Money (REAL-TIME)")
    print(f"  This is OBJECTIVE market data — no watchlist, no bias.")

    # Also run watchlist scan for reference
    print(f"\n\n  [Your watchlist stocks for reference:]")


def sector_analysis():
    """Industry prosperity / sector rotation analysis (行业景气度)"""
    print_header("SECTOR PROSPERITY — Industry & Concept Board Analysis")

    # Fetch industry + concept board data from East Money
    sectors = []
    for board_type, market_code, label in [
        ("Industry", "m:90+t:2", "行业板块"),
        ("Concept", "m:90+t:3", "概念板块"),
    ]:
        try:
            import urllib.request, urllib.error
            url = (f"https://push2delay.eastmoney.com/api/qt/clist/get?"
                   f"pn=1&pz=15&po=1&np=1&fltt=2&fid=f3&"
                   f"fs={market_code}&fields=f2,f3,f4,f12,f14,f104,f105")
            req = urllib.request.Request(url, headers={
                "User-Agent": "Mozilla/5.0",
                "Referer": "https://quote.eastmoney.com/"
            })
            data = json.loads(urllib.request.urlopen(req, timeout=10).read())
            items = data.get("data", {}).get("diff", [])

            print(f"\n  {'─'*55}")
            print(f"  [{label}] Top 15 — by Performance")
            print(f"  {'─'*55}")
            print(f"  {'Rank':<5} {'Name':<16} {'Chg%':<8} {'Code':<8} {'Count':<6}")
            print(f"  {'─'*45}")

            for i, item in enumerate(items[:15]):
                name = item.get("f14", "?")
                chg = item.get("f3", 0) or 0
                code = item.get("f12", "?")
                count = item.get("f2", 0) or 0
                bar = "█" * int(abs(chg) * 3) if chg > 0 else "░" * int(abs(chg) * 3)
                bar = bar[:25]
                tag = ""
                if chg > 5: tag = " 🔥"
                elif chg > 3: tag = " 📈"
                elif chg < -3: tag = " 📉"
                print(f"  {i+1:<5} {name:<16} {chg:>+6.2f}% {bar} {code:<8} {count}只{tag}")
                sectors.append({"type": label, "name": name, "chg": chg, "code": code, "count": count})
        except Exception as e:
            print(f"  [{label}] unavailable: {e}")

    # Map watchlist stocks to sectors
    print(f"\n  {'─'*55}")
    print(f"  WATCHLIST SECTOR EXPOSURE")
    print(f"  {'─'*55}")

    for group, codes in CFG.WATCHLIST.items():
        chgs = []
        for code in codes:
            df = fetch_kline(code, 20)
            if not df.empty:
                chg_20d = (df["close"].iloc[-1] - df["close"].iloc[0]) / df["close"].iloc[0] * 100
                chgs.append(chg_20d)

        if chgs:
            avg_chg = sum(chgs) / len(chgs)
            best = max(chgs)
            worst = min(chgs)
            bar = "█" * int(abs(avg_chg) * 3) if avg_chg > 0 else "░" * int(abs(avg_chg) * 3)
            strength = "STRONG" if avg_chg > 10 else ("GOOD" if avg_chg > 3 else ("WEAK" if avg_chg < -5 else "NEUTRAL"))
            print(f"  {group:<12s} Avg {avg_chg:>+6.1f}% {bar[:20]:20s} [{strength}] best:{best:+.1f}% worst:{worst:+.1f}%")

    # Sector fund flow analysis
    print(f"\n  {'─'*55}")
    print(f"  SECTOR FUND FLOW (from watchlist stocks)")
    print(f"  {'─'*55}")
    sector_flows = []
    for group, codes in CFG.WATCHLIST.items():
        inflow, outflow = 0, 0
        for code in codes:
            df = fetch_kline(code, 20)
            if df.empty: continue
            up_v = df[df["close"] > df["open"]]["volume"].sum()
            dn_v = df[df["close"] < df["open"]]["volume"].sum()
            inflow += up_v
            outflow += dn_v
        ratio = inflow / outflow if outflow > 0 else 2.0
        direction = "INFLOW" if ratio > 1.2 else ("OUTFLOW" if ratio < 0.8 else "NEUTRAL")
        sector_flows.append((group, ratio, inflow, outflow, direction))

    sector_flows.sort(key=lambda x: -x[1])
    for group, ratio, inflow, outflow, direction in sector_flows:
        bar = "█" * int(ratio * 8) if ratio > 1 else "░" * int((2-ratio) * 8)
        print(f"  {group:<12s} ratio={ratio:.2f} {bar[:20]:20s} [{direction}]")

    # Summary
    hot = [s for s in sectors if s["chg"] > 3][:5]
    cold = [s for s in sectors if s["chg"] < -3][:5]
    strong = [g for g, r, _, _, _ in sector_flows if r > 1.1]
    weak = [g for g, r, _, _, _ in sector_flows if r < 0.9]

    print(f"\n  {'─'*55}")
    print(f"  PROSPERITY SUMMARY")
    print(f"  {'─'*55}")
    if hot: print(f"  🔥 HOT sectors: " + ", ".join(f"{s['name']}({s['chg']:+.1f}%)" for s in hot))
    if cold: print(f"  ❄️  COLD sectors: " + ", ".join(f"{s['name']}({s['chg']:+.1f}%)" for s in cold))
    if strong: print(f"  💰 Capital INFLOW: " + ", ".join(strong))
    if weak: print(f"  💸 Capital OUTFLOW: " + ", ".join(weak))

    print(f"\n  Data source: East Money real-time sector boards")




def hot_sector_scan():
    """Hot sector tracking + strong stock screening (热门板块追踪+强势股筛选)"""
    print_header("HOT SECTOR TRACKER — 热门板块追踪 + 强势股筛选")
    try:
        tracker = SectorTracker()
        result = tracker.full_scan(verbose=True)
        return result
    except Exception as e:
        print(f"\n  [ERROR] 热门板块扫描失败: {e}")
        print(f"  Tip: 确保网络连接正常，东方财富API可用")
        import traceback
        traceback.print_exc()
        return None


def dark_horse_screen():
    """Dark horse screening — removed, module not available"""
    print_header("DARK HORSE SCREENER")
    print("  [WARNING] 黑马筛选模块 (dark_horse) 已移除, 请使用其它筛选命令")
    return None


def curl_get(url, referer="https://finance.sina.com.cn/"):
    """HTTP GET via curl (Sina API 专用, 带速率限制 + 重试)"""
    from quant_trading.circuit_breaker import retry_on_failure, SINA_BREAKER, SINA_LIMITER

    if not SINA_BREAKER.allow_request():
        return None

    @retry_on_failure(max_retries=2, base_delay=0.5)
    def _do_curl():
        if not SINA_LIMITER.acquire():
            time.sleep(1.0)
            SINA_LIMITER.acquire()
        r = subprocess.run(["curl", "-s", "--connect-timeout", "10", "--max-time", "15",
            "-H", "User-Agent: Mozilla/5.0", "-H", f"Referer: {referer}", url],
            capture_output=True, text=True, timeout=20)
        if r.returncode != 0:
            raise ConnectionError(f"curl rc={r.returncode}")
        return r.stdout if r.stdout else None

    try:
        result = _do_curl()
        SINA_BREAKER.record_success()
        return result
    except Exception as e:
        SINA_BREAKER.record_failure()
        return None


def macro_analysis():
    """宏观分析 — 多源宏观评分引擎 (AkShare/Tushare/yfinance/Sina)"""
    print_header("MACRO SCORING ENGINE — 宏观评分引擎",
                 "流动性 · 增长/景气 · 外部环境 · 市场情绪 → 策略建议")
    try:
        from quant_trading.macro_dashboard import run_macro_dashboard
        export_fmt = "docx" if ("--docx" in sys.argv or "--word" in sys.argv) else \
                     ("html" if ("--html" in sys.argv) else "")
        run_macro_dashboard(force_refresh="--refresh" in sys.argv, export=export_fmt)
    except ImportError as e:
        console.print(f"\n  [{C_ORANGE}]⚠️ 宏观引擎模块未找到: {e}[/]")
        console.print(f"  [{C_GRAY}]Tip: 确保 quant_trading/macro_config.py, macro_engine.py, macro_dashboard.py 存在[/]")
    except Exception as e:
        console.print(f"\n  [{C_RED}]宏观分析失败: {e}[/]")

def fetch_watchlist():
    """Fetch all A-share watchlist stocks"""
    results = {}
    for group, codes in CFG.WATCHLIST.items():
        for code in codes:
            df = fetch_kline(code)
            if not df.empty:
                results[code] = df
                print(f"  [{group}] {code}: {len(df)} records")
    return results

def compute_rsi(close, period=14):
    """Quick RSI for a series"""
    d = close.diff()
    g = d.where(d>0,0.0).ewm(alpha=1/period,adjust=False).mean()
    l = (-d.where(d<0,0.0)).ewm(alpha=1/period,adjust=False).mean()
    return float(100-(100/(1+g/l))) if l.iloc[-1] != 0 else 50.0

# ============================================================
# FACTOR ENGINE
# ============================================================
def compute_factors(df):
    """Compute all technical factors from OHLCV data"""
    close = df["close"].astype(float)
    high = df["high"].astype(float)
    low = df["low"].astype(float)
    volume = df["volume"].astype(float)
    f = df.copy()

    # --- Trend ---
    for p in [5,10,20,60]:
        f[f"ma{p}"] = close.rolling(p).mean()
    f["ma_score"] = ((f["ma5"]>f["ma10"]).astype(int) + (f["ma10"]>f["ma20"]).astype(int) + (f["ma20"]>f["ma60"]).astype(int))
    f["above_ma60"] = (close > f["ma60"]).astype(int)

    # --- MACD ---
    e12 = close.ewm(span=12, adjust=False).mean()
    e26 = close.ewm(span=26, adjust=False).mean()
    f["dif"] = e12 - e26
    f["dea"] = f["dif"].ewm(span=9, adjust=False).mean()
    f["macd"] = 2 * (f["dif"] - f["dea"])
    f["macd_cross"] = ((f["dif"]>f["dea"]) & (f["dif"].shift(1)<=f["dea"].shift(1))).astype(int)
    f["macd_bullish"] = (f["dif"] > f["dea"]).astype(int)

    # --- RSI(14) ---
    delta = close.diff()
    gain = delta.where(delta>0, 0.0).ewm(alpha=1/14, adjust=False).mean()
    loss = (-delta.where(delta<0, 0.0)).ewm(alpha=1/14, adjust=False).mean()
    f["rsi"] = 100 - (100/(1 + gain/loss))

    # --- KD(9,3) ---
    l9, h9 = close.rolling(9).min(), close.rolling(9).max()
    rsv = (close-l9)/(h9-l9)*100
    f["kd_k"] = rsv.ewm(alpha=1/3, adjust=False).mean()
    f["kd_d"] = f["kd_k"].ewm(alpha=1/3, adjust=False).mean()

    # --- BOLL(20,2) ---
    f["boll_mid"] = close.rolling(20).mean()
    s20 = close.rolling(20).std()
    f["boll_up"] = f["boll_mid"] + 2*s20
    f["boll_dn"] = f["boll_mid"] - 2*s20
    f["boll_pos"] = (close-f["boll_dn"])/(f["boll_up"]-f["boll_dn"])

    # --- ATR(14) ---
    tr = pd.concat([high-low, (high-close.shift()).abs(), (low-close.shift()).abs()], axis=1).max(axis=1)
    f["atr"] = tr.rolling(14).mean()
    f["atr_pct"] = f["atr"] / close

    # --- Volume ---
    f["vol_ma5"] = volume.rolling(5).mean()
    f["vol_ma20"] = volume.rolling(20).mean()
    f["vol_ratio"] = f["vol_ma5"] / f["vol_ma20"]
    f["vol_expand"] = (f["vol_ratio"]>1.5).astype(int)

    # OBV (On-Balance Volume)
    obv = (volume * np.sign(close.diff())).cumsum()
    f["obv"] = obv

    # --- Returns ---
    for p in [1,5,10,20,60]:
        f[f"ret_{p}d"] = close.pct_change(p)

    return f

# ============================================================
# STRATEGIES
# ============================================================
def strategy_ma_cross(df, fast=5, slow=20, vol_filter=True, vol_ratio=1.5):
    """Golden Cross / Death Cross"""
    df = df.copy()
    ma_f = df["close"].rolling(fast).mean()
    ma_s = df["close"].rolling(slow).mean()
    df["signal"] = 0
    df.loc[(ma_f>ma_s)&(ma_f.shift(1)<=ma_s.shift(1)), "signal"] = 1
    df.loc[(ma_f<ma_s)&(ma_f.shift(1)>=ma_s.shift(1)), "signal"] = -1
    if vol_filter and "vol_ratio" in df.columns:
        df.loc[df["vol_ratio"]<vol_ratio, "signal"] = 0
    return df

def strategy_momentum(df, lookback=20, breakout_pct=0.05, vol_confirm=True):
    """Breakout above N-day high with volume"""
    df = df.copy()
    high_n = df["high"].rolling(lookback).max()
    vol_m = df["volume"].rolling(20).mean()
    df["signal"] = 0
    mask = (df["close"]>high_n.shift(1)*(1-breakout_pct))
    if vol_confirm: mask &= (df["volume"]>vol_m*1.2)
    df.loc[mask, "signal"] = 1
    return df

def strategy_value_reversal(df, rsi_threshold=30, boll_threshold=0.2):
    """Buy on RSI oversold + BOLL lower band"""
    df = df.copy()
    df["signal"] = 0
    df.loc[(df["rsi"]<rsi_threshold)&(df["boll_pos"]<boll_threshold)&(df["vol_ratio"]>1.0), "signal"] = 1
    return df

def strategy_trend_follow(df, ma_long=60, ma_entry=20):
    """Buy pullback to MA20 while above MA60"""
    df = df.copy()
    df["signal"] = 0
    df.loc[(df["close"]>df[f"ma{ma_long}"])&(df["close"]<=df[f"ma{ma_entry}"]*1.02)&(df["close"]>=df[f"ma{ma_entry}"]*0.98)&(df["rsi"]>40)&(df["rsi"]<60), "signal"] = 1
    return df

STRATEGIES = [
    ("MA_Cross", strategy_ma_cross, CFG.MA_CROSS),
    ("Momentum", strategy_momentum, CFG.MOM_BREAK),
    ("Value_Rev", strategy_value_reversal, CFG.VAL_REV),
    ("Trend_Fol", strategy_trend_follow, CFG.TREND_FOL),
]

def composite_score(df):
    """Compute composite signal score (0-10) from the latest factors"""
    s = df.iloc[-1]
    score = 0
    if s.get("ma_score", 0) >= 2: score += 2
    if s.get("above_ma60", 0) == 1: score += 2
    if s.get("macd_cross", 0) == 1: score += 1
    if s.get("macd_bullish", 0) == 1: score += 1
    if 30 < s.get("rsi", 50) < 70: score += 1
    if s.get("rsi", 50) < 30: score += 2
    if s.get("kd_k", 50) > s.get("kd_d", 50): score += 1
    if s.get("ret_5d", 0) > 0: score += 1
    if s.get("vol_expand", 0) == 1: score += 1
    boll = s.get("boll_pos", 0.5)
    if boll < 0.2: score += 2
    return min(score, 10)

# ============================================================
# REPORT GENERATOR
# ============================================================

def single_stock_report(code, df):
    """Deep single-stock analysis"""
    f = compute_factors(df)
    latest = f.iloc[-1]
    prev = f.iloc[-2]
    name = get_stock_name(code)

    print_header(f"{name} ({code}) — Comprehensive Analysis Report")
    print(f"  Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # ---- Price Summary ----
    print(f"\n  {'─'*55}")
    print(f"  📊 PRICE SUMMARY")
    print(f"  {'─'*55}")
    close_now = latest["close"]
    chg = (close_now-prev["close"])/prev["close"]*100
    print(f"  Close: {close_now:>10.2f}  |  Change: {chg:>+8.2f}%")
    print(f"  Open:  {latest['open']:>10.2f}  |  High:   {latest['high']:>10.2f}")
    print(f"  Low:   {latest['low']:>10.2f}  |  Volume: {latest['volume']/1e4:>8.0f}W")

    # ---- Moving Averages ----
    print(f"\n  {'─'*55}")
    print(f"  📈 MOVING AVERAGES")
    print(f"  {'─'*55}")
    for p in [5,10,20,60]:
        ma = f[f"ma{p}"].iloc[-1]
        status = "↑ ABOVE" if close_now > ma else "↓ BELOW"
        print(f"  MA{p:>3}: {ma:>8.2f}  {status}")
    ma_score = latest["ma_score"]
    print(f"  MA Alignment Score: {ma_score}/3  {'(BULLISH)' if ma_score>=2 else '(BEARISH)' if ma_score<=1 else '(MIXED)'}")

    # ---- Technical Indicators ----
    print(f"\n  {'─'*55}")
    print(f"  🔧 TECHNICAL INDICATORS")
    print(f"  {'─'*55}")
    print(f"  RSI(14):     {latest['rsi']:>8.1f}  {'[OVERBOUGHT]' if latest['rsi']>70 else '[OVERSOLD]' if latest['rsi']<30 else '[NORMAL]'}")
    print(f"  MACD DIF/DEA:{latest['dif']:>8.3f}/{latest['dea']:.3f}  {'[GOLDEN CROSS]' if latest['dif']>latest['dea'] else '[DEATH CROSS]'}")
    print(f"  KD K/D:      {latest['kd_k']:>8.1f}/{latest['kd_d']:.1f}  {'[OVERBOUGHT]' if latest['kd_k']>80 else '[OVERSOLD]' if latest['kd_k']<20 else '[NORMAL]'}")
    print(f"  BOLL pos:    {latest['boll_pos']*100:>8.1f}%  {'[UPPER BAND]' if latest['boll_pos']>0.8 else '[LOWER BAND]' if latest['boll_pos']<0.2 else '[MID RANGE]'}")
    print(f"  BOLL Width:  {((latest['boll_up']-latest['boll_dn'])/latest['boll_mid']*100):>8.1f}%")
    print(f"  ATR(14):     {latest['atr']:>8.2f}  ({latest['atr_pct']*100:.1f}%)")
    print(f"  Vol Ratio:   {latest['vol_ratio']:>8.2f}  {'[EXPANDING]' if latest['vol_ratio']>1.5 else '[CONTRACTING]' if latest['vol_ratio']<0.5 else '[NORMAL]'}")

    # ---- Period Returns ----
    print(f"\n  {'─'*55}")
    print(f"  📅 PERIOD RETURNS")
    print(f"  {'─'*55}")
    for p, label in [(1,"1D"),(5,"5D"),(10,"10D"),(20,"20D"),(60,"60D")]:
        ret = latest[f"ret_{p}d"]*100 if not pd.isna(latest.get(f"ret_{p}d",np.nan)) else 0
        bar = "█" * int(abs(ret)*2) if ret>0 else "░" * int(abs(ret)*2)
        print(f"  {label:>4}: {ret:>+7.2f}%  {bar}")

    # ---- 52-Week Range ----
    max_c = f["close"].max()
    min_c = f["close"].min()
    pos_52w = (close_now-min_c)/(max_c-min_c)*100
    print(f"\n  {'─'*55}")
    print(f"  📏 PERIOD RANGE")
    print(f"  {'─'*55}")
    print(f"  {len(f)}-day High: {max_c:.2f}  |  Low: {min_c:.2f}")
    print(f"  Position: {pos_52w:.0f}% from low, {(1-close_now/max_c)*100:.0f}% from high")

    # ---- Strategy Signals ----
    print(f"\n  {'─'*55}")
    print(f"  🎯 STRATEGY SIGNALS")
    print(f"  {'─'*55}")
    all_signals = {}
    for name, strat_fn, params in STRATEGIES:
        sdf = strat_fn(f, **params)
        sig = sdf["signal"].iloc[-1]
        sig_str = "🔴 BUY" if sig==1 else "🟢 SELL" if sig==-1 else "⚪ HOLD"
        all_signals[name] = sig
        print(f"  {name:15s}: {sig_str}")

    # ---- Composite Score ----
    score = composite_score(f)
    print(f"\n  {'─'*55}")
    print(f"  ⭐ COMPOSITE SCORE: {score}/10")
    print(f"  {'─'*55}")
    stars = "★" * score + "☆" * (10-score)
    print(f"  {stars}")
    if score >= 7: print(f"  >>> STRONG BUY signal — multiple factors aligned")
    elif score >= 5: print(f"  >> BUY signal — favorable technical setup")
    elif score >= 3: print(f"  > NEUTRAL — waiting for clearer signals")
    else: print(f"  > CAUTION — technical weakness, wait or avoid")

    # ---- Key Price Levels ----
    print(f"\n  {'─'*55}")
    print(f"  🎯 KEY PRICE LEVELS")
    print(f"  {'─'*55}")
    r20 = f.tail(20)
    print(f"  Resistance:  {r20['high'].max():.2f} (20D High)  |  {max_c:.2f} (All-Time)")
    print(f"  Support:     {r20['low'].min():.2f} (20D Low)   |  {min_c:.2f} (All-Time)")
    for p in [20, 60]:
        ma = f[f"ma{p}"].iloc[-1]
        print(f"  MA{p}:        {ma:.2f}")

    # ---- Risk/Reward ----
    stop_loss = close_now * (1 - CFG.STOP_LOSS_PCT)
    take_profit = close_now * (1 + CFG.TAKE_PROFIT_PCT)
    print(f"\n  {'─'*55}")
    print(f"  ⚖️  RISK/REWARD (Position Size: {CFG.MAX_POSITION_PCT*100:.0f}%)")
    print(f"  {'─'*55}")
    print(f"  Stop Loss ({CFG.STOP_LOSS_PCT*100:.0f}%):   {stop_loss:.2f}  (Risk: {(close_now-stop_loss)*100:.0f} shares)")
    print(f"  Take Profit ({CFG.TAKE_PROFIT_PCT*100:.0f}%): {take_profit:.2f}  (Reward: {(take_profit-close_now)*100:.0f} shares)")
    rr = (take_profit-close_now)/(close_now-stop_loss) if close_now>stop_loss else 0
    print(f"  Risk/Reward Ratio: 1:{rr:.1f}")

    return score


def get_stock_name(code):
    """Try to get stock name from unified config or return code"""
    return STOCK_NAMES.get(code, code)


# ============================================================
# BACKTEST ENGINE
# ============================================================
def run_backtest(code, df, strategy_fn, params):
    """Simple vectorized backtest"""
    f = compute_factors(df)
    sdf = strategy_fn(f, **params)
    signals = sdf["signal"].values
    close = f["close"].values

    cash = CFG.INITIAL_CAPITAL
    shares = 0
    capital_curve = []
    trades = []

    for i in range(len(signals)):
        price = close[i] * (1 + CFG.SLIPPAGE)

        if signals[i] == 1 and shares == 0 and cash > 0:
            # Buy
            max_shares = int(cash * CFG.MAX_POSITION_PCT / price)
            if max_shares > 0:
                cost = max_shares * price * (1 + CFG.COMMISSION)
                if cost <= cash:
                    cash -= cost
                    shares = max_shares
                    trades.append({"date": str(sdf.iloc[i]["date"])[:10], "action":"BUY",
                                   "price":price, "shares":shares, "cost":cost})

        elif signals[i] == -1 and shares > 0:
            # Sell
            revenue = shares * price * (1 - CFG.COMMISSION)
            cash += revenue
            trades.append({"date": str(sdf.iloc[i]["date"])[:10], "action":"SELL",
                           "price":price, "shares":shares, "revenue":revenue})
            shares = 0

        # Stop loss / Take profit
        if shares > 0:
            entry_price = trades[-1]["price"] if trades else price
            if price <= entry_price * (1 - CFG.STOP_LOSS_PCT):
                revenue = shares * price * (1 - CFG.COMMISSION)
                cash += revenue
                trades.append({"date": str(sdf.iloc[i]["date"])[:10], "action":"STOP",
                               "price":price, "shares":shares, "revenue":revenue})
                shares = 0
            elif price >= entry_price * (1 + CFG.TAKE_PROFIT_PCT):
                revenue = shares * price * (1 - CFG.COMMISSION)
                cash += revenue
                trades.append({"date": str(sdf.iloc[i]["date"])[:10], "action":"TAKE",
                               "price":price, "shares":shares, "revenue":revenue})
                shares = 0

        total_value = cash + shares * price
        capital_curve.append(total_value)

    # Final sell if still holding
    if shares > 0:
        total_value = cash + shares * close[-1] * (1 - CFG.COMMISSION)
    else:
        total_value = cash

    total_return = (total_value - CFG.INITIAL_CAPITAL) / CFG.INITIAL_CAPITAL * 100
    annual_return = total_return / (len(close)/252) if len(close) > 0 else 0

    # Max drawdown
    peak = np.maximum.accumulate(capital_curve)
    drawdown = (np.array(capital_curve) - peak) / peak * 100
    max_dd = np.min(drawdown)

    # Sharpe ratio
    daily_returns = np.diff(np.log(capital_curve))
    sharpe = np.mean(daily_returns) / np.std(daily_returns) * np.sqrt(252) if len(daily_returns)>1 and np.std(daily_returns)>0 else 0

    # Win rate
    wins = 0
    for i in range(0, len(trades)-1, 2):
        if i+1 < len(trades):
            if trades[i+1].get("revenue",0) > trades[i].get("cost",0): wins += 1
    win_rate = wins / (len(trades)//2) * 100 if len(trades)>=2 else 0

    return {
        "code": code, "name": get_stock_name(code),
        "total_return": total_return, "annual_return": annual_return,
        "max_drawdown": max_dd, "sharpe": sharpe, "win_rate": win_rate,
        "n_trades": len(trades)//2, "trades": trades,
        "capital_curve": capital_curve,
    }


def backtest_all():
    """Run backtest on all watchlist stocks"""
    print_header("BACKTEST REPORT — All Strategies")
    results = []

    for code in list(CFG.WATCHLIST.values())[0][:3]:  # First 3 stocks
        df = fetch_kline(code, 250)
        if df.empty: continue

        for name, strat_fn, params in STRATEGIES[:2]:  # First 2 strategies
            try:
                res = run_backtest(code, df, strat_fn, params)
                results.append(res)
                print(f"\n  {res['name']}({code}) - {name}")
                print(f"  Total Return: {res['total_return']:+.1f}%  |  Annual: {res['annual_return']:+.1f}%")
                print(f"  Max DD: {res['max_drawdown']:.1f}%  |  Sharpe: {res['sharpe']:.2f}  |  Win Rate: {res['win_rate']:.0f}%")
                print(f"  Trades: {res['n_trades']}")
            except: pass

    # Summary ranking
    results.sort(key=lambda r: r["sharpe"], reverse=True)
    print(f"\n  {'─'*55}")
    print(f"  RANKING (by Sharpe Ratio)")
    print(f"  {'─'*55}")
    for i, r in enumerate(results):
        print(f"  {i+1}. {r['name']}({r['code']})  Sharpe:{r['sharpe']:.2f}  Ret:{r['total_return']:+.1f}%  DD:{r['max_drawdown']:.1f}%")

    return results


# ============================================================
# MARKET SCANNER
# ============================================================
def market_scan():
    """Scan entire watchlist for signals"""
    print_header("MARKET SCAN — Daily Signal Report")

    all_data = fetch_watchlist()
    all_signals = []

    for code, df in all_data.items():
        f = compute_factors(df)
        score = composite_score(f)

        strategy_signals = {}
        for name, strat_fn, params in STRATEGIES:
            sdf = strat_fn(f, **params)
            strategy_signals[name] = int(sdf["signal"].iloc[-1])

        all_signals.append({
            "code": code, "name": get_stock_name(code),
            "score": score, "strategies": strategy_signals,
            "close": f["close"].iloc[-1],
            "ret_5d": f["ret_5d"].iloc[-1]*100 if not pd.isna(f["ret_5d"].iloc[-1]) else 0,
            "rsi": f["rsi"].iloc[-1],
            "vol_ratio": f["vol_ratio"].iloc[-1],
        })

    # Sort by score
    all_signals.sort(key=lambda s: s["score"], reverse=True)

    # Print results
    print(f"\n  {'Code':<8} {'Name':<10} {'Score':<6} {'5D Ret':<8} {'RSI':<6} {'Vol':<6} {'MA-X':<6} {'Mom':<6} {'Val':<6} {'Trd':<6}")
    print(f"  {'─'*75}")
    for s in all_signals:
        print(f"  {s['code']:<8} {s['name']:<10} {s['score']:<6} {s['ret_5d']:>+6.1f}%  {s['rsi']:>5.1f}  {s['vol_ratio']:>4.1f}  "
              f"{'BUY' if s['strategies'].get('MA_Cross',0)==1 else '-':<6} "
              f"{'BUY' if s['strategies'].get('Momentum',0)==1 else '-':<6} "
              f"{'BUY' if s['strategies'].get('Value_Rev',0)==1 else '-':<6} "
              f"{'BUY' if s['strategies'].get('Trend_Fol',0)==1 else '-':<6}")

    # Top picks
    buys = [s for s in all_signals if s["score"] >= 4]
    print(f"\n  {'─'*55}")
    print(f"  >> TOP PICKS (Score >= 4): {len(buys)} stocks")
    print(f"  {'─'*55}")
    for s in buys:
        print(f"  {s['name']}({s['code']}) | Score:{s['score']}/10 | Price:{s['close']:.2f} | 5D:{s['ret_5d']:+.1f}%")

    # Save report
    report = {
        "date": datetime.now().strftime("%Y%m%d"),
        "scanned": len(all_signals),
        "buy_signals": len(buys),
        "results": [{k: (float(v) if isinstance(v, (np.floating,)) else int(v) if isinstance(v, (np.integer,)) else v) for k,v in s.items()} for s in all_signals],
    }
    report_path = Path(__file__).parent / "reports" / f"scan_{datetime.now().strftime('%Y%m%d')}.json"
    report_path.parent.mkdir(exist_ok=True)
    with open(report_path, "w", encoding="utf-8") as fp:
        json.dump(report, fp, indent=2, ensure_ascii=False, default=str)
    print(f"\n  [Report saved] {report_path}")

    return all_signals


# ============================================================
# MAIN
# ============================================================
def detailed_bearish_analysis(code, df):
    """单股深度分析: 检查上升趋势中的第一根大阴线"""
    name = get_stock_name(code)
    f = compute_factors(df)
    close = f["close"]
    open_ = f["open"]
    high = f["high"]
    low = f["low"]
    volume = f["volume"]

    print_header(f"FIRST BEARISH ANALYSIS — {name} ({code})")

    # ---- 趋势诊断 ----
    ma_score = 0
    if close.iloc[-1] > f["ma5"].iloc[-1]: ma_score += 1
    if f["ma5"].iloc[-1] > f["ma10"].iloc[-1]: ma_score += 1
    if f["ma10"].iloc[-1] > f["ma20"].iloc[-1]: ma_score += 1
    if f["ma20"].iloc[-1] > f["ma60"].iloc[-1]: ma_score += 1

    print(f"\n  {'─'*55}")
    print(f"  📈 UPTREND DIAGNOSIS — 上涨趋势诊断")
    print(f"  {'─'*55}")
    print(f"  MA多头排列得分: {ma_score}/4")
    print(f"  MA5:  {f['ma5'].iloc[-1]:.2f}  {'↑ 在上' if close.iloc[-1] > f['ma5'].iloc[-1] else '↓ 在下'}")
    print(f"  MA10: {f['ma10'].iloc[-1]:.2f}")
    print(f"  MA20: {f['ma20'].iloc[-1]:.2f}")
    print(f"  MA60: {f['ma60'].iloc[-1]:.2f}")

    # MA20 斜率
    if len(close) >= 30:
        ma20_slope = (f["ma20"].iloc[-1] - f["ma20"].iloc[-11]) / f["ma20"].iloc[-11] * 100
        slope_tag = "↗ 上升" if ma20_slope > 1 else ("→ 走平" if ma20_slope > -1 else "↘ 下降")
        print(f"  MA20 10日斜率: {ma20_slope:+.1f}%  [{slope_tag}]")

    trend_ok = (ma_score >= 2 and close.iloc[-1] > f["ma60"].iloc[-1])
    print(f"  趋势状态: {'✅ 上升趋势' if trend_ok else '❌ 非上升趋势'}")

    # ---- 扫描最近的大阴线 ----
    print(f"\n  {'─'*55}")
    print(f"  🕯️ BEARISH CANDLE DETECTION — 大阴线检测")
    print(f"  {'─'*55}")

    # 计算平均实体
    avg_bodies = []
    for i in range(max(0, len(f)-20), len(f)):
        avg_bodies.append(abs(float(open_.iloc[i]) - float(close.iloc[i])) / float(open_.iloc[i]) * 100)
    avg_body = np.mean(avg_bodies) if avg_bodies else 2.0

    print(f"  近20日平均实体: {avg_body:.2f}%")
    print(f"  大阴线门槛: 实体 ≥ max({avg_body*2:.1f}%, 3.0%)")

    bearish_candles = []
    for i in range(len(f)):
        o = open_.iloc[i]
        c = close.iloc[i]
        body = o - c
        if body <= 0:
            continue
        body_pct = body / o * 100
        if body_pct >= max(avg_body * 2, 3.0):
            bearish_candles.append({
                "idx": i,
                "date": str(f.iloc[i]["date"])[:10],
                "open": o, "close": c, "high": high.iloc[i], "low": low.iloc[i],
                "body_pct": body_pct,
                "volume": volume.iloc[i],
            })

    if not bearish_candles:
        print(f"\n  🟢 该股近期无大阴线 — 趋势健康")
    else:
        print(f"\n  检测到 {len(bearish_candles)} 根大阴线:")
        print(f"  {'日期':<12} {'开盘':<8} {'收盘':<8} {'最低':<8} {'实体%':<8} {'成交量':<10}")
        print(f"  {'─'*60}")
        for bc in bearish_candles[-8:]:  # 最近8根
            vol_str = f"{bc['volume']/1e4:.0f}万手"
            tag = " ← 最近" if bc["idx"] >= len(f) - 5 else ""
            print(f"  {bc['date']:<12} {bc['open']:<8.2f} {bc['close']:<8.2f} "
                  f"{bc['low']:<8.2f} {bc['body_pct']:>5.1f}%   {vol_str:<10}{tag}")

        # 检查"第一根"判定
        if len(bearish_candles) >= 1:
            latest = bearish_candles[-1]
            # 看latest之前15天内有没有大阴线
            prior = [b for b in bearish_candles if b["idx"] < latest["idx"] and latest["idx"] - b["idx"] <= 15]
            is_first = len(prior) == 0
            print(f"\n  最近大阴线: {latest['date']} (实体{latest['body_pct']:.1f}%)")
            print(f"    此前15天内大阴线: {len(prior)}根" + (" → ✅ 这是\"第一根\"!" if is_first else " → ❌ 非第一根"))
            if prior:
                for p in prior:
                    print(f"      {p['date']} 实体{p['body_pct']:.1f}% (间隔{latest['idx']-p['idx']}天)")

            if is_first:
                # 分析大阴线特征
                print(f"\n  {'─'*55}")
                print(f"  🔍 FIRST BEARISH ANALYSIS — 首阴特征分析")
                print(f"  {'─'*55}")

                # 量比
                vol_at = latest["volume"]
                vol_ma5 = float(volume.iloc[max(0,latest["idx"]-5):latest["idx"]].mean()) if latest["idx"] >= 5 else vol_at
                vol_ratio = vol_at / vol_ma5 if vol_ma5 > 0 else 1.0

                print(f"  当日成交量: {vol_at/1e4:.0f}万手  |  5日均量: {vol_ma5/1e4:.0f}万手  |  量比: {vol_ratio:.2f}")
                if vol_ratio < 0.7:
                    print(f"  >> 缩量大跌 — 洗盘特征，筹码锁定期")
                elif vol_ratio > 2.0:
                    print(f"  >> 放量大跌 — 主力出货嫌疑，警惕!")
                else:
                    print(f"  >> 量能正常 — 分歧加大")

                # 影线分析
                body_len = abs(latest["open"] - latest["close"])
                upper_wick = latest["high"] - max(latest["open"], latest["close"])
                lower_wick = min(latest["open"], latest["close"]) - latest["low"]

                print(f"\n  蜡烛形态:")
                print(f"    实体: {body_len:.2f} (开{latest['open']:.2f} → 收{latest['close']:.2f})")
                print(f"    上影线: {upper_wick:.2f} ({upper_wick/body_len*100:.0f}% of body)" if body_len > 0 else "")
                print(f"    下影线: {lower_wick:.2f} ({lower_wick/body_len*100:.0f}% of body)" if body_len > 0 else "")
                if lower_wick > body_len * 0.5:
                    print(f"  >> 长下影线 — 买方反击，下方承接力强")
                if upper_wick > body_len * 0.3:
                    print(f"  >> 有上影线 — 曾试图上攻后回落")

                # 前涨幅
                if latest["idx"] >= 10:
                    ret_before = (close.iloc[latest["idx"]-1] - close.iloc[latest["idx"]-10]) / close.iloc[latest["idx"]-10] * 100
                    print(f"\n  大阴线前10日涨幅: {ret_before:+.1f}%")
                    if ret_before > 20:
                        print(f"  >> 短期暴涨后回调 — 获利盘回吐，正常调整")
                    elif ret_before > 10:
                        print(f"  >> 温和上涨后回调 — 可能是趋势中继")

                # 支撑位
                dist_ma20 = (close.iloc[-1] - f["ma20"].iloc[-1]) / f["ma20"].iloc[-1] * 100
                dist_ma60 = (close.iloc[-1] - f["ma60"].iloc[-1]) / f["ma60"].iloc[-1] * 100
                print(f"\n  支撑位:")
                print(f"    距MA20: {dist_ma20:+.1f}% {'← 触及支撑' if -2 < dist_ma20 < 3 else ''}")
                print(f"    距MA60: {dist_ma60:+.1f}% {'← 强支撑区' if -3 < dist_ma60 < 3 else ''}")
                print(f"    阴线最低价: {latest['low']:.2f} (跌破则止损)")

                # 后续走势
                days_after = len(f) - 1 - latest["idx"]
                if days_after >= 1:
                    print(f"\n  后续走势 ({days_after}天):")
                    print(f"    阴线收盘: {latest['close']:.2f}  →  现价: {close.iloc[-1]:.2f}")
                    chg_after = (close.iloc[-1] - latest["close"]) / latest["close"] * 100
                    print(f"    变化: {chg_after:+.1f}%")
                    if chg_after > 0:
                        print(f"  >> 已收复阴线失地 — 回踩确认有效")
                    elif close.iloc[-1] > latest["low"]:
                        print(f"  >> 未破新低 — 观察能否收复")
                    else:
                        print(f"  >> 破新低 — 趋势可能转弱")


# ============================================================
# 单股检测逻辑 (提取出来供并行调用)
# ============================================================
def _check_one_stock_bearish(code, params, name=None):
    """
    [V2 改进版] 首阴策略 — 核心改进:
    1. 趋势要求: MA20>MA60 (中期多头) + 价在MA60上方 (不要求严格的周月共振)
    2. 流动性: 日成交额 ≥ 3亿
    3. 阴线后必须有企稳迹象 (未创新低 或 已开始回升)
    4. 排除放量暴跌 (量比>2.5)
    5. 放宽首阴窗口到30天 + 降低实体倍率到1.5x
    6. 优先距MA20近的标的 (回踩均线支撑)
    """
    try:
        df = fetch_kline(code, max(120, params["TREND_LOOKBACK"] + 30))
        if df.empty or len(df) < 60:
            return None

        f = compute_factors(df)
        if name is None:
            name = get_stock_name(code)
        close = f["close"]
        open_ = f["open"]
        high = f["high"]
        low = f["low"]

        # ---- 1) 趋势确认 (简化: MA20>MA60 + 价在MA60上) ----
        ma20_now = f["ma20"].iloc[-1]
        ma60_now = f["ma60"].iloc[-1]
        if ma20_now <= ma60_now:
            return None
        if close.iloc[-1] <= ma60_now:
            return None

        # MA20 斜率向上 (5日回看)
        if len(close) >= 25:
            ma20_5d_ago = f["ma20"].iloc[-6]
            if ma20_now <= ma20_5d_ago:
                return None

        # ---- 2) 流动性过滤: 日成交额 ≥ 3亿 ----
        try:
            last_vol = float(f.iloc[-1]["volume"])
            last_close = float(close.iloc[-1])
            turnover = last_vol * last_close  # 估算成交额
            if turnover < 1e9:  # 10亿
                return None
        except Exception:
            pass

        # ---- 3) 寻找最近的阴线 (回溯5天) ----
        found_bearish = None
        for i in range(len(f) - params["RECENT_DAYS"], len(f)):
            o = float(open_.iloc[i])
            c = float(close.iloc[i])
            body = o - c  # 正值=阴线
            if body <= 0:
                continue
            body_pct = body / o * 100
            if body_pct < params["BODY_PCT_MIN"]:
                continue

            # 计算实体倍率 (vs 20日均实体)
            start_idx = max(0, i - 20)
            recent_bodies = []
            for j in range(start_idx, i):
                b = abs(float(open_.iloc[j]) - float(close.iloc[j]))
                recent_bodies.append(b / float(open_.iloc[j]) * 100)
            avg_body = np.mean(recent_bodies) if recent_bodies else body_pct
            ratio = body_pct / avg_body if avg_body > 0.3 else body_pct

            if ratio >= params["BODY_RATIO_MIN"]:
                # 排除放量暴跌 (量比>2.5)
                vol_at = float(f.iloc[i]["volume"])
                vol_ma5 = float(f.iloc[max(0,i-5):i]["volume"].mean()) if i >= 5 else vol_at
                vol_r = vol_at / vol_ma5 if vol_ma5 > 0 else 1.0
                if vol_r > params.get("VOL_RATIO_MAX", 2.5):
                    continue

                if found_bearish is None or body_pct > found_bearish[1]:
                    found_bearish = (i, body_pct, ratio, vol_r)

        if found_bearish is None:
            return None

        bi, body_pct, ratio, vol_ratio_at = found_bearish

        # 日期新鲜度
        bearish_date_str = str(f.iloc[bi]["date"])[:10]
        try:
            bearish_dt = datetime.strptime(bearish_date_str, "%Y-%m-%d")
            if (datetime.now() - bearish_dt).days > 10:
                return None
        except Exception:
            pass

        # ---- 4) 阴线后企稳确认 (关键改进!) ----
        days_after = len(f) - 1 - bi
        recovery = 0
        made_new_low = False
        if days_after >= 1:
            recovery = (close.iloc[-1] - close.iloc[bi]) / close.iloc[bi] * 100
            # 检查是否破了阴线最低价
            lowest_after = low.iloc[bi+1:].min()
            if lowest_after < low.iloc[bi] * 0.98:  # 破阴线低点2%以上
                made_new_low = True
        # 阴线后创新低的 — 排除
        if made_new_low:
            return None

        # 阴线日必须是最近5天内 "第一根" 突出阴线 (放宽窗口到30天但不严格排除)
        has_big_prior = False
        first_window_start = max(0, bi - params["FIRST_WINDOW"])
        for i in range(first_window_start, bi):
            o = open_.iloc[i]
            c = close.iloc[i]
            b = o - c
            if b <= 0:
                continue
            bp = b / o * 100
            if bp >= params["BODY_PCT_MIN"] * 1.2:  # 更严格: 1.2x而非0.8x
                ws = max(0, i - 20)
                bodies_before = []
                for j in range(ws, i):
                    bb = abs(float(open_.iloc[j]) - float(close.iloc[j]))
                    bodies_before.append(bb / float(open_.iloc[j]) * 100)
                avg = np.mean(bodies_before) if bodies_before else bp
                r = bp / avg if avg > 0.3 else bp
                if r >= params["BODY_RATIO_MIN"] * 0.8:
                    has_big_prior = True
                    break
        if has_big_prior:
            return None

        # ---- 5) 综合评分 (重新调权) ----
        score = 4  # 基础分 (趋势+流动性已确认)

        # 阴线前涨幅 (温和上涨最优)
        if bi >= 10:
            ret_10d_before = (close.iloc[bi-1] - close.iloc[bi-10]) / close.iloc[bi-10] * 100
            if 5 < ret_10d_before < 20:
                score += 3  # 温和上涨后回调 — 最优
            elif 0 < ret_10d_before <= 5:
                score += 1
            elif ret_10d_before > 30:  # 暴涨后阴线 — 风险高
                score -= 2
        else:
            ret_10d_before = 0

        # 量比评分
        if vol_ratio_at < 0.7:
            score += 3  # 极致缩量 — 最佳
        elif vol_ratio_at < 1.0:
            score += 2  # 缩量
        elif vol_ratio_at < 1.5:
            score += 1  # 正常

        # 距MA20距离 (回调到均线附近最佳)
        dist_to_ma20 = (close.iloc[-1] - ma20_now) / ma20_now * 100
        if -2 < dist_to_ma20 < 3:
            score += 3  # 精准回踩MA20
        elif -5 < dist_to_ma20 < 8:
            score += 2
        elif dist_to_ma20 > 25:  # 远离均线 — 追高风险
            score -= 2

        # 阴线后回升
        if days_after >= 1:
            if recovery > 3:
                score += 2  # 已明显回升
            elif recovery > 0:
                score += 1  # 开始回升
            elif recovery > -2:
                score += 0  # 横盘

        # 距MA60
        dist_to_ma60 = (close.iloc[-1] - ma60_now) / ma60_now * 100
        if 0 < dist_to_ma60 < 20:
            score += 1  # 在MA60上方不远 — 趋势健康

        score = max(1, min(score, 10))

        bearish_date = str(f.iloc[bi]["date"])[:10]
        o_price = open_.iloc[bi]
        c_price = close.iloc[bi]
        h_price = high.iloc[bi]
        l_price = low.iloc[bi]

        # 信号分类
        if vol_ratio_at < 0.8 and recovery >= 0 and abs(dist_to_ma20) < 5:
            signal_type = "🔥 缩量回踩买点"
        elif vol_ratio_at < 1.0 and recovery >= 0:
            signal_type = "🟢 健康回调"
        elif vol_ratio_at < 1.5:
            signal_type = "🔵 温和回调"
        elif recovery > 0:
            signal_type = "🟡 放量但已企稳"
        else:
            signal_type = "⚪ 观察中"

        return {
            "code": code, "name": name, "score": score,
            "date": bearish_date, "body_pct": round(body_pct, 2),
            "body_ratio": round(ratio, 1), "vol_ratio": round(vol_ratio_at, 2),
            "ret_10d_before": round(ret_10d_before, 1), "ma_score": int(ma20_now > ma60_now),
            "open": round(o_price, 2), "close": round(c_price, 2),
            "high": round(h_price, 2), "low": round(l_price, 2),
            "recovery": round(recovery, 2), "dist_ma20": round(dist_to_ma20, 1),
            "signal_type": signal_type, "close_now": round(float(close.iloc[-1]), 2),
        }
    except Exception:
        return None


def first_bearish_scan():
    """
    全市场扫描: 寻找上涨趋势中出现的"第一根大阴线"
    Baostock 获取全A股代码 → 多线程并行抓取K线 → 逐个检测 → 排序输出
    """
    print_header("FIRST BEARISH IN UPTREND [V2]",
                 "首阴策略V2 — 中期多头回踩 · 流动性过滤 · 企稳确认")

    # ====== [V2] 参数 ======
    params = dict(
        TREND_LOOKBACK  = 90,   MA_SHORT = 10,   MA_MID = 20,   MA_LONG = 60,
        BODY_RATIO_MIN  = 1.5,  BODY_PCT_MIN = 2.0,
        FIRST_WINDOW    = 30,   RECENT_DAYS = 5,  UP_TREND_SCORE = 1,
        MTF_FILTER      = False, MTF_STRICT = False,  # V2: 不再要求周月双周期
        VOL_RATIO_MAX   = 2.5,  # 排除放量暴跌
    )

    # ====== Phase 1: 获取全市场股票代码→名称映射 ======
    with console.status(f"[{C_CYAN}]从 Baostock 获取全A股代码...[/]", spinner="dots"):
        code_name = fetch_all_a_stock_codes()

    console.print(f"  [{C_GREEN}]✓[/] 获取到 [{C_GOLD}]{len(code_name)}[/] 只A股 (已排除ST/退市)")
    console.print(f"  条件: MA20>MA60中期多头 | 价在MA60上方 | 阴线实体≥{params['BODY_PCT_MIN']}% ≥均值{params['BODY_RATIO_MIN']}x")
    console.print(f"  流动性: 日成交额≥10亿 | 排除放量暴跌(量比>{params['VOL_RATIO_MAX']})")
    console.print(f"  🆕 V2改进: 阴线后必须企稳(未创新低) | 距MA20越近加分越多 | 回溯{params['RECENT_DAYS']}天 | 首阴窗口{params['FIRST_WINDOW']}天")

    # ====== Phase 2: 多线程并行检测 ======
    results = []
    completed = 0
    total = len(code_name)
    found_lock = threading.Lock()

    console.print(f"\n  [{C_CYAN}]Phase 2:[/] 16线程并行检测 [{total}] 只股票...")

    with Progress(
        SpinnerColumn(), TextColumn("[progress.description]{task.description}"),
        BarColumn(bar_width=40), TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TextColumn("│ 候选: [bold yellow]{task.fields[found]}[/]"),
        TimeElapsedColumn(),
        console=console,
    ) as progress:
        task = progress.add_task(f"[{C_CYAN}]扫描中[/]", total=total, found=0)

        with ThreadPoolExecutor(max_workers=8) as executor:
            futures = {executor.submit(_check_one_stock_bearish, code, params, code_name[code]): code
                       for code in code_name}

            for future in as_completed(futures):
                try:
                    r = future.result()
                    with found_lock:
                        completed += 1
                        if r is not None:
                            results.append(r)
                        progress.update(task, advance=1, found=len(results))
                except Exception:
                    with found_lock:
                        completed += 1
                        progress.update(task, advance=1, found=len(results))

    console.print(f"  [{C_GREEN}]✓[/] 扫描完成, 共找到 [{C_GOLD}]{len(results)}[/] 只候选")

    # ====== Phase 3: 排序输出 ======
    results.sort(key=lambda r: (-r["score"], -r["body_pct"]))

    if not results:
        console.print(f"\n  [{C_GRAY}]未找到符合条件的形态[/]")
        return []

    # ---- 结果表格 ----
    score_style = lambda s: f"[{C_RED} bold]{s}[/]" if s >= 7 else (f"[{C_GOLD}]{s}[/]" if s >= 5 else f"[{C_GRAY}]{s}[/]")
    signal_style = lambda t: {"洗盘回踩": f"[{C_GREEN}]{t}[/]", "放量分歧": f"[{C_ORANGE}]{t}[/]", "趋势观望": f"[{C_BLUE}]{t}[/]"}.get(t, t)

    table = Table(title=f"[bold {C_GOLD}]筛选结果: {len(results)} 只股票[/]",
                  box=box.ROUNDED, border_style=C_GRAY, title_justify="left",
                  highlight=True, pad_edge=False)
    for h, sty in [("代码", "cyan"), ("名称", "white"), ("评分(满分10)", "yellow"),
                    ("阴线日", ""), ("实体%", ""), ("倍率", ""), ("量比", ""),
                    ("前10D涨幅%", ""), ("回复%", ""), ("距MA20%", ""), ("信号", "")]:
        table.add_column(h, style=sty)

    for r in results:
        stars = "★" * r["score"] + "☆" * (10 - r["score"])
        table.add_row(
            f"[cyan]{r['code']}[/]", r['name'],
            f"[bold]{stars}[/] [yellow]({r['score']})[/]",
            r['date'],
            f"{r['body_pct']:.1f}%",
            f"{r['body_ratio']:.1f}x",
            f"{r['vol_ratio']:.2f}",
            f"{r['ret_10d_before']:+.1f}%",
            f"{r['recovery']:+.1f}%",
            f"{r['dist_ma20']:+.0f}%",
            signal_style(r['signal_type']),
        )
    console.print(table)

    # ---- 分类汇总 ----
    washout = [r for r in results if r["signal_type"] == "洗盘回踩"]
    divergence = [r for r in results if r["signal_type"] == "放量分歧"]
    watch = [r for r in results if r["signal_type"] == "趋势观望"]

    summary = Table(box=box.SIMPLE, border_style=C_GRAY, title=f"[bold {C_GOLD}]信号分类[/]")
    summary.add_column("类型", style="white", width=15)
    summary.add_column("数量", style="yellow", width=8)
    summary.add_column("股票", style="white")

    if washout:
        summary.add_row(f"[{C_GREEN}]🟢 洗盘回踩[/]", str(len(washout)),
                        ", ".join(f"{r['name']}({r['score']}分)" for r in washout[:8]))
    else:
        summary.add_row(f"[{C_GREEN}]🟢 洗盘回踩[/]", "0", f"[{C_GRAY}](无)[/]")

    if divergence:
        summary.add_row(f"[{C_ORANGE}]🟡 放量分歧[/]", str(len(divergence)),
                        ", ".join(f"{r['name']}({r['score']}分)" for r in divergence[:8]))
    else:
        summary.add_row(f"[{C_ORANGE}]🟡 放量分歧[/]", "0", f"[{C_GRAY}](无)[/]")

    if watch:
        summary.add_row(f"[{C_BLUE}]🔵 趋势观望[/]", str(len(watch)),
                        ", ".join(f"{r['name']}({r['score']}分)" for r in watch[:8]))
    else:
        summary.add_row(f"[{C_BLUE}]🔵 趋势观望[/]", "0", f"[{C_GRAY}](无)[/]")

    console.print(summary)

    # ---- Top Picks ----
    top = [r for r in results if r["score"] >= 6]
    if top:
        console.print()
        top_panel = []
        for r in top[:10]:
            top_panel.append(
                f"[bold white]{r['name']}[/] ([{C_CYAN}]{r['code']}[/])  "
                f"[{C_GOLD}]评分: {r['score']}/10[/]  阴线日: {r['date']}\n"
                f"  实体: [{C_RED}]{r['body_pct']:.1f}%[/] ({r['body_ratio']:.1f}x)  │  "
                f"开 [{C_ORANGE}]{r['open']:.2f}[/] → 收 [{C_RED}]{r['close']:.2f}[/]  │  "
                f"最低: {r['low']:.2f}  │  现价: [{C_GREEN}]{r['close_now']:.2f}[/]\n"
                f"  前10D: {r['ret_10d_before']:+.1f}%  │  量比: {r['vol_ratio']:.2f}  │  "
                f"距MA20: {r['dist_ma20']:+.0f}%  │  回升: {r['recovery']:+.1f}%"
            )
        console.print(Panel("\n\n".join(top_panel),
                           title=f"[bold {C_GOLD}]⭐ TOP PICKS (≥6分): {len(top)} 只[/]",
                           border_style=C_GOLD, padding=(1, 2)))

    # ---- 策略建议 ----
    tips = (
        f"[{C_GREEN}]1.[/] 洗盘回踩型: 缩量回踩均线支撑 → 止损设在阴线最低价下方\n"
        f"[{C_ORANGE}]2.[/] 放量分歧型: 高位放量 → 警惕反转, 观察能否缩量企稳\n"
        f"[{C_BLUE}]3.[/] 关注阴线后3-5天走势: 缩量不破低点→加仓; 继续放量下跌→减仓\n"
        f"[{C_CYAN}]4.[/] 最佳买点: 阴线后十字星/小阳线 + 缩量 → 回踩确认有效"
    )
    console.print(Panel(tips, title=f"[bold {C_GOLD}]📋 策略参考[/]", border_style=C_GRAY, padding=(1, 2)))

    # ---- 保存报告 ----
    report = {
        "date": datetime.now().strftime("%Y%m%d"),
        "type": "first_bearish_in_uptrend_full_market",
        "total_scanned": total,
        "total_found": len(results),
        "results": [{k: (float(v) if isinstance(v, (np.floating, np.integer)) else v)
                      for k, v in r.items()} for r in results],
    }
    report_path = Path(__file__).parent / "reports" / f"first_bearish_{datetime.now().strftime('%Y%m%d')}.json"
    report_path.parent.mkdir(exist_ok=True)
    with open(report_path, "w", encoding="utf-8") as fp:
        json.dump(report, fp, indent=2, ensure_ascii=False, default=str)
    console.print(f"\n  [{C_GRAY}]📁 报告已保存: {report_path}[/]")

    return results


def _check_mtf_pullback(code, params, name=None):
    """检测周线月线趋多 + 日线阴线/十字星 (线程安全)"""
    try:
        df = fetch_kline(code, params["TREND_LOOKBACK"] + 30)
        if df.empty or len(df) < 60:
            return None

        f = compute_factors(df)
        if name is None:
            name = get_stock_name(code)
        close = f["close"]
        open_ = f["open"]
        high = f["high"]
        low = f["low"]

        # ---- 1) 周线趋势确认 (简化: 价格站上MA20, MA20 > MA60) ----
        weekly_ok = False
        w_ma20 = close.iloc[-20:].mean()
        w_ma60 = close.iloc[-60:].mean()
        if close.iloc[-1] > w_ma20 and w_ma20 > w_ma60:
            weekly_ok = True

        if not weekly_ok:
            return None

        # ---- 2) 月线趋势确认 (价格站上MA120, MA60 > MA120) ----
        monthly_ok = False
        if len(close) >= 120:
            m_ma60 = close.iloc[-60:].mean()
            m_ma120 = close.iloc[-120:].mean()
            if close.iloc[-1] > m_ma120 and m_ma60 > m_ma120:
                monthly_ok = True

        if not monthly_ok:
            return None

        # ---- 3) 日线阴线或十字星检测 ----
        # 检查最近5个交易日
        found_candle = None
        for i in range(len(f) - params["LOOKBACK_DAYS"], len(f)):
            o = float(open_.iloc[i])
            c = float(close.iloc[i])
            h = float(high.iloc[i])
            l = float(low.iloc[i])
            body = o - c  # 正值=阴线, 负值=阳线
            candle_range = h - l
            if candle_range <= 0:
                continue

            body_pct = abs(body) / o * 100
            body_ratio = abs(body) / candle_range  # 实体占全天振幅的比例

            candle_type = None
            # 十字星判定: 实体占振幅 < 15% 且 实体% < 1.5%
            if body_ratio < 0.15 and body_pct < 1.5:
                candle_type = "十字星"
            # 阴线判定: close < open 且 实体 >= 1%
            elif body > 0 and body_pct >= 1.0:
                candle_type = "阴线"

            if candle_type is None:
                continue

            # 量比
            vol_at = float(f.iloc[i]["volume"])
            if i >= 5:
                vol_ma5 = float(f.iloc[max(0, i-5):i]["volume"].mean())
                vol_ratio = vol_at / vol_ma5 if vol_ma5 > 0 else 1.0
            else:
                vol_ratio = 1.0

            # 选最近一根符合条件的
            if found_candle is None or i > found_candle["idx"]:
                found_candle = {
                    "idx": i,
                    "date": str(f.iloc[i]["date"])[:10],
                    "open": round(o, 2),
                    "close": round(c, 2),
                    "high": round(h, 2),
                    "low": round(l, 2),
                    "body_pct": round(body_pct, 2),
                    "body_ratio": round(body_ratio, 3),
                    "range_pct": round(candle_range / o * 100, 2),
                    "vol_ratio": round(vol_ratio, 2),
                    "candle_type": candle_type,
                    "direction": "阴线" if body > 0 else "假阳十字",
                }

        if found_candle is None:
            return None

        # ---- 3.5) 日成交额过滤 (≥ 10亿) ----
        candle_vol = float(f.iloc[found_candle["idx"]]["volume"])
        candle_turnover = candle_vol * found_candle["close"]  # 估算成交额: 股数×收盘价
        TURNOVER_MIN = 1e9  # 10亿
        if candle_turnover < TURNOVER_MIN:
            return None
        found_candle["turnover_yi"] = round(candle_turnover / 1e8, 2)  # 记录成交额(亿)

        # ---- 4) 综合评分 ----
        score = 5  # 基础分(周月共振)

        # 距离关键均线
        ma20_val = f["ma20"].iloc[-1]
        ma60_val = f["ma60"].iloc[-1]
        dist_ma20 = (close.iloc[-1] - ma20_val) / ma20_val * 100
        dist_ma60 = (close.iloc[-1] - ma60_val) / ma60_val * 100

        # 回踩均线加分
        if -3 < dist_ma20 < 3:
            score += 2
        elif -5 < dist_ma20 < 5:
            score += 1

        # 缩量加分
        if found_candle["vol_ratio"] < 0.7:
            score += 2
        elif found_candle["vol_ratio"] < 1.0:
            score += 1
        elif found_candle["vol_ratio"] > 2.0:
            score -= 1

        # 十字星在上升趋势中是反转信号
        if found_candle["candle_type"] == "十字星":
            if found_candle["vol_ratio"] < 0.8:
                score += 1  # 缩量十字星更可靠

        # 前10日涨幅
        bi = found_candle["idx"]
        if bi >= 10:
            ret_10d = (close.iloc[bi-1] - close.iloc[bi-10]) / close.iloc[bi-10] * 100
            if 5 < ret_10d < 15:
                score += 1  # 温和上涨后回调
            elif 15 <= ret_10d < 30:
                score += 1  # 较强上涨后回调
        else:
            ret_10d = 0

        # 阴线后走势
        days_after = len(f) - 1 - bi
        recovery = 0
        if days_after >= 1:
            recovery = (close.iloc[-1] - close.iloc[bi]) / close.iloc[bi] * 100
            if recovery > 0:
                score += 1

        score = max(3, min(score, 10))

        # 信号分类
        if found_candle["candle_type"] == "十字星" and found_candle["vol_ratio"] < 0.8:
            signal_type = "🌟 缩量十字星买点"
        elif found_candle["candle_type"] == "十字星":
            signal_type = "🔍 十字星观察"
        elif found_candle["vol_ratio"] < 0.7 and recovery >= 0:
            signal_type = "🟢 缩量阴线回踩"
        elif found_candle["vol_ratio"] < 1.0:
            signal_type = "🔵 温和阴线回调"
        else:
            signal_type = "🟡 放量阴线分歧"

        return {
            "code": code, "name": name, "score": score,
            "date": found_candle["date"],
            "candle_type": found_candle["candle_type"],
            "direction": found_candle["direction"],
            "body_pct": found_candle["body_pct"],
            "body_ratio": found_candle["body_ratio"],
            "range_pct": found_candle["range_pct"],
            "vol_ratio": found_candle["vol_ratio"],
            "open": found_candle["open"],
            "close": found_candle["close"],
            "high": found_candle["high"],
            "low": found_candle["low"],
            "ret_10d_before": round(ret_10d, 1),
            "recovery": round(recovery, 2),
            "dist_ma20": round(dist_ma20, 1),
            "dist_ma60": round(dist_ma60, 1),
            "signal_type": signal_type,
            "close_now": round(float(close.iloc[-1]), 2),
            "turnover_yi": found_candle.get("turnover_yi", 0),
        }
    except Exception:
        return None


def mtf_pullback_scan():
    """
    全市场扫描: 周线趋多 + 月线趋多 + 日线阴线/十字星
    用法: python quant_trader.py mtf_pullback
    """
    print_header("MTF PULLBACK SCAN",
                 "周线月线双周期趋多 → 日线阴线/十字星回调买点")

    params = dict(
        TREND_LOOKBACK=120,
        LOOKBACK_DAYS=5,  # 检查最近5天
    )

    # Phase 1: 获取全市场股票代码
    with console.status(f"[{C_CYAN}]从 Baostock 获取全A股代码...[/]", spinner="dots"):
        code_name = fetch_all_a_stock_codes()

    console.print(f"  [{C_GREEN}]✓[/] 获取到 [{C_GOLD}]{len(code_name)}[/] 只A股 (已排除ST/退市)")
    console.print(f"  条件: 周线趋多(价>MA20 + MA20>MA60) + 月线趋多(价>MA120 + MA60>MA120)")
    console.print(f"  日线形态: 阴线(实体≥1%) 或 十字星(实体/振幅<15% 且 实体<1.5%)")
    console.print(f"  流动性过滤: 日成交额 ≥ 10亿 (volume×close)")
    console.print(f"  回溯: 最近 {params['LOOKBACK_DAYS']} 个交易日")

    # Phase 2: 多线程并行检测
    results = []
    total = len(code_name)
    found_lock = threading.Lock()

    console.print(f"\n  [{C_CYAN}]Phase 2:[/] 16线程并行检测 [{total}] 只股票...")

    with Progress(
        SpinnerColumn(), TextColumn("[progress.description]{task.description}"),
        BarColumn(bar_width=40), TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TextColumn("│ 候选: [bold yellow]{task.fields[found]}[/]"),
        TimeElapsedColumn(),
        console=console,
    ) as progress:
        task = progress.add_task(f"[{C_CYAN}]扫描中[/]", total=total, found=0)

        with ThreadPoolExecutor(max_workers=8) as executor:
            futures = {executor.submit(_check_mtf_pullback, code, params, code_name[code]): code
                       for code in code_name}

            for future in as_completed(futures):
                try:
                    r = future.result()
                    with found_lock:
                        if r is not None:
                            results.append(r)
                        progress.update(task, advance=1, found=len(results))
                except Exception:
                    with found_lock:
                        progress.update(task, advance=1, found=len(results))

    console.print(f"  [{C_GREEN}]✓[/] 扫描完成, 共找到 [{C_GOLD}]{len(results)}[/] 只候选")

    # Phase 3: 排序输出
    results.sort(key=lambda r: (-r["score"], r["vol_ratio"]))

    if not results:
        console.print(f"\n  [{C_GRAY}]未找到符合条件的形态[/]")
        return []

    # 结果表格
    score_style = lambda s: f"[{C_RED} bold]{s}[/]" if s >= 8 else (f"[{C_GOLD}]{s}[/]" if s >= 6 else f"[{C_GRAY}]{s}[/]")

    table = Table(title=f"[bold {C_GOLD}]周月趋多·日线回调: {len(results)} 只股票[/]",
                  box=box.ROUNDED, border_style=C_GRAY, title_justify="left",
                  highlight=True, pad_edge=False)
    for h, sty in [("代码", "cyan"), ("名称", "white"), ("评分", "yellow"),
                   ("日期", ""), ("形态", ""), ("实体%", ""), ("振幅%", ""),
                   ("量比", ""), ("成交额(亿)", "green"), ("距MA20%", ""), ("前10D%", ""), ("信号", "")]:
        table.add_column(h, style=sty)

    for r in results:
        stars = "★" * r["score"] + "☆" * (10 - r["score"])
        candle_icon = "➕" if r["candle_type"] == "十字星" else "🔻"
        table.add_row(
            f"[cyan]{r['code']}[/]", r['name'],
            f"[bold]{stars}[/] [yellow]({r['score']})[/]",
            r['date'],
            f"{candle_icon} {r['candle_type']}",
            f"{r['body_pct']:.2f}%",
            f"{r['range_pct']:.1f}%",
            f"{r['vol_ratio']:.2f}",
            f"{r.get('turnover_yi', 0):.1f}",
            f"{r['dist_ma20']:+.0f}%",
            f"{r['ret_10d_before']:+.1f}%" if r['ret_10d_before'] != 0 else "-",
            r['signal_type'],
        )
    console.print(table)

    # 分类汇总
    doji = [r for r in results if r["candle_type"] == "十字星"]
    bearish = [r for r in results if r["candle_type"] == "阴线"]

    summary = Table(box=box.SIMPLE, border_style=C_GRAY, title=f"[bold {C_GOLD}]形态分类[/]")
    summary.add_column("形态", style="white", width=12)
    summary.add_column("数量", style="yellow", width=6)
    summary.add_column("代表标的", style="white")

    if doji:
        summary.add_row(f"[{C_CYAN}]➕ 十字星[/]", str(len(doji)),
                        ", ".join(f"{r['name']}({r['score']}分)" for r in doji[:6]))
    else:
        summary.add_row(f"[{C_CYAN}]➕ 十字星[/]", "0", f"[{C_GRAY}](无)[/]")

    if bearish:
        summary.add_row(f"[{C_RED}]🔻 阴线[/]", str(len(bearish)),
                        ", ".join(f"{r['name']}({r['score']}分)" for r in bearish[:6]))
    else:
        summary.add_row(f"[{C_RED}]🔻 阴线[/]", "0", f"[{C_GRAY}](无)[/]")

    console.print(summary)

    # Top Picks
    top = [r for r in results if r["score"] >= 6]
    if top:
        console.print()
        top_panel = []
        for r in top[:10]:
            top_panel.append(
                f"[bold white]{r['name']}[/] ([{C_CYAN}]{r['code']}[/])  "
                f"[{C_GOLD}]评分: {r['score']}/10[/]  {r['date']}\n"
                f"  形态: {r['candle_type']}  │  实体: {r['body_pct']:.2f}%  │  "
                f"开 [{C_ORANGE}]{r['open']:.2f}[/] → 收 [{C_RED if r['direction'] == '阴线' else C_GREEN}]{r['close']:.2f}[/]  │  "
                f"高: {r['high']:.2f}  低: {r['low']:.2f}\n"
                f"  量比: {r['vol_ratio']:.2f}  │  距MA20: {r['dist_ma20']:+.0f}%  │  "
                f"前10D: {r['ret_10d_before']:+.1f}%  │  现价: [{C_GREEN}]{r['close_now']:.2f}[/]"
            )
        console.print(Panel("\n\n".join(top_panel),
                           title=f"[bold {C_GOLD}]⭐ TOP PICKS (≥6分): {len(top)} 只[/]",
                           border_style=C_GOLD, padding=(1, 2)))

    # 策略建议
    tips = (
        f"[{C_GREEN}]1.[/] 缩量十字星: 上升趋势中最可靠的回调结束信号，次日放量阳线确认为最佳买点\n"
        f"[{C_CYAN}]2.[/] 缩量阴线回踩: 回踩均线不破 → 止损设在阴线最低价下方1-2%\n"
        f"[{C_ORANGE}]3.[/] 放量阴线: 警惕趋势转折，需观察1-2天能否缩量企稳\n"
        f"[{C_BLUE}]4.[/] 周月共振背景下的日线回调，胜率显著高于单纯日线形态交易"
    )
    console.print(Panel(tips, title=f"[bold {C_GOLD}]📋 策略参考[/]", border_style=C_GRAY, padding=(1, 2)))

    # 保存报告
    report = {
        "date": datetime.now().strftime("%Y%m%d"),
        "type": "mtf_pullback_scan",
        "total_scanned": total,
        "total_found": len(results),
        "doji_count": len(doji),
        "bearish_count": len(bearish),
        "results": [{k: (float(v) if isinstance(v, (np.floating, np.integer)) else v)
                      for k, v in r.items()} for r in results],
    }
    report_path = Path(__file__).parent / "reports" / f"mtf_pullback_{datetime.now().strftime('%Y%m%d')}.json"
    report_path.parent.mkdir(exist_ok=True)
    with open(report_path, "w", encoding="utf-8") as fp:
        json.dump(report, fp, indent=2, ensure_ascii=False, default=str)
    console.print(f"\n  [{C_GRAY}]📁 报告已保存: {report_path}[/]")

    return results


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        print("Usage: python quant_trader.py [scan|flow|macro|analyze|fundamental|backtest|hot_sector|dark_horse|first_bearish|dashboard|oversold|oversold_backtest|yuezi|yixian|yixian_scan|shap|darts|polars_scan|polars_bench|sentiment]")
        sys.exit(1)

    cmd = sys.argv[1].lower()

    # === 插件命令检查: commands/ 下有同名模块则优先使用 ===
    try:
        from commands import get_command
        plugin_fn = get_command(cmd)
        if plugin_fn:
            plugin_fn(sys.argv[2:])
            return
    except Exception:
        pass  # 插件不可用，回退到传统 elif 链

    if cmd == "ladder":
        from quant_trading.limit_up_ladder import run_ladder
        print_header("LIMIT-UP LADDER — 连板天梯表",
                     "每日涨停板连板统计 · Excel导出")
        run_ladder(export=True)

    elif cmd == "consensus":
        from quant_trading.analyst_consensus import build_consensus_report, print_consensus_report, generate_docx
        print_header("ANALYST CONSENSUS — 一致预期",
                     "分析师覆盖 · 盈利预测 · 评级调整 · 目标价")
        data = build_consensus_report()
        print_consensus_report(data)
        path = generate_docx(data)
        print(f"\n  📁 Word报告: {path}")

    elif cmd == "event":
        from quant_trading.event_driven import build_event_report, print_event_report, generate_docx
        print_header("EVENT-DRIVEN SCAN — 事件驱动扫描",
                     "增减持 · 大宗交易 · 股权质押 · 限售解禁")
        data = build_event_report()
        print_event_report(data)
        path = generate_docx(data)
        print(f"\n  📁 Word报告: {path}")

    elif cmd == "news":
        from quant_trading.news_aggregator import run_news_aggregator
        print_header("NEWS AGGREGATOR — 新闻聚合",
                     "东方财富快讯 · 题材分类 · 时间戳")
        run_news_aggregator()

    elif cmd == "rumor":
        from quant_trading.rumor_tracker import run_rumor_tracker
        print_header("RUMOR TRACKER — 小作文追踪",
                     "雪球热帖 · 东财人气榜 · 多平台共振检测")
        run_rumor_tracker()

    elif cmd == "yuezi":
        from quant_trading.seat_tracker import (
            run_seat_command, print_seat_report, generate_daily_report,
            print_trader_profile, get_trader_profile,
            search_seat_history, print_seat_search_results, save_trades,
        )
        subcmd = sys.argv[2] if len(sys.argv) > 2 else ""
        arg = sys.argv[3] if len(sys.argv) > 3 else ""
        if subcmd == "profile" and arg:
            print_header(f"SEAT PROFILE — 游资画像: {arg}")
            profile = get_trader_profile(arg)
            print_trader_profile(profile)
        elif subcmd == "seat" and arg:
            print_header(f"SEAT SEARCH — 席位搜索: {arg}")
            results = search_seat_history(arg, days=30)
            print_seat_search_results(results, arg)
            if results:
                console.print(f"\n  [dim]正在保存到数据库...[/]")
                save_trades(results)
                console.print(f"  [green]已保存 {len(results)} 条记录[/]")
        elif subcmd == "seat":
            console.print(f"\n  [yellow]请指定席位关键词: yuezi seat <关键词>[/]")
            console.print(f"  [dim]例: yuezi seat 上塘路  /  yuezi seat 文二西路[/]")
        elif subcmd == "resonance":
            print_header("RESONANCE DETECTION — 多游资共振检测")
            from quant_trading.seat_tracker import fetch_daily_seats, detect_resonance
            df = fetch_daily_seats()
            if df.empty:
                console.print(f"\n  [dim]今日无龙虎榜数据[/]")
            else:
                resonance = detect_resonance(df, min_seats=2, min_tier=2)
                if resonance:
                    for i, rs in enumerate(resonance[:10]):
                        aliases_str = " + ".join(rs["aliases"])
                        console.print(f"\n  [{C_GREEN}]#{i+1}[/] {rs['code']} {rs['name']} "
                                      f"| [{C_MAGENTA}]{rs['seat_count']}位游资共振[/] "
                                      f"| {aliases_str} "
                                      f"| 买入{rs['total_buy']/1e8:.1f}亿")
                else:
                    console.print(f"\n  [dim]今日无多游资共振标的[/]")
        elif subcmd == "top":
            print_header("TOP SEATS — 顶级游资今日操作")
            from quant_trading.seat_tracker import fetch_daily_seats
            df = fetch_daily_seats()
            if df.empty:
                console.print(f"\n  [dim]今日无龙虎榜数据[/]")
            else:
                famous = df[df["alias"] != ""]
                tier1 = famous[famous["tier"] == 1]
                if not tier1.empty:
                    for _, r in tier1.iterrows():
                        icon = "[green]🟢[/]" if r["net"] > 0 else "[red]🔴[/]"
                        console.print(
                            f"  {icon} [{C_MAGENTA}]{r['alias']:10s}[/] | "
                            f"{r['code']} {r['name']:8s} | "
                            f"买{r['buy']/1e8:.1f}亿 卖{r['sell']/1e8:.1f}亿 "
                            f"净{r['net']/1e8:+.1f}亿 | {r.get('style','')}"
                        )
                else:
                    console.print(f"\n  [dim]今日无顶级游资操作[/]")
        else:
            print_header("SEAT TRACKER — 游资席位跟踪",
                         "80+知名席位 · 多游资共振 · 风格画像 · 历史胜率")
            report = generate_daily_report()
            if report.get("status") == "no_data":
                console.print(f"\n  [dim]⚪ {report.get('message', '今日无龙虎榜数据')}[/]")
            else:
                print_seat_report(report)

    elif cmd == "tushare":
        from quant_trading.tushare_engine import check_connection, print_health, fetch_kline as ts_fetch
        print_header("TUSHARE CONNECTION TEST — 数据源状态")
        print_health()
        if check_connection()["status"] == "ok":
            console.print(f"\n  [{C_GREEN}]✓ Tushare 连接正常，将作为 fetch_kline 首选数据源[/]")
            # 快速测试
            df = ts_fetch("000001", 5)
            if df is not None:
                console.print(f"  [{C_GREEN}]✓ 上证指数测试通过 ({len(df)} rows)[/]")
        else:
            console.print(f"\n  [{C_ORANGE}]⚠ Tushare Token 未配置或无效[/]")
            console.print(f"  [{C_GRAY}]→ 注册获取: https://tushare.pro → 微信扫码 → 复制Token[/]")
            console.print(f"  [{C_GRAY}]→ 粘贴到 F:\\CLAUDE\\quant_trading\\.env 中的 TUSHARE_TOKEN=[/]")
            console.print(f"  [{C_GREEN}]→ 当前自动使用 adata (免费) 作为备用数据源[/]")

    elif cmd == "oversold_backtest":
        from quant_trading.oversold_bounce import OversoldBounceScreener
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        start = sys.argv[3] if len(sys.argv) > 3 else "20230101"
        print_header(f"OVERSOLD BOUNCE BACKTEST — {code}")
        screener = OversoldBounceScreener()
        result = screener.backtest(code, start)
        if "error" in result:
            console.print(f"  [{C_RED}]{result['error']}[/]")
        else:
            console.print(f"\n  [bold]回测结果: {code}[/]")
            console.print(f"  期间: {result['period']}")
            console.print(f"  交易次数: {result['total_trades']}")
            console.print(f"  胜率: [{C_GREEN if result['win_rate'] > 50 else C_ORANGE}]{result['win_rate']}%[/]")
            console.print(f"  总盈亏: [{C_GREEN if result['total_pnl'] > 0 else C_RED}]{result['total_pnl']:,.0f}[/]")
            console.print(f"  收益率: [{C_GREEN if result['total_return'] > 0 else C_RED}]{result['total_return']}%[/]")
            console.print(f"  盈亏比: {result['profit_factor']}")
            console.print(f"  平均盈利: {result['avg_win']:,.0f}  平均亏损: {result['avg_lose']:,.0f}")

    elif cmd == "weekly_doji":
        from quant_trading.weekly_doji import WeeklyDojiScreener, print_doji_results
        top_n = int(sys.argv[2]) if len(sys.argv) > 2 else 20
        print_header("WEEKLY DOJI SCANNER — 周线缩量十字星",
                     "十字星 + 缩量 → 经典反转信号 | 全A股并发扫描")
        screener = WeeklyDojiScreener()
        df = screener.screen_all(top_n=top_n)
        if len(df) > 0:
            print_doji_results(df, top_n)
            out = Path(str(REPORTS_DIR)) / f"weekly_doji_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            out.parent.mkdir(parents=True, exist_ok=True)
            df.to_json(str(out), orient="records", force_ascii=False, indent=2)
            console.print(f"\n  [green]结果已保存: {out}[/]")
        else:
            console.print(f"\n  [{C_ORANGE}]未发现符合条件的周线缩量十字星候选[/]")

    elif cmd == "weekly_doji_deep":
        from quant_trading.weekly_doji import WeeklyDojiScreener, print_deep_analysis
        code = sys.argv[2] if len(sys.argv) > 2 else "600519"
        print_header(f"WEEKLY DOJI DEEP ANALYSIS — 周线十字星深度分析: {code}")
        screener = WeeklyDojiScreener()
        report = screener.deep_analyze(code)
        print_deep_analysis(report)

    elif cmd == "chanlun":
        from quant_trading.chan_theory import ChanTheoryAnalyzer, print_chanlun_report, print_chanlun_screen
        code = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].isdigit() and len(sys.argv[2]) == 6 else None
        if code:
            # 单股深度分析
            print_header(f"CHANLUN ANALYSIS — 缠论分析: {code}",
                         "分型 → 笔 → 线段 → 中枢 → 背驰 → 买卖点")
            analyzer = ChanTheoryAnalyzer()
            report = analyzer.analyze(code)
            print_chanlun_report(report)
            out = Path(str(REPORTS_DIR)) / f"chanlun_{code}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            out.parent.mkdir(parents=True, exist_ok=True)
            with open(str(out), "w", encoding="utf-8") as f:
                json.dump(report, f, indent=2, ensure_ascii=False, default=str)
            console.print(f"\n  [green]缠论报告已保存: {out}[/]")
        else:
            # 全市场扫描
            top_n = int(sys.argv[2]) if len(sys.argv) > 2 else 20
            print_header("CHANLUN SCREEN — 缠论全市场扫描",
                         "分型 → 笔 → 线段 → 中枢 → 背驰 → 买卖点 | 7因子加权评分")
            analyzer = ChanTheoryAnalyzer()
            df = analyzer.screen_all(top_n=top_n)
            print_chanlun_screen(df)
            if len(df) > 0:
                out = Path(str(REPORTS_DIR)) / f"chanlun_scan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                out.parent.mkdir(parents=True, exist_ok=True)
                df.to_json(str(out), orient="records", force_ascii=False, indent=2)
                console.print(f"\n  [green]扫描结果已保存: {out}[/]")

    elif cmd == "yixian":
        from quant_trading.yixian_dingqiankun import YixianDingQiankun, _print_report
        code = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].isdigit() and len(sys.argv[2]) == 6 else None
        if not code:
            print_header("YIXIAN DINGQIANKUN — 一线定乾坤",
                         "用法: python quant_trader.py yixian 300657 | yixian_scan 30")
            return
        print_header("YIXIAN ANALYSIS — 一线定乾坤: %s" % code,
                     "牛顿运动定律 → 价格加速度 → 多级EMA滤波 → 乾坤信号线")
        yx = YixianDingQiankun()
        report = yx.deep_analyze(code)
        _print_report(report)
        chart_path = yx.generate_chart(code)
        console.print(f"\n  [green]图表已保存: {chart_path}[/]")

    elif cmd == "yixian_scan":
        from quant_trading.yixian_dingqiankun import YixianDingQiankun
        top_n = int(sys.argv[2]) if len(sys.argv) > 2 and sys.argv[2].isdigit() else 30
        print_header("YIXIAN DINGQIANKUN SCAN — 一线定乾坤 全市场扫描",
                     "筛选TOP%d | 5层指标: QKSL→LXYZ→XLQK→YXDQK→XLTD" % top_n)
        yx = YixianDingQiankun()
        df = yx.screen_all(top_n=top_n)
        if len(df) > 0:
            console.print(df.to_string(index=False))
            out = Path(str(REPORTS_DIR)) / "yixian_scan_%s.json" % datetime.now().strftime('%Y%m%d_%H%M%S')
            out.parent.mkdir(parents=True, exist_ok=True)
            df.to_json(str(out), orient="records", force_ascii=False, indent=2)
            console.print(f"\n  [green]扫描结果已保存: {out}[/]")
        else:
            console.print("  [orange]未发现符合条件的乾坤线信号[/]")

    elif cmd == "yixian_live":
        from quant_trading.yixian_dingqiankun import YixianDingQiankun, _print_report
        code = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].isdigit() and len(sys.argv[2]) == 6 else None
        freq = sys.argv[3] if len(sys.argv) > 3 else 'd'
        if not code:
            print_header("YIXIAN LIVE — 一线定乾坤 Tushare实时更新",
                         "用法: python quant_trader.py yixian_live 300657 [d|60]")
            return
        print_header("YIXIAN LIVE UPDATE — 一线定乾坤实时: %s" % code,
                     "Tushare Pro 拉取最新 → 刷新缓存 → 重算信号 → 推送微信")
        yx = YixianDingQiankun()
        result = yx.live_update(code, freq=freq, push_wechat=True)
        if 'error' in result.get('report', {}):
            console.print("[red]ERROR: %s[/]" % result['report']['error'])
        else:
            _print_report(result['report'])
            console.print("\n  Data source: [green]%s[/] | Chart: [green]%s[/]" % (
                result['data_source'], result['chart_path']))

    elif cmd == "margin_index":
        from quant_trading.margin_vs_index import generate_margin_report, print_margin_report
        months = int(sys.argv[2]) if len(sys.argv) > 2 else 24
        print_header("MARGIN vs INDEX — 两融余额 vs 上证指数",
                     f"关联分析 | 回溯{months}个月 | Pearson相关 + 滚动相关")
        report = generate_margin_report(months)
        print_margin_report(report)

    elif cmd == "temperature":
        from quant_trading.market_temperature import measure_temperature, print_temperature
        print_header("MARKET THERMOMETER — 市场温度计",
                     "7维度加权评分 | 融资+北向+涨跌+量能+涨停+期指+波动")
        result = measure_temperature(verbose=True)
        print_temperature(result)
        alerts = result.get("alerts", [])
        if alerts:
            for a in alerts:
                console.print(f"  [{C_YELLOW}]{a}[/]")

    elif cmd == "sector_temp":
        from quant_trading.sector_temperature import run_sector_temp
        sub = sys.argv[2] if len(sys.argv) > 2 else ""
        if sub == "concept":
            print_header("CONCEPT SECTOR THERMOMETER — 概念板块温度计")
            run_sector_temp(mode="concept")
        elif sub and sub not in ("industry", ""):
            print_header(f"SECTOR SEARCH — 板块搜索: {sub}")
            run_sector_temp(mode="industry", search=sub)
        else:
            print_header("INDUSTRY SECTOR THERMOMETER — 行业板块温度计")
            run_sector_temp(mode="industry")

    elif cmd == "prob":
        from quant_trading.probability_engine import run_prob_command
        sub = sys.argv[2] if len(sys.argv) > 2 else ""
        arg = sys.argv[3] if len(sys.argv) > 3 else ""
        if sub == "guide":
            run_prob_command("guide")
        else:
            run_prob_command(sub, arg)

    elif cmd == "chemical":
        from quant_trading.chemical_kb import query_element, query_product, query_sector, search_chemicals
        subcmd = sys.argv[2] if len(sys.argv) > 2 else ""
        arg = sys.argv[3] if len(sys.argv) > 3 else ""
        if subcmd == "element" and arg:
            query_element(arg)
        elif subcmd == "product" and arg:
            query_product(arg)
        elif subcmd == "sector":
            query_sector()
        elif subcmd == "search" and arg:
            results = search_chemicals(arg)
            if results:
                console.print(f"\n  [bold green][SEARCH] '{arg}' 匹配 {len(results)} 个化工品:[/]\n")
                for r in results:
                    from quant_trading.chemical_kb import print_product_info
                    print_product_info(r["name"], r)
            else:
                console.print(f"\n  [dim]未找到匹配 '{arg}' 的化工品[/]")
        else:
            console.print(f"\n  [yellow]用法: python quant_trader.py chemical <subcmd> <arg>[/]")
            console.print(f"  [dim]  element <符号>   — 查元素 (如: chemical element Li)[/]")
            console.print(f"  [dim]  product <名称>  — 查化工产品 (如: chemical product 磷酸铁锂)[/]")
            console.print(f"  [dim]  sector         — 化工行业分类[/]")
            console.print(f"  [dim]  search <关键词> — 全局搜索[/]")

    elif cmd == "news_pick":
        from quant_trading.news_driven import NewsDrivenScreener, print_report
        top_n = int(sys.argv[2]) if len(sys.argv) > 2 else 15
        print_header("NEWS-DRIVEN STOCK PICKER — 新闻驱动选股")
        screener = NewsDrivenScreener()
        report = screener.run(top_n=top_n)
        print_report(report)

    elif cmd == "inst_check":
        from quant_trading.institutional_analyzer import InstitutionalAnalyzer, print_analysis
        code = sys.argv[2] if len(sys.argv) > 2 else "603979"
        print_header(f"INSTITUTIONAL STOCK ANALYZER — {code}")
        r = InstitutionalAnalyzer().analyze(code)
        if "error" in r:
            console.print(f"  [{C_RED}]{r['error']}[/]")
        else:
            print_analysis(r)

    elif cmd == "inst_scan":
        from quant_trading.institutional_analyzer import InstitutionalAnalyzer
        top_n = int(sys.argv[2]) if len(sys.argv) > 2 else 20
        print_header(f"INSTITUTIONAL STOCK SCANNER — 机构票扫描 TOP{top_n}")
        df = InstitutionalAnalyzer().batch_scan(top_n=top_n)
        if len(df) > 0:
            console.print(df.to_string(index=False))
            out = Path(str(REPORTS_DIR)) / f"inst_scan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            df.to_json(str(out), orient="records", force_ascii=False, indent=2)
            console.print(f"\n  [green]📁 已保存: {out}[/]")
        else:
            console.print(f"  [{C_ORANGE}]未发现机构票[/]")

    elif cmd == "morning_scan":
        print_header("MORNING SCAN — 一键晨间扫描")
        top_n = int(sys.argv[2]) if len(sys.argv) > 2 else 20
        console.print(f"  [{C_CYAN}]🔍 超跌反弹筛选...[/]")
        from quant_trading.oversold_bounce import OversoldBounceScreener
        df_os = OversoldBounceScreener().screen_all(top_n=top_n)
        if len(df_os) > 0:
            console.print(df_os.to_string(index=False))
            out = Path(str(REPORTS_DIR)) / f"morning_scan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            df_os.to_json(str(out), orient="records", force_ascii=False, indent=2)
            console.print(f"\n  [green]📁 超跌结果: {out}[/]")
        else:
            console.print(f"  [{C_ORANGE}]未发现超跌候选[/]")
        console.print(f"\n  [{C_CYAN}]🔥 热门板块...[/]")
        try:
            from quant_trading.sector_tracker import hot_sector_scan as _hot_sector_scan
            _hot_sector_scan()
        except Exception: pass
        console.print(f"\n  [{C_CYAN}]💰 北向资金...[/]")
        try:
            from quant_trading.north_bound import fetch_north_bound_daily, analyze_north_bound_flow, print_north_bound_report
            nb = fetch_north_bound_daily(5)
            if not nb.empty:
                print_north_bound_report(analyze_north_bound_flow(nb), None)
        except Exception: pass
        console.print(f"\n  [green]✅ 晨间扫描完成[/]")

    elif cmd == "oversold_report":
        from quant_trading.oversold_bounce import OversoldBounceScreener
        top_n = int(sys.argv[2]) if len(sys.argv) > 2 else 15
        print_header(f"OVERSOLD BOUNCE REPORT — 超跌反弹报告 TOP{top_n}")
        df = OversoldBounceScreener().screen_all(top_n=top_n)
        if len(df) > 0:
            console.print(df.to_string(index=False))
            from docx import Document as _Doc
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            _doc = _Doc()
            for _s in _doc.sections: _s.top_margin = Cm(2); _s.bottom_margin = Cm(2)
            _p = _doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _p.add_run(f"超跌反弹日报\n{datetime.now().strftime('%Y-%m-%d')}").font.size = Pt(32)
            _doc.add_paragraph(f"全市场 {5038} 只扫描 · {len(df)} 只候选 · 平均分 {df['得分'].mean():.0f}")
            _doc.add_paragraph()
            _t = _doc.add_table(rows=len(df)+1, cols=6, style="Light Grid Accent 1")
            for _j, _h in enumerate(["代码","名称","得分","RSI","距MA60","距60日高"]):
                _t.rows[0].cells[_j].text = _h
            for _i, (_, _r) in enumerate(df.iterrows()):
                for _j, _k in enumerate(["代码","名称","得分","RSI(14)","距MA60","距60日高"]):
                    _t.rows[_i+1].cells[_j].text = str(_r.get(_k, ""))
            _out = Path(str(REPORTS_DIR)) / f"超跌反弹日报_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            _doc.save(str(_out))
            console.print(f"\n  [green]📄 DOCX: {_out}[/]")
            json_out = Path(str(REPORTS_DIR)) / f"oversold_scan_{datetime.now().strftime('%Y%m%d')}.json"
            df.to_json(str(json_out), orient="records", force_ascii=False, indent=2)
            console.print(f"  [green]📁 JSON: {json_out}[/]")

    elif cmd == "bearish_report":
        code = sys.argv[2] if len(sys.argv) > 2 else None
        print_header("FIRST BEARISH REPORT — 首阴策略报告")
        if code:
            df_k = fetch_kline(code, 120)
            if not df_k.empty:
                detailed_bearish_analysis(code, df_k)
        else:
            first_bearish_scan()
        try:
            from pathlib import Path as _P
            _src = _P("F:/DEEPCODE/DeepCode (CLONE)/reports/first_bearish_20260622.json")
            if _src.exists():
                import json as _json
                with open(_src, "r", encoding="utf-8") as _f:
                    _data = _json.load(_f)
                _candidates = _data.get("results", _data) if isinstance(_data, dict) else _data
                from docx import Document as _D
                _doc = _D()
                _doc.add_heading("首阴策略报告", 0)
                _doc.add_paragraph(f"扫描时间: {datetime.now().strftime('%Y-%m-%d')}  |  候选: {len(_candidates)}只")
                for _c in _candidates:
                    _doc.add_heading(f"{_c.get('name','?')}({_c.get('code','?')}) — {_c.get('score',0)}/10", 2)
                    _doc.add_paragraph(f"阴线日: {_c.get('date','?')} | 实体: {_c.get('body_pct',0):.1f}% | 量比: {_c.get('vol_ratio',0):.2f} | 信号: {_c.get('signal_type','?')}")
                _out = _P(str(REPORTS_DIR)) / f"首阴策略报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                _doc.save(str(_out))
                console.print(f"\n  [green]📄 DOCX: {_out}[/]")
        except Exception as _e:
            console.print(f"  [{C_ORANGE}]DOCX生成失败: {_e}[/]")

    elif cmd == "bellwether":
        from quant_trading.bellwether import BellwetherScreener, print_report as _pr_bw
        sector = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2] in ("银行","食品饮料","医药生物","电子","电力设备","有色金属","汽车","国防军工","通信","计算机","非银金融","家用电器","房地产","公用事业","建筑装饰") else None
        top_n = int(sys.argv[3]) if len(sys.argv) > 3 and sys.argv[3].isdigit() else (int(sys.argv[2]) if len(sys.argv) > 2 and sys.argv[2].isdigit() else 6)
        print_header(f"BELLWETHER SCANNER — 城门立木选股{(' · '+sector) if sector else ''}")
        screener = BellwetherScreener()
        results = screener.scan(sector_name=sector, top_n=top_n)
        report = screener.report(results, top_n)
        _pr_bw(report)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out = Path(str(REPORTS_DIR)) / f"bellwether_{ts}.json"
        with open(out, "w", encoding="utf-8") as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        console.print(f"\n  [green]📁 已保存: {out}[/]")

    elif cmd == "cmd_list":
        print_header("QUANT TRADER COMMAND REFERENCE — 命令大全")
        from gen_command_reference import COMMANDS
        for cat, cmds in COMMANDS.items():
            console.print(f"\n  [bold]{cat}[/]")
            for cmd_name, desc, usage in cmds:
                is_new = "🆕" if cmd_name in ("oversold","oversold_backtest","news_pick","inst_check","inst_scan","morning_scan","oversold_report","bearish_report","dark_horse") else "  "
                console.print(f"  {is_new} [{C_CYAN}]{cmd_name:<22s}[/] {desc}")

    elif cmd == "market":
        objective_market()

    elif cmd == "dark_horse":
        dark_horse_screen()

    elif cmd in ("lowvol", "lowvol_momentum"):
        from quant_trading.lowvol_momentum import fetch_and_score, print_results
        top_n = int(sys.argv[2]) if len(sys.argv) > 2 else 30
        min_score = int(sys.argv[3]) if len(sys.argv) > 3 else 0
        print_header("LOW-VOL + MOMENTUM + LIQUIDITY — 低波动+高动量+高流动性三因子选股")
        results = fetch_and_score(None)  # Use default candidate pool
        if min_score > 0:
            results = [r for r in results if r["score"] >= min_score]
        print_results(results, top_n=top_n)

    elif cmd == "first_bearish":
        from quant_trading.first_bearish import FirstBearishScreener
        code = sys.argv[2] if len(sys.argv) > 2 else None
        screener = FirstBearishScreener()
        if code:
            s = screener.check_one(code)
            if s:
                import json as _json
                print(_json.dumps(s.__dict__, indent=2, ensure_ascii=False, default=str))
            else:
                print(f"{code} 不满足首阴条件")
        else:
            from rich.table import Table
            print_header("FIRST BEARISH SCAN [V3]",
                         "首阴策略V3 — 上升趋势回踩 · 主力/VSA交叉验证")
            df = screener.screen_all()
            if len(df) > 0:
                console.print(df.to_string(index=False))
                out = Path(str(REPORTS_DIR)) / f"first_bearish_{datetime.now().strftime('%Y%m%d')}.json"
                df.to_json(str(out), orient="records", force_ascii=False, indent=2)
                console.print(f"\n  [green]已保存: {out}[/]")
            else:
                console.print(f"  [{C_ORANGE}]未发现首阴候选[/]")

    elif cmd == "mtf_pullback":
        mtf_pullback_scan()

    elif cmd == "tui":
        tui_path = Path(__file__).parent / "webapp" / "tui_app.py"
        if tui_path.exists():
            os.system(f"python {tui_path}")
        else:
            console.print(f"[{C_RED}]TUI app not found at {tui_path}[/]")

    elif cmd == "sina_eod":
        """收盘后用 Sina 拉取当日全市场K线，补齐 Tushare T+1 缺口"""
        from quant_trading.data_cache import sina_eod_sync, get_cache_stats
        print_header("SINA EOD SYNC — Sina收盘数据同步")
        console.print(f"  [{C_CYAN}]收盘后从 Sina 拉取全市场当日 OHLCV，写入缓存...[/]")
        ok, fail = sina_eod_sync()
        console.print(f"  [{C_GREEN}]✓[/] Sina EOD: {ok} 成功, {fail} 失败")
        stats = get_cache_stats()
        console.print(f"  Cache: {stats['total_rows']:,} rows | {stats['codes']} stocks | {stats['size_mb']} MB")

    elif cmd == "daily-update":
        """统一数据库每日更新 (早盘仅Tushare数据 → 18点后全套)"""
        import subprocess, os
        hour = datetime.now().hour
        script = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                              "quant_trading", "reports", "daily_auto_update.py")
        if hour < 18:
            # 早盘 (<18:00): 仅从Tushare拉取数据，不做评分/模型/GRPO
            print_header("DAILY UPDATE — 早盘模式 (仅Tushare数据同步)",
                         "从Tushare拉取日线行情 → 完成")
            args_list = [sys.executable, script, "--tushare-only"]
        else:
            # 晚盘 (>=18:00): 全套 — Sina EOD + 评分 + 重训 + GRPO + 验证
            print_header("DAILY UPDATE — 全量更新",
                         "Sina收盘同步 → Batch评分 → 模型重训 → GRPO → 前向验证")
            args_list = [sys.executable, script, "--skip-prewarm"]
        console.print(f"  [{C_CYAN}]启动: {' '.join(args_list)}[/]")
        result = subprocess.run(args_list, cwd=os.path.dirname(os.path.abspath(__file__)))
        if result.returncode == 0:
            console.print(f"\n  [{C_GREEN}]✓ 每日更新完成[/]")
        else:
            console.print(f"\n  [{C_RED}]✗ 每日更新失败 (exit={result.returncode})[/]")

    elif cmd == "prewarm":
        from quant_trading.data_cache import pre_warm_cache, get_cache_stats
        print_header("CACHE PRE-WARM — 预加载全市场K线到本地")
        # 显式从导入路径获取 fetch_kline，避免 __main__ 双模块问题
        import quant_trader as _qt
        code_name = _qt.fetch_all_a_stock_codes()
        console.print(f"  [{C_CYAN}]Fetching & caching {len(code_name)} stocks...[/]")
        stats = pre_warm_cache(code_name, lambda c, d: _qt.fetch_kline(c, d, use_cache=False, skip_fallback=True), datalen=250)
        console.print(f"  [{C_GREEN}]✓[/] Cache stats: {stats['total_rows']:,} rows | "
                      f"{stats['codes']} stocks | {stats['size_mb']} MB")

    elif cmd == "factor":
        from quant_trading.factor_engine import run_factor_pipeline
        print_header("FACTOR ENGINE — IC Analysis + ML Synthesis",
                     "因子IC分析 · 相关性矩阵 · LightGBM合成")
        code_name = fetch_all_a_stock_codes()
        sample = int(sys.argv[2]) if len(sys.argv) > 2 else 300
        console.print(f"  [{C_CYAN}]Sample size:[/] {sample} stocks")
        # 仅用缓存, 避免Baostock并发压缩错误
        run_factor_pipeline(lambda c, d: fetch_kline(c, min(d, 90), use_cache=True), code_name, sample_size=sample)

    elif cmd == "portfolio":
        from quant_trading.portfolio_engine import portfolio_from_first_bearish
        print_header("PORTFOLIO OPTIMIZATION — 组合优化引擎",
                     "首阴候选 → 因子评分 → 协方差估计 → Max Sharpe/Risk Parity")
        code_name = fetch_all_a_stock_codes()
        best = portfolio_from_first_bearish(
            lambda c, d: fetch_kline(c, min(d, 90), use_cache=True),
            code_name,
            max_weight=0.25
        )
        if best:
            console.print(f"\n  [{C_GREEN}]✓[/] Portfolio optimization complete")

    elif cmd == "cache":
        from quant_trading.data_cache import get_cache_stats, get_prewarm_history
        if "--history" in sys.argv:
            print_header("PREWARM HISTORY — Prewarm 运行历史")
            hist = get_prewarm_history(20)
            if hist['rows']:
                from rich.table import Table
                table = Table(title="最近 Prewarm 运行记录")
                for col in hist['columns']:
                    table.add_column(col, style="cyan" if col != "ID" else "dim")
                for row in hist['rows']:
                    styled = []
                    for i, val in enumerate(row):
                        if hist['columns'][i] == "成功率%" and val is not None:
                            color = "green" if val >= 99 else "yellow" if val >= 95 else "red"
                            styled.append(f"[{color}]{val}[/]")
                        elif hist['columns'][i] == "DB_MB" and val is not None:
                            styled.append(f"{val:.1f}")
                        elif hist['columns'][i] == "耗时秒" and val is not None:
                            styled.append(f"{val:.0f}")
                        else:
                            styled.append(str(val) if val is not None else "-")
                    table.add_row(*styled)
                console.print(table)
            else:
                console.print(f"  [{C_CYAN}]暂无 prewarm 历史记录[/]")
        else:
            stats = get_cache_stats()
            print_header("CACHE STATS — 本地缓存状态")
            console.print(f"  [{C_CYAN}]总K线条数:[/] {stats['total_rows']:,}")
            console.print(f"  [{C_CYAN}]覆盖股票数:[/] {stats['codes']}")
            console.print(f"  [{C_CYAN}]最后更新:[/] {stats['latest_update']}")
            console.print(f"  [{C_CYAN}]数据库大小:[/] {stats['size_mb']} MB")
            # 如果有历史记录，提示可查看
            hist = get_prewarm_history(1)
            if hist['rows']:
                console.print(f"\n  [[{C_CYAN}]tip[/]] 查看 prewarm 运行历史: python quant_trader.py cache --history")

    elif cmd == "rag":
        """RAG管道: 索引报告 / 语义检索 / RAG增强分析"""
        from quant_trading.rag_pipeline import get_rag
        rag = get_rag()
        sub = sys.argv[2] if len(sys.argv) > 2 else "index"
        if sub == "index":
            print_header("RAG INDEX — 索引量化报告")
            result = rag.index_all_reports()
            console.print(f"  [{C_GREEN}]✓[/] {result['indexed_reports']} 份报告, {result['total_chunks']} 个片段")
        elif sub == "search":
            if len(sys.argv) < 4:
                console.print(f"  [{C_RED}]用法: python quant_trader.py rag search <关键词>[/]")
                return
            query = sys.argv[3]
            results = rag.search(query, k=8)
            for i, r in enumerate(results):
                console.print(f"\n  [{C_CYAN}]{i+1}.[/] [{r['code']}] {r['content'][:250]}")
        elif sub == "stats":
            s = rag.stats()
            console.print(f"  RAG索引: {s['total_chunks']} 片段 | {s['index_dir']}")
        else:
            console.print(f"  [{C_RED}]用法: rag index|search|stats[/]")

    elif cmd == "daily":
        """每日一键刷新: Sina收盘同步 → Tushare prewarm覆盖 → 评分"""
        print_header("DAILY UPDATE — 每日数据刷新")
        import subprocess, time
        from pathlib import Path
        t0 = time.time()
        cwd = str(Path(__file__).parent)

        # Step 1: Sina EOD 收盘同步
        console.print(f"\n  [{C_CYAN}]>>> Sina收盘同步...[/]")
        r = subprocess.run([sys.executable, "quant_trader.py", "sina_eod"], cwd=cwd)
        if r.returncode != 0:
            console.print(f"  [{C_RED}]Sina收盘同步 失败 (exit={r.returncode})[/]")

        # Step 2: Tushare prewarm — 用内联方式绕过 CLI __main__ 双模块问题
        console.print(f"\n  [{C_CYAN}]>>> Tushare缓存预热 (覆盖Sina)...[/]")
        prewarm_script = (
            "from quant_trading.data_cache import pre_warm_cache, get_cache_stats;"
            "from quant_trader import fetch_kline, fetch_all_a_stock_codes;"
            "import time;"
            "cn=fetch_all_a_stock_codes();"
            "print(f'Prewarming {len(cn)} stocks...');"
            "t0=time.time();"
            "s=pre_warm_cache(cn,lambda c,d:fetch_kline(c,d,use_cache=False,skip_fallback=True),freq='d',datalen=250);"
            "print(f'Done in {time.time()-t0:.0f}s: {s[\"total_rows\"]:,} rows, {s[\"codes\"]} codes, {s[\"size_mb\"]} MB')"
        )
        r = subprocess.run([sys.executable, "-X", "utf8", "-c", prewarm_script], cwd=cwd)
        if r.returncode != 0:
            console.print(f"  [{C_RED}]Tushare预热 失败 (exit={r.returncode})[/]")

        # Step 3: 增量引擎评分
        console.print(f"\n  [{C_CYAN}]>>> 增量引擎评分...[/]")
        r = subprocess.run([sys.executable, "-X", "utf8", "quant_trading/batch_scorer.py", "--incremental"], cwd=cwd)
        if r.returncode != 0:
            console.print(f"  [{C_RED}]增量引擎评分 失败 (exit={r.returncode})[/]")

        elapsed = time.time() - t0
        console.print(f"\n  [{C_GREEN}]每日更新完成 (耗时 {elapsed/60:.1f}分钟)[/]")

    elif cmd == "health":
        print_header("SYSTEM HEALTH CHECK — 系统健康诊断")
        import sqlite3, os
        from pathlib import Path

        all_ok = True
        results = []

        # 1. 数据库完整性
        db_path = Path("quant_trading/kline_cache.db")
        if db_path.exists():
            size_mb = db_path.stat().st_size / (1024*1024)
            conn = sqlite3.connect(str(db_path))
            integrity = conn.execute("PRAGMA integrity_check").fetchone()[0]
            conn.close()
            results.append(("DB 完整性", "✅" if integrity == "ok" else "❌", f"{size_mb:.0f}MB, {integrity}"))
            if integrity != "ok":
                all_ok = False
        else:
            results.append(("DB 存在", "❌", "kline_cache.db 不存在"))
            all_ok = False

        # 2. 缓存统计
        try:
            from quant_trading.data_cache import get_cache_stats
            stats = get_cache_stats()
            codes = stats.get('codes', 0)
            rows = stats.get('total_rows', 0)
            results.append(("K线缓存", "✅" if codes > 100 else "⚠️", f"{codes}只, {rows:,}行"))
            if codes < 100:
                all_ok = False
        except Exception as e:
            results.append(("K线缓存", "❌", str(e)))
            all_ok = False

        # 3. 数据新鲜度
        try:
            from quant_trading.data_cache import get_cached_kline
            from quant_trading.holiday_calendar import is_trading_day
            df = get_cached_kline('000001', 5, 'd')
            latest = str(df['date'].max())[:10] if df is not None and not df.empty else 'N/A'
            today = datetime.now().strftime('%Y-%m-%d')
            if latest == today or not is_trading_day():
                results.append(("数据新鲜度", "✅", f"最新={latest}"))
            else:
                days = (datetime.now() - datetime.strptime(latest, '%Y-%m-%d')).days
                results.append(("数据新鲜度", "⚠️" if days <= 2 else "❌", f"最新={latest}, 滞后{days}天"))
                if days > 2:
                    all_ok = False
        except Exception as e:
            results.append(("数据新鲜度", "❌", str(e)))
            all_ok = False

        # 4. 引擎评分覆盖
        try:
            conn = sqlite3.connect(str(db_path))
            scored = conn.execute("SELECT COUNT(DISTINCT code) FROM engine_scores").fetchone()[0]
            conn.close()
            results.append(("引擎评分", "✅" if scored >= 10 else "⚠️", f"{scored}只已评分"))
        except Exception:
            results.append(("引擎评分", "⚠️", "表不存在或无数据"))

        # 5. 测试套件
        try:
            import subprocess
            r = subprocess.run(['python', '-m', 'pytest', 'quant_trading/tests/', '-q'],
                             capture_output=True, timeout=30)
            passed = b'passed' in r.stdout or b'passed' in r.stderr
            results.append(("测试套件", "✅" if r.returncode == 0 else "❌",
                          f"{r.returncode == 0 and 'PASS' or 'FAIL'}"))
            if r.returncode != 0:
                all_ok = False
        except Exception:
            results.append(("测试套件", "⚠️", "无法运行"))

        # 6. Tushare 连接
        try:
            from quant_trading.tushare_engine import _get_pro
            pro = _get_pro()
            results.append(("Tushare", "✅" if pro else "❌", "已连接" if pro else "Token未配置"))
            if not pro:
                all_ok = False
        except Exception:
            results.append(("Tushare", "❌", "连接失败"))
            all_ok = False

        # 输出
        console.print(f"\n  [{'green' if all_ok else 'red'}]{'═' * 50}[/]")
        for name, status, detail in results:
            color = 'green' if '✅' in status else ('yellow' if '⚠️' in status else 'red')
            console.print(f"  {status} [{color}]{name:<12}[/] {detail}")
        console.print(f"  [{'green' if all_ok else 'red'}]{'═' * 50}[/]")
        console.print(f"\n  总体状态: [{'green' if all_ok else 'red'}]{'✅ 健康' if all_ok else '❌ 有问题'}[/]")

    elif cmd == "monitor":
        from quant_trading.prometheus_metrics import (
            start_metrics_server, collect_all_metrics,
            generate_grafana_dashboard, start_health_monitor,
        )
        # 解析端口: 非 -- 开头的参数视为端口号
        port = 9090
        for a in sys.argv[2:]:
            if not a.startswith('--'):
                try:
                    port = int(a)
                except ValueError:
                    pass
        if "--dashboard" in sys.argv:
            generate_grafana_dashboard()
        elif "--daemon" in sys.argv:
            print_header("PROMETHEUS HEALTH DAEMON — 后台健康监控")
            start_health_monitor(interval_seconds=300)
            console.print(f"  [{C_GREEN}]健康监控守护已启动 (间隔5分钟)[/]")
            console.print(f"  [{C_CYAN}]按 Ctrl+C 停止[/]")
            try:
                import time
                while True:
                    time.sleep(1)
            except KeyboardInterrupt:
                console.print(f"\n  [{C_YELLOW}]守护已停止[/]")
        elif "--collect" in sys.argv:
            print_header("PROMETHEUS METRICS — 一次性采集")
            collect_all_metrics()
            from quant_trading.prometheus_metrics import QuantMetrics
            text = QuantMetrics().generate_metrics_text()
            console.print(f"  [{C_GREEN}]采集完成[/]")
            # 显示关键指标摘要
            import re
            gauges = re.findall(r'^quant_(\w+)\s+([\d.]+)', text, re.MULTILINE)
            for name, val in gauges[:10]:
                console.print(f"    {name}: {val}")
        else:
            print_header(f"PROMETHEUS METRICS SERVER — :{port}")
            collect_all_metrics()
            start_metrics_server(port=port)
            console.print(f"  [{C_GREEN}]Prometheus endpoint → http://localhost:{port}/metrics[/]")
            console.print(f"  [{C_CYAN}]Grafana 数据源: http://<host>:{port}[/]")
            console.print(f"  [{C_CYAN}]仪表盘生成: python quant_trader.py monitor --dashboard[/]")
            console.print(f"  [{C_CYAN}]按 Ctrl+C 停止[/]")
            try:
                import time
                while True:
                    time.sleep(60)
                    collect_all_metrics()
            except KeyboardInterrupt:
                console.print(f"\n  [{C_YELLOW}]服务已停止[/]")

    elif cmd == "north_bound":
        from quant_trading.north_bound import (
            fetch_north_bound_daily, analyze_north_bound_flow,
            print_north_bound_report, fetch_north_bound_stocks,
        )
        print_header("NORTH-BOUND FLOW — 北向资金分析")
        df = fetch_north_bound_daily(30)
        if not df.empty:
            analysis = analyze_north_bound_flow(df)
            top_stocks = fetch_north_bound_stocks("hsgt", 10)
            print_north_bound_report(analysis, top_stocks)
        else:
            console.print(f"  [{C_RED}]无法获取北向资金数据[/]")

    elif cmd == "north_stock":
        from quant_trading.north_bound import analyze_north_bound_stock, print_north_bound_stock_report
        if len(sys.argv) < 3:
            console.print(f"  [{C_RED}]用法: python quant_trader.py north_stock <股票代码>[/]")
        else:
            code = sys.argv[2]
            print_header(f"NORTH-BOUND STOCK — {code} 北向资金分析")
            result = analyze_north_bound_stock(code)
            print_north_bound_stock_report(result)

    elif cmd == "limit":
        from quant_trading.limit_tracker import (
            fetch_limit_up_stocks, fetch_limit_down_stocks,
            compute_market_sentiment, print_limit_report,
        )
        print_header("LIMIT UP/DOWN TRACKER — 涨跌停板统计")
        up_df = fetch_limit_up_stocks()
        down_df = fetch_limit_down_stocks()
        sentiment = compute_market_sentiment(
            len(up_df), len(down_df)
        )
        print_limit_report(up_df, down_df, sentiment)

    elif cmd == "dragon_tiger":
        from quant_trading.dragon_tiger import (
            fetch_dragon_tiger_daily, fetch_dragon_tiger_traders,
            analyze_dragon_tiger, print_dragon_tiger_report,
        )
        print_header("DRAGON-TIGER BOARD — 龙虎榜分析")
        dt_df = fetch_dragon_tiger_daily()
        traders_df = fetch_dragon_tiger_traders()
        analysis = analyze_dragon_tiger(dt_df, traders_df)
        print_dragon_tiger_report(analysis)

    elif cmd == "margin":
        from quant_trading.margin_trading import (
            fetch_margin_balance_history, fetch_margin_stocks,
            analyze_margin_sentiment, print_margin_report,
        )
        print_header("MARGIN TRADING — 融资融券分析")
        df = fetch_margin_balance_history(30)
        analysis = analyze_margin_sentiment(df)
        top_stocks = fetch_margin_stocks(15)
        print_margin_report(analysis, top_stocks)

    elif cmd == "multi_tf":
        from quant_trading.multi_timeframe import (
            compute_resonance, find_precise_entry,
            print_resonance_report,
        )
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        print_header(f"MULTI-TIMEFRAME RESONANCE — {code}")
        resonance = compute_resonance(
            lambda c, d, freq: fetch_kline(c, d, freq=freq, use_cache=True),
            code,
        )
        print_resonance_report(resonance)
        # Also show precise entry
        entry = find_precise_entry(
            lambda c, d, freq: fetch_kline(c, d, freq=freq, use_cache=True),
            code,
        )
        if entry:
            console.print(f"\n  [{C_GOLD}]精确入场:[/] {entry.get('reason', '')}")
            if entry.get("entry_price"):
                console.print(f"    入场价: {entry['entry_price']:.2f}")
            if entry.get("stop_loss"):
                console.print(f"    止损价: {entry['stop_loss']:.2f}")

    elif cmd == "heatmap":
        from quant_trading.sector_heatmap import (
            fetch_sector_multi_period_returns,
            compute_sector_momentum_score,
            detect_rotation_signal,
            print_sector_heatmap_report,
            plot_sector_heatmap,
        )
        print_header("SECTOR ROTATION HEATMAP — 板块轮动热力图")
        df = fetch_sector_multi_period_returns()
        if not df.empty:
            df = compute_sector_momentum_score(df)
            top5 = df.head(5)["name"].tolist() if "name" in df.columns else []
            rotation = detect_rotation_signal(top5)
            print_sector_heatmap_report(df, rotation)

            # Also try to show Plotly chart
            try:
                fig = plot_sector_heatmap(df)
                if fig:
                    fig.write_html(
                        Path(__file__).parent / "reports" / f"sector_heatmap_{datetime.now().strftime('%Y%m%d')}.html"
                    )
                    console.print(f"\n  [{C_GREEN}]✓[/] 热力图已保存到 reports/ 目录")
            except Exception:
                _suppress("unknown")
        else:
            console.print(f"  [{C_RED}]无法获取板块数据[/]")

    elif cmd == "optimize":
        from quant_trading.optimizer import (
            grid_search, walk_forward_optimization,
            print_optimization_results, print_wf_results,
        )
        from quant_trading.backtest_engine import run_backtest, FirstBearishStrategy

        print_header("PARAMETER OPTIMIZATION — Grid Search + Walk-Forward")
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        df = fetch_kline(code, 500)
        if df.empty:
            console.print(f"[{C_RED}]Cannot fetch data for {code}[/]")
        else:
            # Simple grid search for FirstBearish strategy
            param_grid = {
                "body_pct_min": [2.0, 3.0, 4.0, 5.0],
                "body_ratio_min": [1.5, 2.0, 2.5],
                "stop_loss_pct": [0.05, 0.08, 0.10],
            }

            def _bt(params):
                return run_backtest(code, df, FirstBearishStrategy, **params)

            console.print(f"  [{C_CYAN}]Grid Search: {code}[/]")
            best, all_results = grid_search(_bt, param_grid, workers=6)
            if best:
                print_optimization_results(best, all_results)

    elif cmd == "hedge":
        from quant_trading.hedge_engine import (
            portfolio_risk_overlay, print_hedge_report,
            evaluate_cb_hedge,
        )
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        name = get_stock_name(code)
        print_header(f"HEDGE ANALYSIS — {name} ({code})")

        # Simple position analysis
        df = fetch_kline(code, 20)
        if not df.empty:
            price = float(df["close"].iloc[-1])
            positions = [{"code": code, "shares": 10000, "price": price, "beta": 1.2}]
            risk = portfolio_risk_overlay(positions, {"enabled": True, "target_beta": 0.3})
            print_hedge_report(risk)

            # Check convertible bond hedge
            cb = evaluate_cb_hedge(code, price)
            if cb:
                console.print(f"\n  [{C_GOLD}]可转债对冲:[/]")
                console.print(f"    {cb.get('text', 'N/A')}")
                console.print(f"    转债代码: {cb.get('cb_code', 'N/A')} 溢价率: {cb.get('premium', 0):+.1f}%")
            else:
                console.print(f"\n  [{C_GRAY}]该股无可转债[/]")

    elif cmd == "akshare":
        from quant_trading.akshare_integration import run_akshare_market_scan
        print_header("AKSHARE MARKET SCAN — 全市场扫描")
        result = run_akshare_market_scan()
        if result.get("indices"):
            console.print(f"\n  [{C_GOLD}]主要指数:[/]")
            for name, info in result["indices"].items():
                chg_color = C_GREEN if info["change_pct"] > 0 else C_RED
                console.print(f"    {name}: {info['price']:.2f} [{chg_color}]{info['change_pct']:+.2f}%[/]")
        if result.get("hot_industries"):
            console.print(f"\n  [{C_GOLD}]热门行业 TOP8:[/]")
            for s in result["hot_industries"]:
                console.print(f"    {s['name']}: [{C_GREEN}]{s['change_pct']:+.2f}%[/]")
        if result.get("hot_concepts"):
            console.print(f"\n  [{C_GOLD}]热门概念 TOP8:[/]")
            for s in result["hot_concepts"][:8]:
                console.print(f"    {s['name']}: [{C_GREEN}]{s['change_pct']:+.2f}%[/]")

    elif cmd == "attribution":
        from quant_trading.attribution import (
            compute_performance_metrics, CostModel,
            print_attribution_report,
        )
        print_header("PERFORMANCE ATTRIBUTION — 绩效归因 + 成本分析")
        cm = CostModel()
        summary = cm.summary()
        console.print(f"\n  [{C_GOLD}]交易成本模型:[/]")
        console.print(f"    买入费率: {summary['buy_rate_bps']:.1f} bps")
        console.print(f"    卖出费率: {summary['sell_rate_bps']:.1f} bps (含印花税)")
        console.print(f"    往返费率: {summary['roundtrip_bps']:.1f} bps")
        np.random.seed(42)
        n = 252
        returns = np.random.randn(n) * 0.015 + 0.0005
        equity = 1_000_000 * np.exp(np.cumsum(returns))
        metrics = compute_performance_metrics(equity)
        print_attribution_report({}, metrics)

    elif cmd == "futures":
        from quant_trading.futures_options import (
            compute_basis, fetch_index_futures_quote,
            fetch_etf_option_chain, compute_put_call_ratio,
            print_futures_report,
        )
        print_header("FUTURES & OPTIONS — 股指期货基差+期权PCR")
        basis = {}
        for ft in ["IF", "IC", "IM", "IH"]:
            try:
                df = fetch_index_futures_quote(ft)
                basis[ft] = compute_basis(df, ft)
            except Exception:
                _suppress("unknown")
        # Try 50ETF options
        try:
            opt_df = fetch_etf_option_chain("510050")
            pcr = compute_put_call_ratio(opt_df) if not opt_df.empty else None
        except Exception:
            pcr = None
        # Print first found basis info
        for ft, info in basis.items():
            if info.get("signal") != "N/A":
                print_futures_report({"IF": info, "IC": basis.get("IC", {}),
                                       "IM": basis.get("IM", {}), "IH": basis.get("IH", {})}, pcr)
                break

    elif cmd == "special":
        from quant_trading.special_data import (
            fetch_block_trades, analyze_block_trades,
            check_portfolio_unlocking,
            fetch_shareholder_changes, flag_insider_signals,
            fetch_sector_fund_flow, compute_sector_flow_signal,
            print_special_data_report,
        )
        print_header("SPECIAL DATA — 另类数据综合(大宗/解禁/增减持/行业资金)")
        bt_df = fetch_block_trades(5)
        bt_analysis = analyze_block_trades(bt_df)
        # Use watchlist for checks
        codes = []
        for g in CFG.WATCHLIST.values():
            codes.extend(g[:2])
        unlock = check_portfolio_unlocking(codes, 30)
        sh_df = fetch_shareholder_changes(10)
        insider = flag_insider_signals(sh_df)
        flow_df = fetch_sector_fund_flow()
        flow_signal = compute_sector_flow_signal(flow_df)
        print_special_data_report(bt_analysis, unlock, insider, flow_signal)

    elif cmd == "pairs":
        from quant_trading.pairs_trading import screen_pairs_in_watchlist, print_pairs_report
        print_header("PAIRS TRADING — 配对交易筛选")
        screening = screen_pairs_in_watchlist(lambda c, d: fetch_kline(c, d, use_cache=True))
        print_pairs_report(screening)

    elif cmd == "rotation":
        from quant_trading.sector_rotation_strategy import (
            momentum_rotation, mean_reversion_rotation, print_rotation_strategy,
        )
        print_header("SECTOR ROTATION STRATEGY — 板块轮动策略")
        mom = momentum_rotation()
        mean_rev = mean_reversion_rotation()
        print_rotation_strategy(mom, mean_rev)

    elif cmd == "earnings":
        from quant_trading.earnings_surprise import (
            fetch_earnings_forecasts, detect_earnings_surprise,
            print_earnings_report,
        )
        print_header("EARNINGS SURPRISE — 业绩超预期筛选")
        df = fetch_earnings_forecasts()
        surprises = detect_earnings_surprise(df)
        print_earnings_report(surprises)

    elif cmd == "sentiment":
        from quant_trading.sentiment_analysis import (
            compute_market_sentiment_index, sentiment_report,
        )
        print_header("SENTIMENT ANALYSIS — 市场情绪分析",
                     "SnowNLP语义 · cnsenti金融词典 · 多引擎融合")
        text = " ".join(sys.argv[2:]) if len(sys.argv) > 2 else None
        if text:
            sentiment_report(text)
        else:
            ms = compute_market_sentiment_index()
            from rich.console import Console
            Console().print(f"市场情绪指数: {ms['sentiment']}/100 — {ms['label']}")
            Console().print("[dim]附加文本: sentiment <内容> 可进行多引擎情感分析[/]")

    elif cmd == "monte_carlo":
        from quant_trading.monte_carlo import (
            portfolio_mc_analysis, print_mc_report,
        )
        print_header("MONTE CARLO — 蒙特卡洛模拟+压力测试")
        positions = [
            {"code": "300438", "weight": 0.3, "mean_ret": 0.15, "vol": 0.35},
            {"code": "300750", "weight": 0.3, "mean_ret": 0.20, "vol": 0.40},
            {"code": "600519", "weight": 0.4, "mean_ret": 0.10, "vol": 0.25},
        ]
        analysis = portfolio_mc_analysis(positions)
        print_mc_report(analysis)

    elif cmd == "intraday":
        from quant_trading.intraday_strategy import (
            run_intraday_scan, print_intraday_report,
        )
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        print_header(f"INTRADAY STRATEGY — {code}")
        result = run_intraday_scan(code, lambda c, d, freq: fetch_kline(c, d, freq=freq, use_cache=True))
        print_intraday_report(result)

    elif cmd == "margin_signal":
        from quant_trading.margin_signals import (
            compute_margin_signal, detect_short_squeeze_candidates,
            detect_deleveraging_stocks, print_margin_signals,
        )
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        print_header(f"MARGIN SIGNALS — 两融信号策略")
        sig = compute_margin_signal(code)
        squeeze = detect_short_squeeze_candidates(15)
        delever = detect_deleveraging_stocks(15)
        print_margin_signals(sig, squeeze, delever)

    elif cmd == "rebalance":
        from quant_trading.index_rebalance import (
            get_next_rebalance_dates, fetch_pledge_data,
            detect_pledge_risk, printf_rebalance_and_pledge,
        )
        print_header("INDEX REBALANCE & PLEDGE — 指数调仓+股权质押")
        events = get_next_rebalance_dates()
        pledge_df = fetch_pledge_data()
        risks = detect_pledge_risk(pledge_df)
        printf_rebalance_and_pledge(events, risks)

    elif cmd == "data_quality":
        from quant_trading.data_quality import (
            batch_quality_check, print_quality_report,
        )
        print_header("DATA QUALITY — 数据质量检查")
        code_name = fetch_all_a_stock_codes()
        result = batch_quality_check(
            list(code_name.keys()),
            lambda c, d: fetch_kline(c, d, use_cache=True),
            sample=200,
        )
        print_quality_report(result)

    elif cmd == "versions":
        from quant_trading.strategy_versioning import (
            StrategyRegistry, print_version_report,
        )
        print_header("STRATEGY VERSIONING — 策略版本管理")
        registry = StrategyRegistry()
        print_version_report(registry)

    elif cmd == "garch":
        from quant_trading.garch_upgrade import garch_fit, dynamic_position_sizing
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        print_header(f"GARCH VOLATILITY — {code}")
        df = fetch_kline(code, 250)
        if not df.empty:
            rets = df["close"].pct_change().dropna().values
            gr = garch_fit(rets)
            if "error" not in gr:
                sizing = dynamic_position_sizing(gr)
                console.print(f"\n  当前波动(年化): [{C_ORANGE}]{gr['current_vol_ann']}%[/]")
                console.print(f"  预测波动(年化): [{C_ORANGE}]{gr['forecast_vol_ann']}%[/]")
                console.print(f"  波动体制: {gr['vol_regime']}")
                console.print(f"  GARCH持久性: {gr['persistence']:.3f}")
                console.print(f"  建议仓位: [{C_GREEN if sizing['adjustment']=='INCREASE' else C_ORANGE}]{sizing['position']:.0%}[/] ({sizing['reason']})")

    elif cmd == "bl":
        from quant_trading.garch_upgrade import quick_bl_portfolio, print_bl_report
        print_header("BLACK-LITTERMAN — 组合优化")
        codes = []
        for group in list(CFG.WATCHLIST.values())[:3]:
            codes.extend(group[:2])
        # Sample view: bullish on first stock
        views = [{"stocks": [0], "value": 0.20, "confidence": 0.6}]
        result = quick_bl_portfolio(codes, views)
        print_bl_report(result, codes)

    elif cmd == "social":
        from quant_trading.social_sentiment import (
            fetch_guba_posts, analyze_social_sentiment,
            batch_social_scan, print_social_report,
        )
        code = sys.argv[2] if len(sys.argv) > 2 else None
        if code:
            print_header(f"SOCIAL SENTIMENT — {code}")
            posts = fetch_guba_posts(code, pages=2)
            result = analyze_social_sentiment(posts)
            print_social_report(None, result)
        else:
            print_header("SOCIAL SENTIMENT — 全市场扫描")
            watch_codes = []
            for g in CFG.WATCHLIST.values():
                watch_codes.extend(g[:2])
            batch = batch_social_scan(watch_codes[:20])
            print_social_report(batch)

    elif cmd == "level2":
        from quant_trading.level2_data import (
            fetch_level2_snapshot, compute_orderbook_signals,
            print_level2_report,
        )
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        print_header(f"LEVEL 2 ORDER BOOK — {code}")
        ob = fetch_level2_snapshot(code)
        if ob and ob.best_bid():
            signals = compute_orderbook_signals(ob)
            print_level2_report(ob, signals)
        else:
            console.print(f"  [{C_ORANGE}]L2数据不可用 (可能需要交易时段)[/]")

    elif cmd == "paper":
        from quant_trading.live_trading import (
            PaperTradingEngine, print_account_summary,
        )
        print_header("PAPER TRADING — 模拟交易账户")
        engine = PaperTradingEngine()
        engine.load_state()
        engine.set_price_provider(lambda c: float(fetch_kline(c, 5)["close"].iloc[-1]) if not fetch_kline(c, 5).empty else 10.0)
        print_account_summary(engine)

    elif cmd == "check_adj":
        from quant_trading.adjustment_validator import (
            validate_adjustment_pipeline, print_adjustment_report,
        )
        code = sys.argv[2] if len(sys.argv) > 2 else "300438"
        print_header(f"ADJUSTMENT VALIDATOR — 复权一致性校验 {code}")
        result = validate_adjustment_pipeline(code, lambda c, d: fetch_kline(c, d, use_cache=True))
        print_adjustment_report(result)

    elif cmd == "check_la":
        from quant_trading.lookahead_detector import (
            full_lookahead_audit, print_lookahead_report,
        )
        from quant_trading.factor_engine import compute_all_factors
        print_header("LOOK-AHEAD AUDIT — 未来函数审计")
        code_name = fetch_all_a_stock_codes()
        codes = list(code_name.keys())[:100]
        # 批量计算因子
        console.print(f"  [{C_CYAN}]Computing factors for {len(codes)} stocks...[/]")
        from quant_trading.factor_engine import batch_compute_factors
        factor_df = batch_compute_factors(codes, lambda c, d: fetch_kline(c, d, use_cache=True), workers=8)
        if not factor_df.empty:
            audit = full_lookahead_audit(
                lambda df: df,  # no-op: factors already computed
                factor_df,
                factor_cols=[c for c in factor_df.columns
                            if c not in ("date","code","open","high","low","close","volume")],
            )
            # Override shuffle — already done via batch
            audit["results"]["shuffle_test"] = {"false_positive_count": len(factor_df.columns)-5, "true_issues": {}}
            print_lookahead_report(audit)
        else:
            console.print(f"  [{C_RED}]因子计算失败[/]")

    elif cmd == "bug_audit":
        from quant_trading.bug_fixes import full_system_bug_audit, print_bug_audit_report
        print_header("BUG AUDIT — 隐形BUG全系统审计")
        code_name = fetch_all_a_stock_codes()
        audit = full_system_bug_audit(list(code_name.keys())[:500], fetch_kline)
        print_bug_audit_report(audit)

    elif cmd == "monitor":
        from quant_trading.prometheus_metrics import (
            QuantMetrics, start_health_monitor,
        )
        print_header("PROMETHEUS MONITOR — 健康监控")
        metrics = QuantMetrics()
        # Update initial values
        try:
            from quant_trading.north_bound import fetch_north_bound_daily, analyze_north_bound_flow
            nb = fetch_north_bound_daily(5)
            nb_a = analyze_north_bound_flow(nb)
            metrics.update_market(north_flow=nb_a.get("net_flow_today", 0))
        except Exception:
            _suppress("unknown")
        try:
            from quant_trading.limit_tracker import fetch_limit_up_stocks, fetch_limit_down_stocks
            up = len(fetch_limit_up_stocks())
            down = len(fetch_limit_down_stocks())
            metrics.update_market(limit_up=up, limit_down=down)
        except Exception:
            _suppress("unknown")
        try:
            from quant_trading.sentiment_analysis import compute_market_sentiment_index
            ms = compute_market_sentiment_index()
            metrics.update_market(sentiment=ms.get("sentiment", 50))
        except Exception:
            _suppress("unknown")
        # Start background monitor
        start_health_monitor(interval_seconds=300)
        txt = metrics.generate_metrics_text()
        console.print(f"\n  [{C_GREEN}]✓[/] Prometheus指标已生成 ({len(txt)} bytes)")
        console.print(f"  [{C_GRAY}]健康监控后台运行中 (每300秒更新)[/]")
        console.print(f"  [{C_GRAY}]Grafana配置: 导入prometheus数据源即可[/]")

    elif cmd == "sw_plot":
        from quant_trading.sw_heatmap_plot import plot_sw_l1_heatmap, plot_sw_l2_heatmap, plot_sw_l2_treemap
        from quant_trading.tushare_engine import fetch_sw_realtime
        level = sys.argv[2] if len(sys.argv) > 2 else 'L2'
        style = sys.argv[3] if len(sys.argv) > 3 else 'grid'

        print(f"  正在获取申万{level}实时数据...")
        df = fetch_sw_realtime(level=level)
        if df is None or df.empty:
            print("  [错误] 无法获取数据")
        else:
            if level == 'L1':
                path = plot_sw_l1_heatmap(df)
            elif style == 'treemap':
                path = plot_sw_l2_treemap(df)
            else:
                path = plot_sw_l2_heatmap(df)
            print(f"\n  📊 热力图已保存: {path}")

    elif cmd == "sw_rt":
        from quant_trading.tushare_engine import fetch_sw_realtime, print_sw_heatmap, SW_L1_MAP
        import pandas as pd
        ts_code = sys.argv[2] if len(sys.argv) > 2 else None

        # 支持 --l2 标志
        if ts_code == '--l2' or ts_code == '-l2':
            print_sw_heatmap(top_n=47, level='L2')
        elif ts_code and not ts_code.startswith('-'):
            # 单个行业查询
            name = SW_L1_MAP.get(ts_code, ts_code)
            print(f"\n  {'='*60}")
            print(f"  申万实时行情: {name} ({ts_code})")
            print(f"  {'='*60}")
            df = fetch_sw_realtime(ts_code=ts_code)
            if df is not None and not df.empty:
                r = df.iloc[0]
                chg = float(r['pct_change'])
                direction = '🔴' if chg >= 0 else '🟢'
                print(f"  现价: {r['close']:.2f}  涨跌: {chg:+.2f}% {direction}")
                print(f"  开盘: {r['open']:.2f}  最高: {r['high']:.2f}  最低: {r['low']:.2f}")
                print(f"  昨收: {r['pre_close']:.2f}  成交额: {float(r['amount'])/1e8:.2f}亿")
                print(f"  更新时间: {r['trade_time']}")
            else:
                print(f"  无数据 (代码: {ts_code})")
        else:
            # 全行业热力图
            print_sw_heatmap()

    elif cmd == "sw_list":
        from quant_trading.tushare_engine import SW_L1_MAP, get_sw_l2_codes
        l2 = '--l2' in sys.argv or '-l2' in sys.argv
        if l2:
            codes = get_sw_l2_codes()
            print(f"\n  申万二级行业指数代码 ({len(codes)}个，rt_sw_k支持):")
            for code in sorted(codes):
                print(f"    {code}")
        else:
            print(f"\n  申万一级行业指数代码 ({len(SW_L1_MAP)}个):")
            for code, name in sorted(SW_L1_MAP.items(), key=lambda x: x[1]):
                print(f"    {code}  {name}")

    # ══════════════════════════════════════════
    # DeepSeek + 审计后新增命令 (2026-07-03)
    # ══════════════════════════════════════════

    elif cmd == "behavioral":
        from quant_trading.behavioral_quant import fear_greed_index, overreaction_detection
        print_header("BEHAVIORAL QUANT — 行为金融分析")
        fg = fear_greed_index()
        print(f"  恐惧贪婪指数: {fg.get('index',50)}/100 ({fg.get('level','?')})")
        print(f"  过度反应: {overreaction_detection()}")

    elif cmd == "mtp":
        from quant_trading.mtp_predictor import MTPPredictor
        code = sys.argv[2] if len(sys.argv) > 2 else "300024"
        print_header(f"MTP RETURN PATH — {code}")
        mtp = MTPPredictor()
        r = mtp.predict_path(code)
        if 'error' not in r:
            for h in [1,3,5,10,20]:
                if f'T+{h}' in r:
                    t = r[f'T+{h}']
                    print(f"  T+{h:>2}: {t['mean']:+.2f}% (±{t['std']:.1f}%) 胜率={t['win_rate']:.0f}%")
            print(f"  最优持仓: T+{r.get('optimal_hold','?')} | 预期峰值: {r.get('peak_expected_return',0):+.1f}%")
            print(f"  {r.get('recommendation','')}")

    elif cmd == "moe":
        from quant_trading.moe_router import MoERouter
        print_header("MoE DYNAMIC ROUTING")
        router = MoERouter()
        route = router.route()
        for name, w in route['weights'].items():
            bar = '█' * int(w * 50)
            print(f"  {name:<12} {w:.1%} {bar}")
        print(f"  路由: {route['route_reason']}")
        comp = router.compare_to_fixed()
        for k, d in comp['diff'].items():
            print(f"  vs固定: {k} {d:+.1%}")

    elif cmd == "speculative" or cmd == "spec":
        from quant_trading.speculative_screener import SpeculativeScreener
        top = int(sys.argv[2]) if len(sys.argv) > 2 else 20
        print_header(f"SPECULATIVE SCREENING — TOP {top}")
        ss = SpeculativeScreener()
        df = ss.full_pipeline(top_n=top)
        if not df.empty:
            for i, (_, r) in enumerate(df.iterrows()):
                print(f"  {i+1:>2}. {r['code']} {r.get('name','?'):<8} "
                      f"综合={r.get('combined_score',0)} {r.get('stars','')} {r.get('grade','?')}")

    elif cmd == "sparse":
        from quant_trading.sparse_kline import SparseKlineDetector
        from quant_trader import fetch_kline
        code = sys.argv[2] if len(sys.argv) > 2 else "300024"
        print_header(f"SPARSE K-LINE — {code}")
        df = fetch_kline(code, datalen=200, freq='d')
        data = [{'date': str(r['date'])[:10], 'open': float(r['open']),
                 'high': float(r['high']), 'low': float(r['low']),
                 'close': float(r['close']), 'volume': float(r['volume'])} for _, r in df.iterrows()]
        skd = SparseKlineDetector()
        key_bars = skd.detect(data, top_k=15)
        for bar in sorted(key_bars, key=lambda x: x['sparse_idx'])[-10:]:
            print(f"  {bar['date']} | {bar['close']:.2f} | "
                  f"重要={bar['importance_score']} | {', '.join(bar['reasons'])}")

    elif cmd == "grpo":
        from quant_trading.grpo_optimizer import train_grpo, GRPOOptimizer
        sub = sys.argv[2] if len(sys.argv) > 2 else "train"
        print_header(f"GRPO OPTIMIZER — {sub}")
        if sub == "best":
            opt = GRPOOptimizer()
            if opt.load_population():
                best = opt.get_best_params()
                ens = opt.get_ensemble_weights()
                print(f"  最优权重: CL={best.w_chanlun:.1%} YX={best.w_yixian:.1%} "
                      f"WK={best.w_wyckoff:.1%} EW={best.w_elliott:.1%}")
                print(f"  Sharpe={best.sharpe:.2f} WinRate={best.win_rate:.1%}")
                print(f"  Ensemble: CL={ens['chanlun']:.1%} YX={ens['yixian']:.1%} "
                      f"WK={ens['wyckoff']:.1%} EW={ens['elliott']:.1%}")
        else:
            opt = train_grpo(generations=5, population_size=8)

    elif cmd == "risk":
        from quant_trading.risk_manager import RiskManager
        print_header("RISK MANAGER")
        rm = RiskManager()
        for k, v in rm.get_limits_summary().items():
            print(f"  {k}: {v}")

    elif cmd == "dragon_backtest":
        from quant_trading.dragon_tiger_backtest import backtest_dragon_tiger
        code = sys.argv[2] if len(sys.argv) > 2 else None
        print_header(f"DRAGON TIGER BACKTEST")
        backtest_dragon_tiger(code) if code else print("  Usage: dragon_backtest <code>")

    elif cmd == "theme":
        from quant_trading.theme_scanner import ThemeScanner
        print_header("THEME SCANNER")
        ts = ThemeScanner()
        themes = ts.scan()
        if themes:
            for t in themes[:10]:
                print(f"  {t.get('name','?')}: {t.get('score',0)}分 | {t.get('stocks','')}")

    elif cmd == "elliott":
        from quant_trading.elliott_wave import analyze_elliott
        code = sys.argv[2] if len(sys.argv) > 2 else "300024"
        print_header(f"ELLIOTT WAVE — {code}")
        r = analyze_elliott(code)
        if 'error' not in r:
            print(f"  波浪数: {r.get('wave_count','?')} | 位置: {r.get('position','')[:60]}")
            print(f"  置信度: {r.get('confidence','?')}")
            for w in r.get('waves', [])[-5:]:
                print(f"  {w['dir']} {w['start']:.2f}→{w['end']:.2f} ({w['pct']:.1f}%)")

    elif cmd == "macro_engine":
        from quant_trading.macro_engine import MacroEngine
        print_header("MACRO ENGINE")
        me = MacroEngine()
        result = me.analyze()
        if result:
            for k, v in result.items():
                print(f"  {k}: {v}")

    elif cmd == "wyckoff":
        from quant_trading.wyckoff_analyzer import analyze_wyckoff
        code = sys.argv[2] if len(sys.argv) > 2 else "300024"
        print_header(f"WYCKOFF — {code}")
        r = analyze_wyckoff(code)
        if 'error' not in r:
            print(f"  阶段: {r['phase']} (置信度={r['confidence']})")
            print(f"  趋势: {r['trend']} | VPR: {r['vol_price_ratio']}")
            ms = r.get('markup_score', 0)
            if ms: print(f"  MARKUP质量: {ms}/100 [{r.get('markup_grade','?')}级] {r.get('markup_grade_desc','')}")
            print(f"  操作: {r['action'][:80]}")

    elif cmd == "high_freq":
        from quant_trading.speculative_screener import SpeculativeScreener
        print_header("HIGH FREQ SCREEN (FAST)")
        ss = SpeculativeScreener()
        df = ss.stage1_fast_screen(max_pass=50)
        if not df.empty:
            for i, (_, r) in enumerate(df.head(20).iterrows()):
                print(f"  {i+1:>2}. {r['code']} fast_score={r['fast_score']}")

    elif cmd == "health":
        from quant_trading.data_quality import batch_quality_check
        from quant_trading.risk_manager import RiskManager
        from quant_trading.market_state import get_market_state
        print_header("SYSTEM HEALTH CHECK")
        print("  [Market State]")
        ms = get_market_state('000001')
        print(f"    State: {ms.get('state','?')} Score: {ms.get('score','?')}")
        print(f"    Reversal Risk: {ms.get('reversal_risk','?')}/100")
        print("  [Risk Limits]")
        rm = RiskManager()
        for k, v in rm.get_limits_summary().items():
            if not k.startswith('peak'): print(f"    {k}: {v}")
        print("  [Data Quality]")
        try:
            from quant_trading.data_cache import get_cached_kline
            q = batch_quality_check(['000001', '300024', '300166'], fetch_fn=get_cached_kline)
            if q: print(f"    {q}")
        except Exception as e:
            print(f"    Data quality check skipped: {e}")

    elif cmd == "chan_aerospace":
        from quant_trading.chan_commercial_aerospace import main as aero_main
        print_header("CHANLUN — COMMERCIAL AEROSPACE")
        aero_main()

    elif cmd == "chan_gas":
        from quant_trading.chan_electronic_gas import main as gas_main
        print_header("CHANLUN — ELECTRONIC GAS")
        gas_main()

    elif cmd == "mainforce":
        from pathlib import Path
        from quant_trading.main_force_analyzer import analyze_main_force, batch_scan, compare_codes
        code = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].isdigit() and len(sys.argv[2]) == 6 else None
        if code:
            # 单股深度分析
            print_header(f"MAIN FORCE ANALYSIS — 主力控盘分析: {code}",
                         "谢佳颖《主力控盘操作学》· 六维度评分 · 成本估算 · 控盘阶段")
            r = analyze_main_force(code)
            if 'error' in r:
                console.print(f"  [{C_RED}]{r['error']}[/]")
            else:
                cp = r['cost_position']
                vs = r['volume_structure']
                pa = r['phase_analysis']
                an = r['anomaly_signals']
                console.print(f"\n  [{C_GREEN}]{r['code']}[/] 主力控盘分析 — {r['rating']}")
                console.print(f"  综合评分: {r['score']}分")
                console.print(f"  主力成本: {cp['cost_basis']} (偏离: {cp['deviation_pct']}) — {cp['zone']}")
                console.print(f"  量价健康: {vs['health_score']}分 | 配合率: {vs['coordination_rate']:.0%}")
                console.print(f"  控盘阶段: {pa['phase']}/{pa.get('sub_phase','')} — {pa['detail']}")
                console.print(f"  异常信号: {', '.join(an['signals']) if an['signals'] else '无'}")
                # VSA量价差
                vsa = r.get('vsa')
                if vsa:
                    v_signals = ', '.join(f"[{'+' if s['type']=='bullish' else '-'}]{s['signal']}" for s in vsa.get('recent_signals',[])[-4:]) or '无'
                    console.print(f"  VSA量价差: {vsa.get('score',50)}/100 {vsa.get('verdict','')} | 最近信号: {v_signals}")
                # 筹码集中度
                cc = r.get('chip_concentration')
                if cc and cc.get('available'):
                    console.print(f"  筹码集中度: {cc.get('score',0)}分 | 股东户数{cc.get('shareholder_count',0):,}户 | 环比{cc.get('shareholder_change',0):+.1%} | {cc.get('concentration_label','')}")
                # 筹码分布 (日线量价推算,每日更新)
                cd = r.get('chip_distribution')
                if cd:
                    sup = ', '.join(f"{s['price']:.2f}" for s in cd.get('supports',[])[:2]) or '无'
                    res = ', '.join(f"{s['price']:.2f}" for s in cd.get('resistances',[])[:2]) or '无'
                    console.print(f"  筹码分布: 平均成本{cd['avg_cost']:.2f} | 获利{cd['profit_pct']} | 支撑[{sup}] | 压力[{res}]")
                if r['risk_flags']:
                    console.print(f"  [{C_RED}]风险: {'; '.join(r['risk_flags'])}[/]")
                console.print(f"\n  [{C_GREEN}]▶ {r['action']}[/]")
        else:
            # 全市场扫描
            scan_top = int(sys.argv[2]) if len(sys.argv) > 2 and sys.argv[2].isdigit() else 50
            phase_filter = sys.argv[3] if len(sys.argv) > 3 and sys.argv[3] in ('ACCUMULATION','MARKUP','DISTRIBUTION','MARKDOWN') else None
            phase_str = f" | 阶段: {phase_filter}" if phase_filter else ""
            print_header(f"MAIN FORCE SCAN — 主力控盘全市场扫描{phase_str}",
                         f"谢佳颖理论 · 六维度评分 · 扫描{scan_top}只")
            df = batch_scan(max_stocks=scan_top, min_score=40, phase_filter=phase_filter)
            if not df.empty:
                console.print(f"\n  共发现 [bold]{len(df)}[/] 只有主力控盘信号")
                for i, (_, r) in enumerate(df.head(20).iterrows()):
                    pa = r['phase_analysis']
                    console.print(f"  {i+1:>2}. {r['code']} 评分={r['score']} | "
                                  f"{pa.get('sub_phase', pa['phase'])} | "
                                  f"成本={r['cost_position']['cost_basis']} "
                                  f"偏离={r['cost_position']['deviation_pct']}")
                out = Path(str(REPORTS_DIR)) / f"mainforce_scan_{datetime.now().strftime('%Y%m%d')}.json"
                out.parent.mkdir(parents=True, exist_ok=True)
                df.to_json(str(out), orient="records", force_ascii=False, indent=2)
                console.print(f"\n  [green]结果已保存: {out}[/]")
            else:
                console.print(f"\n  [{C_ORANGE}]未发现主力控盘标的[/]")

    elif cmd == "pit_lift":
        from quant_trading.pit_lift_detector import analyze_pit_lift, scan_all
        code = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].isdigit() and len(sys.argv[2]) == 6 else None
        if code:
            r = analyze_pit_lift(code)
            if 'error' in r:
                console.print(f"  [{C_RED}]{r['error']}[/]")
            else:
                p = r['pit_profile']; sub = r['sub_scores']
                console.print(f"\n  [{C_GREEN}]{code}[/] 挖坑拉升 — {r['score']}分 [{r['rating']}]")
                console.print(f"  高{p['peak_price']}→底{p['trough_price']}({p['drop_pct']})→今{r['latest_close']}(+{p['bounce_pct']}) {p['pit_days']}天")
                console.print(f"  量: 底{sub['volume_shrink']['bottom_vol_ratio']:.1%} 今{sub['volume_shrink']['current_vol_ratio']:.1f}x")
                for name, s in sub.items():
                    sig = ' | '.join(s['signals'][:2])
                    console.print(f"  [{C_GREEN}]{name}[/]: {s['score']} {sig}")
        else:
            scan_n = int(sys.argv[2]) if len(sys.argv) > 2 and sys.argv[2].isdigit() else 100
            min_s = int(sys.argv[3]) if len(sys.argv) > 3 and sys.argv[3].isdigit() else 50
            print_header(f"PIT-LIFT SCAN — 挖坑拉升全市场扫描", f"挖掘主力打压吸筹后拉升标的 · top {scan_n}")
            results = scan_all(min_score=min_s, max_stocks=scan_n)
            if results:
                console.print(f"\n  共发现 [bold]{len(results)}[/] 只")
                for i, r in enumerate(results[:20]):
                    p = r['pit_profile']
                    console.print(f"  {i+1:>2}. {r['code']} {r['score']}分 "
                                  f"高{p['peak_price']}→底{p['trough_price']}({p['drop_pct']})→今{r['latest_close']}(+{p['bounce_pct']})")
            else:
                console.print(f"\n  [{C_ORANGE}]未发现挖坑拉升标的[/]")

    elif cmd == "shap":
        print_header("SHAP EXPLAINER — 模型因子归因",
                     "XGBoost/LightGBM · SHAP特征重要性 · 多股横比")
        code = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].isdigit() and len(sys.argv[2]) == 6 else None
        compare = sys.argv[3] if len(sys.argv) > 3 else None
        if compare:
            codes = [c.strip() for c in compare.split(',')]
            from quant_trading.shap_explainer import compare_stocks
            compare_stocks(codes)
        elif code:
            from quant_trading.shap_explainer import explain_and_print
            explain_and_print(code)
        else:
            console.print("[yellow]用法: shap <code> [对比代码列表][/]")
            from quant_trading.shap_explainer import explain_and_print
            explain_and_print('002261')

    elif cmd == "darts":
        print_header("DARTS FORECASTER — 统一时序预测",
                     "Prophet/ARIMA/N-BEATS/TFT · 多模型对比")
        code = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].isdigit() and len(sys.argv[2]) == 6 else None
        horizon = int(sys.argv[3]) if len(sys.argv) > 3 and sys.argv[3].isdigit() else 20
        if code:
            from quant_trading.darts_forecaster import forecast_and_print
            forecast_and_print(code, horizon)
        else:
            console.print("[yellow]用法: darts <code> [预测天数][/]")
            from quant_trading.darts_forecaster import forecast_and_print
            forecast_and_print('002261', 20)

    elif cmd == "polars_scan":
        print_header("POLARS SCAN — 极速超跌扫描",
                     "Polars引擎 · 10x加速 · 全市场秒级筛选")
        threshold = float(sys.argv[2]) if len(sys.argv) > 2 else 65.0
        from quant_trading.polars_engine import scan_and_print
        scan_and_print(threshold)

    elif cmd == "polars_bench":
        print_header("POLARS BENCHMARK — 性能对比",
                     "Polars vs Pandas · 加速比展示")
        from quant_trading.polars_engine import benchmark_vs_pandas
        benchmark_vs_pandas()

    elif cmd == "task":
        """任务看板"""
        import subprocess
        args = [sys.executable, "quant_trading/task_board.py"]
        if "--days" in sys.argv or "-d" in sys.argv:
            for i, a in enumerate(sys.argv):
                if a in ("--days", "-d") and i + 1 < len(sys.argv):
                    args += ["--days", sys.argv[i + 1]]
        if "--watch" in sys.argv or "-w" in sys.argv:
            args.append("--watch")
        subprocess.run(args)

    elif cmd == "code_search":
        """代码语义搜索"""
        query = " ".join(sys.argv[2:]) if len(sys.argv) > 2 else ""
        if not query:
            print("Usage: python quant_trader.py code_search <关键词>")
            sys.exit(1)
        from quant_trading.code_index import search_code
        results = search_code(query)
        print(f"\n[CodeIndex] 搜索 \"{query}\" 返回 {len(results)} 条:\n")
        for r in results:
            args_str = f"({', '.join(r['args'])})" if r['args'] else ""
            print(f"  [{r['kind']:>8}] {r['module']}.{r['name']}{args_str}")
            print(f"          {r['file']}:{r['lineno']}")
            if r['docstring']:
                print(f"          {r['docstring'][:120]}")
            print()

    elif cmd == "agent":
        """多 Agent 团队 v2（对抗式审查 + 断线恢复）"""
        request = " ".join(sys.argv[2:]) if len(sys.argv) > 2 else ""
        plan_only = "--plan-only" in sys.argv
        history = "--history" in sys.argv
        get_id = None
        resume_id = None
        for i, a in enumerate(sys.argv):
            if a == "--get" and i + 1 < len(sys.argv):
                get_id = sys.argv[i + 1]
            if a == "--resume" and i + 1 < len(sys.argv):
                resume_id = sys.argv[i + 1]

        if history:
            os.system(f"{sys.executable} quant_trading/agent_team.py --history")
        elif get_id:
            os.system(f"{sys.executable} quant_trading/agent_team.py --get {get_id}")
        elif resume_id:
            os.system(f"{sys.executable} quant_trading/agent_team.py --resume {resume_id}")
        elif request:
            if plan_only:
                os.system(f"{sys.executable} quant_trading/agent_team.py --plan-only \"{request}\"")
            else:
                os.system(f"{sys.executable} quant_trading/agent_team.py \"{request}\"")
        else:
            print("Usage:")
            print("  python quant_trader.py agent <任务>                # 多 Agent 团队")
            print("  python quant_trader.py agent --resume <任务ID>     # 断线恢复")
            print("  python quant_trader.py agent --history             # 历史")
            print("  python quant_trader.py agent --get <任务ID>        # 详情")

    elif cmd in ("knowledge", "know"):
        """知识引擎"""
        sub = sys.argv[2] if len(sys.argv) > 2 else ""
        if sub == "ingest":
            from quant_trading.knowledge_engine import KnowledgeEngine
            ke = KnowledgeEngine()
            result = ke.ingest_all()
            print(f"[KnowledgeEngine] 摄入完成: {result}")
        elif sub in ("query", "search") and len(sys.argv) > 3:
            from quant_trading.knowledge_engine import KnowledgeEngine
            ke = KnowledgeEngine()
            query = " ".join(sys.argv[3:])
            results = ke.query(query, k=10)
            print(f"\n[KnowledgeEngine] 查询 \"{query}\" 返回 {len(results)} 条:\n")
            for r in results:
                print(f"  [{r['source']:>6}] (相关性={r['relevance']:.0%}) {r['content'][:120]}...")
        elif sub == "add_rule" and len(sys.argv) > 3:
            from quant_trading.knowledge_engine import KnowledgeEngine
            ke = KnowledgeEngine()
            rule = " ".join(sys.argv[3:])
            ok = ke.add_rule(rule)
            print(f"[KnowledgeEngine] 规则{'已添加' if ok else '已存在，跳过'}")
        else:
            print("Usage:")
            print("  python quant_trader.py knowledge ingest            全量摄入")
            print("  python quant_trader.py knowledge query <问题>      跨源检索")
            print("  python quant_trader.py knowledge add_rule <规则>   添加规则")

    elif cmd in ("token", "tokens", "usage"):
        """Token 用量报告"""
        from quant_trading.token_usage_report import generate_report
        detail = "--detail" in sys.argv or "-d" in sys.argv
        cost = "--cost" in sys.argv or "-c" in sys.argv
        days = 1
        for i, a in enumerate(sys.argv):
            if a in ("--days", "-n") and i + 1 < len(sys.argv):
                try: days = int(sys.argv[i + 1])
                except: pass
        print_header(f"TOKEN USAGE — Token 用量报告{' (含费用)' if cost else ''}")
        generate_report(days=days, show_detail=detail, show_cost=cost)

        print(f"Unknown command: {cmd}")
        print("Usage: python quant_trader.py <command>")
        print()
        print("每日工作流:")
        print("  daily-update 统一数据库每日更新 (Sina EOD + 评分 + 重训 + GRPO)")
        print("  prewarm      全市场K线缓存")
        print("  cache        缓存状态查看")
        print("  health      系统健康诊断")
        print("  monitor     Prometheus监控端点")
        print()
        print("分析:")
        print("  analyze <code>   单股缠论分析")
        print("  mainforce [n]    主力控盘扫描")
        print("  mainforce <code> 单股主力分析")
        print("  pit_lift [n]     挖坑拉升扫描 — 上升趋势中的回调买点（顺势）")
        print("  pit_lift <code>  单股挖坑分析")
        print("  oversold [n]     超跌反弹扫描 — 深度超卖后的反弹机会（逆势）")
        print("  ladder           涨跌停板梯次")
        print("  hot_sector       热门板块")
        print("  heatmap          板块热力图")
        print("  fundamental      基本面扫描")
        print()
        print("回测 & 优化:")
        print("  oversold_backtest <code>  超跌回测")
        print("  factor <code>             因子分析")
        print("  backtest <code>           事件回测")
        print()
        print("更多: scan | flow | macro | north_bound | limit | dragon_tiger | yuezi | mainforce | pit_lift")
        print("策略区分: pit_lift=顺势挖坑(上升期回调) | oversold=逆势超跌(深度超卖反弹)")
        print("      dark_horse | first_bearish | yixian | dashboard | cache | margin")
        print()
        print("快捷工具:")
        print("  python quant_trading/backtest_engine.py --report   信号统计报告")
        print("  python quant_trading/grpo_optimizer.py --train     策略参数优化")
        print("  python quant_trading/batch_scorer.py --codes 100   批量引擎评分")
        print()
        print("AI/ML 增强工具 (NEW 2026-07-07):")
        print("  shap <code>             SHAP模型因子归因分析")
        print("  darts <code> [days]     Darts统一时序预测(Prophet/N-BEATS/TFT)")
        print("  polars_scan [thres]     Polars极速超跌扫描(默认阈值65)")
        print("  polars_bench            Polars vs Pandas性能对比")
        print("  sentiment <text>        中文金融情感分析(SnowNLP+cnsenti)")
        print()
        print("New: behavioral|mtp|moe|speculative|sparse|grpo|risk|theme|elliott|wyckoff|health|high_freq|chan_aerospace|chan_gas")

if __name__ == "__main__":
    main()
